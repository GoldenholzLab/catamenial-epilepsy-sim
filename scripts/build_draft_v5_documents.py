"""Build a clean Draft 5 template and Appendix S1 from structured source text.

This legacy clean-build path deliberately writes to distinct ``clean_rebuild``
filenames so it cannot overwrite the coauthor-edited Draft 5. For the
submission documents, use ``update_draft_v5_direct_alignment.py`` to patch the
coauthor-edited DOCX packages while preserving their Zotero fields and edits.
Citations emitted here are unique tokens for a later Zotero/Word field pass.
"""

from __future__ import annotations

import json
import math
import re
from pathlib import Path
from typing import Iterable, Sequence

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / ".codex_review" / "paper1_v3_072526" / "accepted.docx"
SUPPLEMENT_DIR = ROOT / "outputs" / "draft_v5_supplement"
OUT_DIR = ROOT / "outputs" / "epilepsia_submission"
MAIN_OUT = OUT_DIR / "draft_v5_clean_rebuild.docx"
APPENDIX_OUT = OUT_DIR / "draft_v5_appendix_clean_rebuild.docx"

TITLE = "False Positive Catamenial Epilepsy Classification: A Simulation Study"
RUNNING_TITLE = "False Positive Catamenial Classification"

CITE = {
    "herzog1997_phase": "[[CITE:Herzog1997-phase-labeling]]",
    "herzog1997_ratios": "[[CITE:Herzog1997-ratio-thresholds]]",
    "herzog1997_c3": "[[CITE:Herzog1997-C3-applicability]]",
    "herzog1997_discussion": "[[CITE:Herzog1997-discussion]]",
    "herzog1997_appendix": "[[CITE:Herzog1997-appendix]]",
    "herzog2004_intro": "[[CITE:Herzog2004-introduction]]",
    "herzog2004_methods": "[[CITE:Herzog2004-methods]]",
    "herzog2015": "[[CITE:Herzog2015]]",
    "voinescu2023": "[[CITE:Voinescu2023]]",
    "chocolates_methods": "[[CITE:GoldenholzWestover2023-methods]]",
    "chocolates_model": "[[CITE:GoldenholzWestover2023-model-check]]",
    "chocolates_appendix_1": "[[CITE:GoldenholzWestover2023-appendix-provenance]]",
    "chocolates_appendix_2": "[[CITE:GoldenholzWestover2023-appendix-calibration]]",
    "hormone_methods": "[[CITE:HORMONE-CYCLE-methods]]",
    "hormone_appendix": "[[CITE:HORMONE-CYCLE-appendix]]",
    "menstrual_targets_methods": "[[CITE:Li2023;Bull2019;Stricker2006-methods]]",
    "menstrual_targets_appendix": "[[CITE:Li2023;Bull2019;Stricker2006-appendix]]",
    "stress_methods": "[[CITE:STRESS2019-methods]]",
    "stress_appendix": "[[CITE:STRESS2019-appendix]]",
    "seizure_cycles": "[[CITE:Baud2018;Karoly2021;Karoly2020]]",
    "treatment": "[[CITE:Maguire2021;Herzog2012]]",
    "alshakhouri": "[[CITE:Alshakhouri2024Part2]]",
}


ABSTRACT = {
    "Objective": (
        "To quantify false-positive catamenial epilepsy classification when seizure and "
        "menstrual-cycle diaries are independent."
    ),
    "Methods": (
        "We simulated 100,000 synthetic participants for 36 months in healthy ovulatory and "
        "heterogeneous menstruating-age cohorts. CHOCOLATES generated seizure diaries and the "
        "new open-source HORMONE-CYCLE simulator generated independent menstrual and hormone "
        "diaries. We applied the phase definitions and pattern-specific seizure-frequency ratio "
        "thresholds introduced by Herzog and colleagues (the Herzog criteria) across calendar, "
        "complete-cycle, and full-diary windows. A negative-binomial C1/C2 analysis was a "
        "model-concordant statistical calibration check. An exploratory C3 negative-binomial "
        "check used the retained 1% daily audit sample."
    ),
    "Results": (
        "In 36-month strict-Herzog windows, windowed thresholds classified 11.2% of healthy "
        "ovulatory and 36.3% of heterogeneous-cohort participants under independence. In the "
        "heterogeneous cohort, the C1/C2 union was 11.8%, whereas C3 positivity occurred in "
        "37.0% of applicable windows. C1/C2 negative-binomial false-positive rates were 4.1% "
        "and 4.2%. The exploratory C3 model classified 6 of 201 classifiable audit-sample "
        "participants (3.0%; 95% Wilson confidence interval, 1.4%–6.4%). Three-month windowed "
        "Herzog rates were 41.3% and 51.2%."
    ),
    "Significance": (
        "False-positive classification can be common under independence, particularly when C3 "
        "logic is applied to heterogeneous cycles. Studies should report patterns separately and "
        "prespecify diary duration, minimum-data rules, denominator handling, phase labeling, and "
        "reproducibility. These rates are conditional on the specified simulators and are not "
        "empirical prevalence estimates."
    ),
}

MAIN_SECTIONS = [
    (
        "Introduction",
        [
            (
                "Catamenial epilepsy describes seizure exacerbation in relation to the menstrual "
                "cycle. Perimenstrual (C1), periovulatory (C2), and inadequate-luteal-phase (C3) "
                "patterns are clinically recognized, but reported prevalence varies with phase "
                "definitions, ratio thresholds, diary duration, and eligibility rules. Herzog and "
                "colleagues reported 39.1% using three complete cycles and a two-of-three-cycle "
                f"rule, and 44.2% has been reported among participants enrolled in the NIH "
                f"Progesterone Treatment Trial. {CITE['herzog2004_intro']} {CITE['herzog2015']}"
            ),
            (
                "Prospective diary classifications may disagree with participants’ prior reports "
                "of catamenial seizure exacerbation. Such discordance can reflect unrecognized "
                "diary-defined patterns, perceived patterns not confirmed by the selected "
                "quantitative criterion, or instability across limited observation windows; "
                f"neither source is an unquestioned reference standard. {CITE['voinescu2023']}"
            ),
            (
                "Seizures are temporally structured through clustering, circadian rhythms, and "
                "multidien rhythms. When these processes are independent of menstrual cycling, "
                "finite windows can still align by chance. Ratio definitions are especially "
                "vulnerable with sparse events: a small comparator count can yield a large or "
                f"undefined ratio. {CITE['seizure_cycles']}"
            ),
            (
                "We conducted a reproducible independence simulation to estimate false-positive "
                "catamenial classification across definitions and observation windows. The study "
                "evaluates specificity-like behavior only. We separately report C1/C2 and C3 and "
                "describe associations between simulated participant features and apparent "
                "classification. We also used a model-concordant negative-binomial C1/C2 analysis "
                "as a statistical calibration check. An exploratory, separately specified C3 "
                "negative-binomial analysis is reported in Appendix S1."
            ),
        ],
    ),
    (
        "Methods",
        [
            (
                "Study design. For each synthetic participant, seizure and menstrual/hormone "
                "diaries were generated independently for 36 months. Analyses were stratified "
                "by cohort and phase-labeling mode."
            ),
            (
                "Cohorts. The healthy ovulatory cohort included 50,000 participants aged "
                "18–45 years with ovulatory cycling enforced through the adapter. The "
                "heterogeneous menstruating-age cohort included 50,000 participants aged "
                "13.0–54.9 years and allowed anovulation and irregularity. Polycystic ovary "
                "syndrome, peri-menarche, perimenopause, and dysmenorrhea were sampled from "
                "investigator-configured stress-test rates because the hormone simulator did not "
                "contain a natural population prevalence sampler. These rates are assumptions, "
                "not population estimates (Appendix S1)."
            ),
            (
                "Diary simulators. CHOCOLATES generated seizure counts using a hierarchical "
                "gamma–Poisson/negative-binomial-like process with participant heterogeneity, "
                "clustering, multidien cycles, and interseizure constraints. "
                f"{CITE['chocolates_methods']} HORMONE-CYCLE is a new custom, data-driven simulator "
                "developed for the present project and released as an open-source community tool "
                f"at https://github.com/GoldenholzLab/catamenial-epilepsy-sim. {CITE['hormone_methods']} "
                "It samples participant traits and cycle realizations, then generates daily "
                "bleeding, ovulation, estradiol, and progesterone trajectories. Both simulators "
                "were designed in a data-driven manner to recapitulate statistical features "
                "observed within and across real patient populations. The exact targets, "
                "sampling levels, realized distributions, and assumption-by-assumption rationale "
                f"are reported in Appendix S1. {CITE['menstrual_targets_methods']} {CITE['stress_methods']}"
            ),
            (
                "Phase labeling and Herzog criteria. Strict Herzog phase labeling was the primary "
                "phase-definition scheme. Menstrual phase comprised forward days 1–3 and backward "
                "days −3 to −1; follicular phase, forward days 4–9; ovulatory phase, forward days "
                "≥10 through backward day −13; and luteal phase, backward days −12 to −4. Days "
                "outside these phases were unlabeled. A luteal-anchored sensitivity fixed the "
                "periovulatory interval at backward days −16 to −13 and assigned cycle-length "
                f"variation to the follicular phase. {CITE['herzog1997_phase']}"
            ),
            (
                "Windowed ratios were C1=ADSF(M)/ADSF(F+L) with threshold 1.69, "
                "C2=ADSF(O)/ADSF(F+L) with threshold 1.83, and "
                "C3=ADSF(O+L+M)/ADSF(F) with threshold 1.62 in applicable "
                f"inadequate-luteal-phase cycles/windows. These exact distributional inflection "
                f"values, rather than a rounded twofold rule, were used. {CITE['herzog1997_ratios']} "
                "A ratio was undefined when its numerator or comparator phase was absent or when "
                "both phase rates were zero; a positive numerator with zero comparator was "
                "treated as positive infinity."
            ),
            (
                "C3 applicability. In the heterogeneous cohort, the adapter assigned a "
                "simulator-generated inadequate-luteal-phase indicator to all nonovulatory cycles "
                "and cycles without a valid ovulation day or progesterone field. For other "
                "ovulatory cycles, it used the maximum progesterone concentration on days 5–9 "
                "after ovulation, falling back to all postovulation days if that window was "
                "absent, and designated values <5.0 ng/mL as inadequate. The C3 ratio then used "
                "only designated cycles. The 5.0-ng/mL cutoff follows the operational ovulatory "
                f"criterion in the original Herzog study; implementation details and sensitivity "
                f"limitations are in Appendix S1. {CITE['herzog1997_c3']}"
            ),
            (
                "Classification definitions. Exact Herzog 2004 required exactly three complete "
                "23–35-day cycles and the same positive pattern in at least two cycles. "
                f"{CITE['herzog2004_methods']} Windowed Herzog pooled the phase-specific ratios in each "
                "calendar, cycle-count, or full-diary window. The investigator-selected "
                "minimum-data sensitivity required at least four calendar months or six complete "
                "cycles and at least four seizure days. These pragmatic thresholds represent an "
                "investigator-selected minimum-data sensitivity intended to reduce classifications "
                "dominated by extremely sparse event counts; they were not biologically derived. "
                "Nearby event thresholds were examined without searching for an optimal cutoff "
                "(Appendix S1). Cycle reproducibility required the same positive pattern in at "
                "least two-thirds of eligible cycles and a positive pooled ratio."
            ),
            (
                "Model-concordant statistical calibration checks. Recurrent heterogeneous seizure "
                "counts are overdispersed, motivating negative-binomial rather than simple Poisson "
                f"models. {CITE['chocolates_model']} The prespecified C1/C2 check used within-participant "
                "daily-count generalized linear models with log link, menstrual and ovulatory "
                "indicators, and cycle fixed effects where estimable. One-sided Wald tests were "
                "Holm-adjusted across C1 and C2; positivity required adjusted P<.05 and the "
                "corresponding Herzog rate-ratio threshold. Because CHOCOLATES uses a "
                "negative-binomial-like count process, this is a model-concordant statistical "
                "calibration check—not an alternative clinical diagnostic method or independent "
                "simulator validation."
            ),
            (
                "A separately specified exploratory C3 model used the deterministic 1% daily "
                "audit sample. It restricted the heterogeneous cohort to complete cycles carrying "
                "the simulator-generated inadequate-luteal-phase indicator and contrasted combined "
                "O+L+M with F days using cycle fixed effects. Positivity required a one-sided "
                "P<.05 and rate ratio ≥1.62. This standalone exploratory family was not pooled "
                "with C1/C2 and therefore received no cross-family multiplicity adjustment "
                "(Appendix S1)."
            ),
            (
                "Outcomes and statistics. The primary outcome was the participant-window "
                "false-positive classification rate among classifiable windows under independence. "
                "We also report rates among all attempted windows, explicit indeterminate reasons, "
                "and mutually exclusive pattern categories. An indeterminate window was one in "
                "which an applicable rule could not return positive or negative because required "
                "phases, cycles, events, or denominator information were insufficient; "
                "nonapplicability was kept separate. Definition-by-window-by-cohort estimates with "
                "fewer than 1,000 classifiable participant-windows were flagged as unstable. "
                "Wilson intervals quantify Monte Carlo uncertainty in a simulated proportion and "
                "do not include model-form uncertainty. Ten thousand simulated studies of 30, 50, "
                "or 100 participants sampled one random valid 3-month window per participant."
            ),
            (
                "Ethics. All participants were synthetic; institutional review board approval, "
                "consent, and clinical-trial registration were not applicable."
            ),
        ],
    ),
    (
        "Results",
        [
            (
                "Simulation cohort. The completed run included 100,000 synthetic participants. "
                "The healthy ovulatory cohort had 100.0% ovulatory cycles by design; the "
                "heterogeneous cohort had 79.0% ovulatory cycles, greater within-participant "
                "cycle-length variability, and similar seizure burden (Table 1)."
            ),
            (
                "Full-diary classification. In strict-Herzog 36-month windows, windowed Herzog "
                "thresholds classified 11.2% of healthy ovulatory and 36.3% of heterogeneous "
                "participants under independence (Table 2). In the heterogeneous cohort, the "
                "C1/C2 union was 11.8%, similar to the healthy cohort’s 11.2%, whereas C3 "
                "positivity occurred in 37.0% of applicable windows. The excess was therefore "
                "primarily associated with the simulated inadequate-luteal-phase C3 mechanism "
                "(Figure 2 and Appendix S1)."
            ),
            (
                "Observation windows. Three-month windowed Herzog false-positive rates were 41.3% "
                "and 51.2% in the healthy and heterogeneous cohorts, respectively. Rates declined "
                "with longer monitoring but remained definition-dependent (Figure 1). Exact Herzog "
                "2004 applied to three complete cycles yielded 49.7% and 51.2% among classifiable "
                "windows, while many attempted windows were indeterminate. Appendix S1 reports C3 "
                "across every saved calendar and complete-cycle duration, minimum-data threshold "
                "sensitivities, and cumulative simulated C1, C2, and C3 ratio distributions."
            ),
            (
                "Calibration checks. The full-diary C1/C2 negative-binomial false-positive rate "
                "was 4.1% in the healthy cohort and 4.2% in the heterogeneous cohort, close to the "
                "prespecified 5% Type I error rate. In the 1% daily audit sample, 431 of 500 "
                "heterogeneous participants had a C3-applicable ratio window, but only 201 met the "
                "exploratory C3 model’s four-complete-ILP-cycle and four-seizure-day requirements. "
                "Six were positive (3.0%; 95% Wilson confidence interval, 1.4%–6.4%; 1.2% of all "
                "500 attempted participants)."
            ),
            (
                "Study-level and feature summaries. Across 10,000 simulated 3-month studies, "
                "apparent prevalence varied widely, particularly at n=30 (Table 3 and Figure 3). "
                "Associations with seizure burden, cycle variability, ovulatory fraction, age, and "
                "configured modifiers are reported in Appendix S1 as relationships between "
                "simulator inputs or realizations and simulated classifications, not clinical risk "
                "factors."
            ),
        ],
    ),
    (
        "Discussion",
        [
            (
                "This independence simulation shows that ratio-based catamenial classifications "
                "can arise frequently without simulated seizure–menstrual coupling. The magnitude "
                "depends on pattern, observation length, analyzability rules, and menstrual-cycle "
                "assumptions. The largest heterogeneous-cohort excess was C3-driven, whereas C1/C2 "
                "behavior was similar between cohorts."
            ),
            (
                "The model-concordant negative-binomial findings address implementation-level "
                "calibration under related model assumptions, not clinical usefulness or "
                "superiority. C1/C2 rates near 5% are expected for tests constructed around a 5% "
                "Type I error rate; ratio rules have no analogous guarantee. The exploratory C3 "
                "estimate was imprecise and used stricter eligibility in a 1% audit sample, so it "
                "should not be read as a direct replacement for the ratio-based C3 result. Its "
                "purpose is to show that a separately specified C3 contrast can be calibrated "
                "without implying independent validation of the simulator."
            ),
            (
                "Earlier studies established biologically plausible and clinically important C1, "
                f"C2, and C3 patterns, and treatment evidence remains limited and mixed. "
                f"{CITE['herzog1997_discussion']} {CITE['treatment']} The present study instead estimates a "
                "diagnostic background rate generated by chance alignment of independently "
                "simulated structured processes. It does not estimate true prevalence, "
                "sensitivity, positive predictive value, or treatment response."
            ),
            (
                "Several design implications follow. First, short-window ratio classifications "
                "require caution when seizure counts are sparse. Second, C1, C2, and C3 should be "
                "reported separately. Third, classifiable and all-attempted denominators should be "
                "shown together. Fourth, study comparisons should match the original design: the "
                "Herzog 2004 benchmark is aligned with the exact three-complete-cycle rule, not a "
                "random three-calendar-month window."
            ),
            (
                "Strict Herzog labeling was retained for historical comparability despite "
                "recognized limitations of fixed menstrual-phase assignment, particularly the "
                "expanding periovulatory interval in longer cycles. The luteal-anchored sensitivity "
                f"therefore tested a fixed periovulatory definition. {CITE['alshakhouri']}"
            ),
            (
                "The study is limited by simulator and adapter assumptions. CHOCOLATES and "
                "HORMONE-CYCLE were calibrated to published features, but summary-level calibration "
                "does not substitute for external validation using paired real diaries. The "
                "configured heterogeneous-cohort modifiers are stress-test assumptions, the "
                "5.0-ng/mL C3 rule operationalizes one historical definition, and no diary "
                "missingness was simulated. Monte Carlo intervals do not include model-form "
                "uncertainty. Medication changes and secular seizure-rate trends were outside the "
                "scope of this paper and could invalidate stationarity assumptions or alter "
                "chance-alignment rates. The study does not evaluate false negatives or adjudicated "
                "seizure types."
            ),
            (
                "Future work should compare these simulation-based rates with empirical-null "
                "distributions derived from longitudinal paired seizure and menstrual diaries, "
                "using complementary preregistered within-person time shifts and matched "
                "cross-person pairing that preserve temporal structure, missingness, and "
                "participant-level dependence. A clinical-facing calculator is deferred unless "
                "external validation and uncertainty decomposition support such use."
            ),
            (
                "In conclusion, apparent catamenial epilepsy classification can be common when "
                "structured seizure and menstrual-cycle diaries are independent. Diagnostic and "
                "interventional studies should prespecify diary length, phase labeling, "
                "minimum-data thresholds, zero-denominator handling, ovulatory-status handling, "
                "and reproducibility criteria before interpreting apparent menstrual clustering "
                "as hormone-linked seizure exacerbation."
            ),
        ],
    ),
]


KEY_POINTS = [
    "Windowed Herzog criteria classified 11.2% of healthy and 36.3% of heterogeneous simulated participants under independence in 36-month diaries.",
    "The heterogeneous-cohort excess was primarily C3/inadequate-luteal-phase driven; C1/C2 rates were similar between cohorts.",
    "Three-month observation windows produced substantially higher false-positive classification than longer monitoring.",
    "Pattern-specific reporting, transparent denominators, and prespecified minimum-data and phase-label rules are essential.",
]


def strip_citations(text: str) -> str:
    return re.sub(r"\[\[CITE:[^\]]+\]\]", "", text)


def word_count(text: str) -> int:
    cleaned = strip_citations(text)
    return len(re.findall(r"\b[\w]+(?:[’'-][\w]+)*\b", cleaned, flags=re.UNICODE))


def build_main_document() -> tuple[Document, dict[str, int]]:
    abstract_text = " ".join(ABSTRACT.values())
    main_text = " ".join(p for _, paragraphs in MAIN_SECTIONS for p in paragraphs)
    counts = {
        "title_chars": len(TITLE),
        "abstract_words": word_count(abstract_text),
        "main_text_words": word_count(main_text),
    }

    doc = _new_document(landscape=False)
    _add_title_page(doc, counts)
    doc.add_page_break()

    _add_heading(doc, "Abstract", level=1)
    for label, text in ABSTRACT.items():
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(text)
    p = doc.add_paragraph()
    p.add_run("Keywords: ").bold = True
    p.add_run("catamenial epilepsy; false-positive classification; menstrual cycle; seizure diary; simulation")

    _add_heading(doc, "Key Points", level=1)
    for point in KEY_POINTS:
        doc.add_paragraph(point, style="List Bullet")

    for heading, paragraphs in MAIN_SECTIONS:
        _add_heading(doc, heading, level=1)
        for text in paragraphs:
            doc.add_paragraph(text)

    _add_heading(doc, "Acknowledgments", level=1)
    doc.add_paragraph(
        "The authors thank the people whose menstrual and seizure diaries contributed to the "
        "published statistical targets used by the two simulators."
    )
    _add_heading(doc, "Study Funding", level=1)
    doc.add_paragraph(
        "This work was supported by NIH K23NS124656, NIH R21NS142800, and American Board of "
        "Psychiatry and Neurology support to DMG. The funders had no role in study design, data "
        "generation, analysis, interpretation, manuscript preparation, or the decision to submit."
    )
    _add_heading(doc, "Conflict of Interest", level=1)
    doc.add_paragraph(
        "DMG reports funding from NIH K23NS124656, NIH R21NS142800, and the American Board of "
        "Psychiatry and Neurology. SC, WTK, and RLS report no disclosures relevant to the "
        "manuscript. MBW disclosure confirmation is pending."
    )
    _add_heading(doc, "Author Contributions", level=1)
    doc.add_paragraph(
        "Daniel M. Goldenholz: Conceptualization, methodology, software, formal analysis, "
        "visualization, writing—original draft, writing—review and editing, funding acquisition. "
        "Wesley T. Kerr: Conceptualization, methodology, clinical interpretation, "
        "writing—review and editing. Sharon Chiang: Methodology, statistical interpretation, "
        "writing—review and editing. M. Brandon Westover: Methodology, software concept, clinical "
        "interpretation, writing—review and editing. Rachael L. Sumner: Methodology, reproductive "
        "neuroendocrinology interpretation, writing—review and editing."
    )
    _add_heading(doc, "Data and Code Availability", level=1)
    doc.add_paragraph(
        "Code, configuration, derived tabular outputs, and the run manifest are available at "
        "https://github.com/GoldenholzLab/catamenial-epilepsy-sim. The completed run used master "
        "seed 20260505; outputs/manifest.json records configuration values, file sizes, and SHA-256 "
        "checksums. The new Appendix S1 tables were derived without altering primary outputs."
    )
    _add_heading(doc, "Ethical Publication Statement", level=1)
    doc.add_paragraph(
        "We confirm that we have read the Journal’s position on issues involved in ethical "
        "publication and affirm that this report is consistent with those guidelines."
    )
    _add_heading(doc, "Supporting Information", level=1)
    doc.add_paragraph(
        "Appendix S1 contains the simulator source-to-parameter map, assumption review, C3 "
        "algorithm, cumulative simulated Herzog-ratio panels, window and minimum-data "
        "sensitivities, exploratory C3 negative-binomial calibration check, pattern decomposition, "
        "feature associations, and Supplementary Figures S1–S5."
    )

    _add_heading(doc, "References", level=1)
    doc.add_paragraph("[[ZOTERO_BIBLIOGRAPHY]]")

    doc.add_page_break()
    _add_heading(doc, "Tables", level=1)
    _add_main_table1(doc)
    doc.add_page_break()
    _add_main_table2(doc)
    doc.add_page_break()
    _add_main_table3(doc)

    doc.add_page_break()
    _add_heading(doc, "Figure Legends and Figures", level=1)
    _add_figure(
        doc,
        SUPPLEMENT_DIR / "fig1_false_positive_by_window_v5.png",
        (
            "Figure 1. False-positive rates by length of observation window. Bars show the "
            "percentage of classifiable calendar-window participant-windows classified by strict "
            "Herzog phase labeling under independence. The dashed 5% line is a visual reference to "
            "the prespecified regression Type I error rate; it is not an expected target for the "
            "Herzog ratio rules."
        ),
        width=Inches(6.7),
    )
    _add_figure(
        doc,
        SUPPLEMENT_DIR / "fig2_pattern_decomposition_v5.png",
        (
            "Figure 2. Pattern decomposition of false-positive catamenial classifications. "
            "Stacked bars show mutually exclusive participant-window categories in 36-month "
            "strict-Herzog windows. Percentages use all attempted windows. C3 is not applicable in "
            "the healthy ovulatory cohort; zero-height C3 segments there are not measured zero rates."
        ),
        width=Inches(6.7),
    )
    _add_figure(
        doc,
        ROOT / "outputs" / "fig3_study_prevalence_distribution_3month.png",
        (
            "Figure 3. False-positive prevalence in simulated 3-month studies. Distributions show "
            "apparent prevalence in 10,000 studies using strict-Herzog windowed thresholds for "
            "n=30, 50, and 100. Vertical lines mark the historical 39.1% and 44.2% benchmarks."
        ),
        width=Inches(6.7),
    )
    _set_core_properties(doc, TITLE)
    return doc, counts


def build_appendix_document() -> Document:
    doc = _new_document(landscape=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Appendix S1")
    run.bold = True
    run.font.size = Pt(18)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph(
        "This appendix is part of draft_v5. It reports only the completed simulations and "
        "derived analyses; no previously published participant data are included in the "
        "cumulative Herzog-ratio table."
    )

    _add_heading(doc, "S1. Reproducibility and provenance", level=1)
    doc.add_paragraph(
        "The primary simulation run (100,000 participants; 50,000 per cohort; 36 months) is "
        "immutable. Its manifest records master seed 20260505 and output hashes. Draft-v5 "
        "supplemental tables were derived from outputs/window_results.parquet, "
        "participant_summary.parquet, study_level_3month.parquet, and the deterministic 1% "
        "audit_daily_sample.parquet. HORMONE-CYCLE version 0.1.0 and the integrated analysis were "
        f"built at repository commit de6a716cfff753583a4831fcecdf5af0847d99b4. "
        f"{CITE['hormone_appendix']} {CITE['chocolates_appendix_1']} {CITE['stress_appendix']}"
    )
    doc.add_paragraph(
        "The completed manifest does not store a run-time Git commit, so the commit above is the "
        "checked-out source provenance associated with the retained outputs rather than a value "
        "embedded in the run manifest. This limitation should remain visible if the outputs are "
        "archived."
    )

    _add_heading(doc, "S2. Simulator inputs, realized features, and sources", level=1)
    doc.add_paragraph(
        "Both CHOCOLATES and HORMONE-CYCLE were designed to reproduce published statistical "
        "features within and across patient populations. CHOCOLATES encodes between-person seizure "
        "burden, overdispersion, clustering, rhythms, and interseizure constraints; HORMONE-CYCLE "
        "encodes between-person traits and within-person cycle-to-cycle realizations. The table "
        "distinguishes literature-calibrated targets, simulator-native behavior, "
        "investigator-configured stress assumptions, adapter-derived fields, and realized outputs. "
        f"{CITE['chocolates_appendix_2']} {CITE['menstrual_targets_appendix']}"
    )
    _add_parameter_table(doc)

    _add_heading(doc, "S3. Assumption-by-assumption review", level=1)
    assumption_rows = [
        ("Independent generators", "Defines the false-positive estimand; no participant-level seizure–hormone coupling was introduced.", "Does not model true catamenial effects or false negatives."),
        ("Direct calendar-day alignment", "Preserves the separately seeded seizure and hormone generator outputs without reordering.", "Independence follows from separate random streams; no artificial wraparound boundary is introduced."),
        ("36-month diaries", "Provides long follow-up and permits shorter nested window comparisons.", "Longer-term medication changes and secular trends were not modeled."),
        ("Healthy ovulatory cohort", "Provides a lower-heterogeneity reference in which C3 is not applicable.", "Not intended to represent all menstruating people with epilepsy."),
        ("Heterogeneous cohort", "Stress-tests irregularity, anovulation, and configured modifiers.", "Modifier rates are assumptions, not population prevalence estimates."),
        ("Strict Herzog phases", "Preserves historical comparability with the original criteria.", "The ovulatory interval expands in longer cycles; the luteal-anchored mode tests a fixed interval."),
        ("Unlabeled days", "Matches the historical phase partition and prevents implicit reassignment.", "Reduces observed denominators and can make short windows indeterminate."),
        ("Exact thresholds 1.69/1.83/1.62", "Uses the published distributional inflection points instead of rounded twofold shorthand.", "Threshold performance remains dependent on sparse numerator and comparator counts."),
        ("Zero-comparator handling", "A positive numerator over zero comparator is represented as +infinity; 0/0 is undefined.", "Makes the sparse-event vulnerability transparent rather than silently adding a continuity correction."),
        ("ILP threshold <5 ng/mL", "Operationalizes the ovulatory threshold used in the original Herzog phase study.", "A single operational cutoff does not exhaust biological definitions of luteal adequacy."),
        ("Nonovulatory cycles designated ILP", "Matches the implemented C3 adapter and the C3 concept of inadequate luteal-phase cycles.", "Makes C3 mechanically sensitive to anovulation prevalence and should not be called an independent validation."),
        ("Minimum-data rule", "Investigator-selected sensitivity to reduce classifications dominated by very sparse seizure counts.", "The 4-month/6-cycle and 4-seizure-day cutoffs are pragmatic rather than biologically derived."),
        ("NB C1/C2 calibration", "Checks rejection behavior with a model family concordant with the CHOCOLATES count process.", "Not an alternative clinical diagnostic method or independent simulator validation."),
        ("Exploratory NB C3", "Provides a separately specified F versus O+L+M contrast in ILP cycles.", "Uses only the saved 1% daily audit sample and stricter eligibility."),
        ("No missingness", "Keeps the primary estimand focused on finite-window stochastic alignment.", "Real diary missingness could affect phase and event denominators."),
        ("No medication/secular changes", "Keeps the null stationary and interpretable.", "Limits transportability to clinical diaries with treatment changes or disease evolution."),
    ]
    _add_table(
        doc,
        ["Assumption", "Rationale", "Limitation / implication"],
        assumption_rows,
        widths=[1.8, 4.3, 4.3],
        font_size=7.5,
    )

    _add_heading(doc, "S4. C3/ILP algorithm and applicability", level=1)
    doc.add_paragraph(
        "The adapter operates at the cycle level. A nonovulatory cycle, a cycle without a valid "
        "ovulation day, or a cycle without a usable progesterone field receives the "
        "simulator-generated inadequate-luteal-phase indicator. For another ovulatory cycle, the "
        "adapter identifies days 5–9 after ovulation and uses the maximum simulated progesterone "
        "concentration in that interval. If the interval is absent, it falls back to all "
        "postovulation days. A value below 5.0 ng/mL is designated inadequate. C3 then compares "
        "ADSF(O+L+M) with ADSF(F) using only designated cycles. This is the implemented algorithm, "
        "and it supersedes the earlier draft sentence suggesting that ordinary nonovulatory cycles "
        f"without an ILP indicator were excluded. {CITE['herzog1997_appendix']}"
    )
    doc.add_paragraph(
        "C3 applicability is window-specific: a pooled window is applicable when it includes at "
        "least one designated cycle, but its ratio can still be undefined because labeled "
        "numerator/comparator days or events are insufficient. The cycle-reproducibility rule "
        "requires at least the prespecified number of complete ILP cycles. The exploratory C3 "
        "negative-binomial model requires at least four complete designated cycles and four "
        "seizure days in labeled designated-cycle data. C3 is not applicable in the healthy "
        "ovulatory cohort."
    )

    _add_heading(doc, "S5. Cumulative simulated Herzog seizure-frequency ratios", level=1)
    doc.add_paragraph(
        "Supplementary Table S1. Cumulative distributions of simulated Herzog seizure-frequency "
        "ratios in strict three-complete-cycle windows under independence. Each panel has one row "
        "per participant-window. The first column is the Herzog ratio threshold. Positive infinity "
        "is defined and exceeds every finite threshold; undefined 0/0 or missing-phase ratios are "
        "excluded from the defined-ratio denominator. The ≥0 row therefore equals 100%. Only the "
        "study simulations are tabulated; the published trial data that motivated the layout are "
        "not reproduced or combined."
    )
    cumulative = pd.read_csv(SUPPLEMENT_DIR / "tableS1_cumulative_herzog_ratios.csv")
    _add_cumulative_panels(doc, cumulative)

    _add_heading(doc, "S6. C3 window-duration sensitivity", level=1)
    doc.add_paragraph(
        "Supplementary Table S2. Windowed Herzog C3-only false-positive classifications in the "
        "heterogeneous cohort. The healthy ovulatory cohort is not applicable and is omitted."
    )
    c3 = pd.read_csv(SUPPLEMENT_DIR / "tableS2_c3_window_sensitivity.csv")
    _add_c3_window_table(doc, c3)

    _add_heading(doc, "S7. Investigator-selected minimum-data sensitivity", level=1)
    doc.add_paragraph(
        "The minimum-data rule is a pragmatic investigator-selected sensitivity, not a Nyquist "
        "sampling requirement. The table varies the seizure-day minimum without selecting an "
        "optimal value and reports both classifiable and all-attempted rates. Calendar windows "
        "shorter than four months and cycle windows shorter than six cycles do not meet the "
        "duration component and are not included below."
    )
    minimum = pd.read_csv(SUPPLEMENT_DIR / "tableS3_minimum_data_sensitivity.csv")
    _add_minimum_sensitivity_table(doc, minimum)

    _add_heading(doc, "S8. Exploratory C3 negative-binomial calibration check", level=1)
    doc.add_paragraph(
        "The specification was frozen before inspecting results (docs/c3_nb_calibration_spec.md). "
        "The estimand is the within-participant daily seizure-count rate ratio for O+L+M versus F "
        "days in complete simulator-designated ILP cycles. The log-link negative-binomial model "
        "uses participant-full-diary method-of-moments dispersion and cycle fixed effects; robust "
        "Poisson is a reported fallback. A positive result requires a one-sided P<.05 and rate "
        "ratio ≥1.62. This is a standalone exploratory family, so it is not Holm-combined with the "
        "primary C1/C2 family."
    )
    c3_summary = pd.read_csv(SUPPLEMENT_DIR / "tableS6_c3_nb_exploratory_summary.csv").iloc[0]
    _add_table(
        doc,
        ["Attempted", "Ratio C3 applicable", "NB classifiable", "Positive", "Classifiable rate (95% Wilson CI)", "All-attempted rate"],
        [
            [
                int(c3_summary["n_attempted_audit_participants"]),
                int(c3_summary["n_ratio_c3_applicable"]),
                int(c3_summary["n_nb_classifiable"]),
                int(c3_summary["positives"]),
                f"{100*c3_summary['false_positive_rate_classifiable']:.1f}% "
                f"({100*c3_summary['wilson95_low']:.1f}%–{100*c3_summary['wilson95_high']:.1f}%)",
                f"{100*c3_summary['positive_rate_all_attempted']:.1f}%",
            ]
        ],
        font_size=8,
    )
    reasons = json.loads(c3_summary["reason_counts"])
    doc.add_paragraph(
        "Reason counts: "
        + "; ".join(f"{key.replace('_', ' ')}={value}" for key, value in reasons.items())
        + ". The narrower classifiable denominator reflects the model’s minimum four complete "
        "ILP cycles and four seizure days, not a discrepancy with pooled-ratio applicability."
    )

    _add_heading(doc, "S9. Mutually exclusive pattern decomposition", level=1)
    _add_pattern_table(doc)

    _add_heading(doc, "S10. Simulator-feature associations", level=1)
    doc.add_paragraph(
        "Supplementary Table S4 (machine-readable CSV) reports stratified apparent-classification "
        "rates and 95% Wilson intervals across prespecified quintiles of age, seizure burden, "
        "cycle length and variability, ovulatory fraction, and dominant seizure-cycle period, plus "
        "configured binary modifiers. These are descriptive associations between simulator "
        "inputs/realizations and simulated outcomes. They are not clinical patient risk factors, "
        "causal effects, or a patient-level prediction model. A clinical-facing calculator was "
        "not built."
    )

    _add_heading(doc, "S11. Realized distribution and association figures", level=1)
    figure_specs = [
        (
            "figS1_seizure_process_distributions.png",
            "Supplementary Figure S1. Realized seizure-process distributions. Panels show seizures/month, seizure days/month, daily count variance/mean, and a next-day clustering propensity in the completed run or retained audit sample.",
        ),
        (
            "figS2_seizure_rhythm_distributions.png",
            "Supplementary Figure S2. Realized seizure-rhythm distributions. Panels show generator-reported dominant period, audit-sample spectral period and amplitude, and latent monthly seizure burden.",
        ),
        (
            "figS3_menstrual_cycle_distributions.png",
            "Supplementary Figure S3. Realized menstrual-cycle distributions. Panels show mean cycle length, within-person cycle-length SD, ovulatory fraction, and ILP fraction. ILP is displayed only for the heterogeneous audit sample because C3 is not applicable in the healthy cohort.",
        ),
        (
            "figS4_age_and_modifier_distributions.png",
            "Supplementary Figure S4. Age and configured modifiers. Modifier frequencies are investigator-configured stress-test assumptions, not population estimates.",
        ),
        (
            "figS5_simulated_classification_associations.png",
            "Supplementary Figure S5. Associations with simulated apparent classification across within-cohort feature quintiles. The display is descriptive and not a clinical risk model.",
        ),
    ]
    for index, (filename, caption) in enumerate(figure_specs):
        # Figure S1 fills its page exactly; an explicit break after it is
        # pushed onto the next page by Word and creates an empty page.
        # Let Figure S2 flow naturally, then force subsequent figures.
        if index > 1:
            doc.add_page_break()
        _add_figure(doc, SUPPLEMENT_DIR / filename, caption, width=Inches(8.8))

    doc.add_page_break()
    _add_heading(doc, "S12. Interpretation of simulator calibration", level=1)
    doc.add_paragraph(
        "The figure set intentionally distinguishes completed-run distributions from external "
        "empirical validation. Where participant-level empirical microdata were unavailable, the "
        "figures show simulated distributions and the source table identifies the published "
        "summary targets. They should be described as distributional calibration outputs. External "
        "validation requires paired real seizure and menstrual diaries analyzed with aligned "
        "definitions, participant-level resampling, and complementary within-person shift and "
        "matched cross-person null designs."
    )

    _add_heading(doc, "References", level=1)
    doc.add_paragraph("[[ZOTERO_BIBLIOGRAPHY]]")
    _set_core_properties(doc, f"Appendix S1 — {TITLE}")
    return doc


def _new_document(landscape: bool) -> Document:
    if SOURCE.exists():
        doc = Document(SOURCE)
        body = doc._element.body
        for child in list(body):
            if child.tag != qn("w:sectPr"):
                body.remove(child)
    else:
        doc = Document()
    section = doc.sections[0]
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
    else:
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
    _configure_styles(doc)
    _add_page_number(section)
    return doc


def _configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in [
        ("Title", 18, "1F4E79"),
        ("Heading 1", 14, "1F4E79"),
        ("Heading 2", 11.5, "2F5597"),
        ("Heading 3", 10.5, "2F5597"),
    ]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)
    if "Table Note" not in [style.name for style in doc.styles]:
        style = doc.styles.add_style("Table Note", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Times New Roman"
        style.font.size = Pt(8)
        style.paragraph_format.space_after = Pt(4)


def _add_title_page(doc: Document, counts: dict[str, int]) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Title"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(TITLE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Research Article").bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(
        "Daniel M. Goldenholz, MD, PhD¹; Wesley T. Kerr, MD, PhD²˒³˒⁴; "
        "Sharon Chiang, MD, PhD⁵; M. Brandon Westover, MD, PhD⁶; "
        "Rachael L. Sumner, PhD⁷˒⁸"
    )
    affiliations = [
        "¹ Department of Neurology, Beth Israel Deaconess Medical Center, Harvard Medical School, Boston, Massachusetts, USA",
        "² Department of Neurology, University of Pittsburgh, Pittsburgh, Pennsylvania, USA",
        "³ Department of Biomedical Informatics, University of Pittsburgh, Pittsburgh, Pennsylvania, USA",
        "⁴ Department of Psychiatry, University of Pittsburgh, Pittsburgh, Pennsylvania, USA",
        "⁵ Department of Neurology and Weill Institute for Neurosciences, University of California, San Francisco, San Francisco, California, USA",
        "⁶ Department of Neurology and Neurological Sciences, Stanford University School of Medicine, Stanford, California, USA",
        "⁷ Department of Biomedicine and Medical Diagnostics, Auckland University of Technology, Auckland, New Zealand",
        "⁸ School of Pharmacy, University of Auckland, Auckland, New Zealand",
    ]
    for affiliation in affiliations:
        doc.add_paragraph(affiliation)
    p = doc.add_paragraph()
    p.add_run("Corresponding author: ").bold = True
    p.add_run(
        "Daniel M. Goldenholz, MD, PhD, Department of Neurology, Beth Israel Deaconess "
        "Medical Center, 330 Brookline Avenue, Boston, MA 02215, USA. Email: author confirmation required."
    )
    metadata = [
        ("Running title", RUNNING_TITLE),
        ("Title character count (including spaces)", str(counts["title_chars"])),
        ("Abstract word count", str(counts["abstract_words"])),
        ("Main-text word count (Introduction through Discussion)", str(counts["main_text_words"])),
        ("References", "16"),
        ("Main tables / figures", "3 / 3 (6 combined)"),
        ("Supporting information", "Appendix S1; Tables S1–S4 and Figures S1–S5"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(value)


def _add_main_table1(doc: Document) -> None:
    doc.add_paragraph("Table 1. Simulated cohort characteristics.", style="Heading 2")
    headers = [
        "Cohort",
        "N",
        "Age, y\nMean",
        "Age, y\nSD",
        "Cycle length, d\nMean",
        "Cycle length, d\nSD",
        "Ovulatory cycles",
        "Seizure days/month",
        "Seizures/month",
    ]
    rows = [
        ["Healthy ovulatory", "50,000", "31.5", "7.8", "29.15", "3.39", "100.0%", "2.46", "6.83"],
        ["Heterogeneous menstruating-age", "50,000", "34.0", "12.1", "30.92", "4.74", "79.0%", "2.45", "6.80"],
    ]
    _add_table(doc, headers, rows, font_size=8)
    _table_note(
        doc,
        "Values are cohort-level means from the completed run. Cycle-length SD is the mean "
        "within-participant cycle-length SD. N denotes simulated participants; y, years; d, days.",
    )


def _add_main_table2(doc: Document) -> None:
    doc.add_paragraph("Table 2. Participant-window false-positive results.", style="Heading 2")
    headers = [
        "Cohort",
        "Analysis window",
        "Definition",
        "Classifiable / attempted",
        "Classifiable, % (95% Wilson CI)",
        "All attempted, %",
        "Indeterminate, %",
    ]
    rows = [
        ["Healthy ovulatory", "36-month full diary", "Windowed Herzog thresholds", "49,605 / 50,000", "11.2 (11.0–11.5)", "11.2", "0.8"],
        ["", "", "Windowed Herzog C1/C2 union", "49,605 / 50,000", "11.2 (11.0–11.5)", "11.2", "0.8"],
        ["", "", "Herzog C1/C2 + ≥4 seizure-day minimum", "48,359 / 50,000", "10.2 (9.9–10.5)", "9.9", "3.3"],
        ["", "", "Negative-binomial C1/C2 calibration", "48,359 / 50,000", "4.1 (3.9–4.2)", "3.9", "3.3"],
        ["", "3-month calendar", "Windowed Herzog thresholds", "45,280 / 50,000", "41.3 (40.8–41.7)", "37.4", "9.4"],
        ["", "3 complete cycles", "Exact Herzog 2004, any pattern", "22,996 / 50,000", "49.7 (49.0–50.3)", "22.8", "54.0"],
        ["Heterogeneous menstruating-age", "36-month full diary", "Windowed Herzog thresholds", "49,587 / 50,000", "36.3 (35.8–36.7)", "36.0", "0.8"],
        ["", "", "Windowed Herzog C1/C2 union", "49,587 / 50,000", "11.8 (11.5–12.1)", "11.7", "0.8"],
        ["", "", "Windowed Herzog C3 only", "38,989 / 50,000", "37.0 (36.5–37.5)", "28.8", "22.0"],
        ["", "", "Herzog C1/C2 + ≥4 seizure-day minimum", "48,333 / 50,000", "10.7 (10.5–11.0)", "10.4", "3.3"],
        ["", "", "Negative-binomial C1/C2 calibration", "48,333 / 50,000", "4.2 (4.0–4.4)", "4.1", "3.3"],
        ["", "3-month calendar", "Windowed Herzog thresholds", "45,297 / 50,000", "51.2 (50.7–51.6)", "46.4", "9.4"],
        ["", "3 complete cycles", "Exact Herzog 2004, any pattern", "16,757 / 50,000", "51.2 (50.4–51.9)", "17.2", "66.5"],
    ]
    _add_table(doc, headers, rows, font_size=7.3)
    _table_note(
        doc,
        "All rows use strict Herzog phase labeling. C3 is not applicable in the healthy "
        "ovulatory cohort, so the prior healthy C3 row is omitted. Indeterminate means that an "
        "applicable rule could not return positive or negative because required phases, cycles, "
        "events, or denominator information were insufficient; nonapplicability is distinct. "
        "Definition-by-window-by-cohort estimates based on fewer than 1,000 classifiable "
        "participant-windows were flagged as unstable.",
    )


def _add_main_table3(doc: Document) -> None:
    doc.add_paragraph("Table 3. Study-level Monte Carlo summaries for 3-month windows.", style="Heading 2")
    headers = [
        "Cohort",
        "N/study",
        "Definition",
        "Mean apparent prevalence",
        "2.5th–97.5th percentiles",
        "Pr(apparent prevalence ≥39.1%)",
        "Pr(apparent prevalence ≥44.2%)",
    ]
    rows = [
        ["Healthy ovulatory", "30", "Windowed Herzog thresholds", "37.5%", "20.0%–53.3%", "45.4%", "20.1%"],
        ["", "30", "Windowed Herzog C1/C2 union", "37.5%", "20.0%–53.3%", "45.4%", "20.1%"],
        ["", "50", "Windowed Herzog thresholds", "37.3%", "24.0%–50.0%", "39.7%", "13.0%"],
        ["", "50", "Windowed Herzog C1/C2 union", "37.3%", "24.0%–50.0%", "39.7%", "13.0%"],
        ["", "100", "Windowed Herzog thresholds", "37.5%", "28.0%–47.0%", "33.8%", "7.1%"],
        ["", "100", "Windowed Herzog C1/C2 union", "37.5%", "28.0%–47.0%", "33.8%", "7.1%"],
        ["Heterogeneous menstruating-age", "30", "Windowed Herzog thresholds", "46.3%", "30.0%–63.3%", "80.5%", "55.9%"],
        ["", "30", "Windowed Herzog C1/C2 union", "38.2%", "20.0%–56.7%", "48.9%", "22.4%"],
        ["", "50", "Windowed Herzog thresholds", "46.1%", "32.0%–60.0%", "84.5%", "55.9%"],
        ["", "50", "Windowed Herzog C1/C2 union", "38.1%", "26.0%–52.0%", "43.9%", "15.3%"],
        ["", "100", "Windowed Herzog thresholds", "46.0%", "36.0%–56.0%", "90.8%", "62.2%"],
        ["", "100", "Windowed Herzog C1/C2 union", "38.0%", "29.0%–47.0%", "38.2%", "8.8%"],
    ]
    _add_table(doc, headers, rows, font_size=7.5)
    _table_note(
        doc,
        "Each row summarizes 10,000 simulated studies using one random valid 3-month window per "
        "selected participant. Ranges are empirical percentiles across simulated studies, not "
        "confidence intervals. The 39.1% benchmark is aligned with the Herzog 2004 "
        "three-cycle/two-of-three rule; 44.2% is the reported NIH Progesterone Trial proportion.",
    )


def _add_parameter_table(doc: Document) -> None:
    source = pd.read_csv(SUPPLEMENT_DIR / "tableS5_simulator_parameters_and_assumptions.csv")
    rows = []
    for row in source.itertuples(index=False):
        rows.append(
            [
                row.domain,
                row.parameter,
                row.cohort,
                str(row.setting_or_sampling_distribution),
                row.sampling_level,
                row.source_or_rationale,
                row.realized_validation_target,
            ]
        )
    _add_table(
        doc,
        ["Domain", "Parameter", "Cohort", "Setting / sampling distribution", "Level", "Source / rationale", "Realized target"],
        rows,
        widths=[0.8, 1.6, 1.0, 2.5, 1.1, 2.4, 1.6],
        font_size=6.6,
    )
    _table_note(
        doc,
        "The full machine-readable table additionally records code/config fields, sensitivity "
        "status, and the run version/commit. Values labeled investigator-configured are "
        "stress-test assumptions. The completed output did not retain CHOCOLATES cluster status "
        "or generator amplitude fields; audit-sample proxies are reported instead.",
    )


def _add_cumulative_panels(doc: Document, data: pd.DataFrame) -> None:
    for pattern, title in [
        ("C1", "Panel A — C1 = ADSF(M)/ADSF(F+L)"),
        ("C2", "Panel B — C2 = ADSF(O)/ADSF(F+L)"),
        ("C3", "Panel C — C3 = ADSF(O+L+M)/ADSF(F)"),
    ]:
        doc.add_paragraph(title, style="Heading 2")
        panel = data[data["pattern"] == pattern]
        if pattern in {"C1", "C2"}:
            healthy = panel[panel["cohort"] == "healthy_ovulatory"].set_index("threshold")
            population = panel[panel["cohort"] == "population"].set_index("threshold")
            rows = []
            for threshold in THRESHOLD_ORDER[pattern]:
                h = healthy.loc[float(threshold)]
                p = population.loc[float(threshold)]
                rows.append(
                    [
                        _fmt_threshold(threshold),
                        f"{int(h.n_at_or_above):,}",
                        f"{h.pct_defined_at_or_above:.2f}",
                        f"{int(p.n_at_or_above):,}",
                        f"{p.pct_defined_at_or_above:.2f}",
                    ]
                )
            _add_table(
                doc,
                ["Herzog ratio threshold (≥)", "Healthy n", "Healthy %", "Heterogeneous n", "Heterogeneous %"],
                rows,
                font_size=7.5,
            )
            h0 = healthy.iloc[0]
            p0 = population.iloc[0]
            _table_note(
                doc,
                f"Healthy attempted={int(h0.n_attempted):,}, defined={int(h0.n_defined):,}, "
                f"undefined={int(h0.n_undefined):,}; heterogeneous attempted={int(p0.n_attempted):,}, "
                f"defined={int(p0.n_defined):,}, undefined={int(p0.n_undefined):,}. "
                "Percentages use the defined-ratio denominator.",
            )
        else:
            population = panel[panel["cohort"] == "population"].set_index("threshold")
            rows = []
            for threshold in THRESHOLD_ORDER[pattern]:
                p = population.loc[float(threshold)]
                rows.append(
                    [
                        _fmt_threshold(threshold),
                        f"{int(p.n_at_or_above):,}",
                        f"{p.pct_defined_at_or_above:.2f}",
                    ]
                )
            _add_table(
                doc,
                ["Herzog ratio threshold (≥)", "Heterogeneous n", "Heterogeneous %"],
                rows,
                font_size=7.5,
            )
            p0 = population.iloc[0]
            _table_note(
                doc,
                f"Attempted={int(p0.n_attempted):,}; applicable={int(p0.n_applicable):,}; "
                f"defined={int(p0.n_defined):,}; undefined={int(p0.n_undefined):,}. "
                "The healthy cohort is not applicable and is not shown.",
            )


THRESHOLD_ORDER = {
    "C1": [0, 1, 1.69, 2, 3, 4, 5, 6, 7, 8, 9, 10],
    "C2": [0, 1, 1.83, 2, 3, 4, 5, 6, 7, 8, 9, 10],
    "C3": [0, 1, 1.62, 2, 3, 4, 5, 6, 7, 8, 9, 10],
}


def _add_c3_window_table(doc: Document, data: pd.DataFrame) -> None:
    order = {"calendar": 0, "cycle": 1, "full": 2}
    data = data.copy()
    data["_order"] = data["window_type"].map(order)
    data["_value"] = pd.to_numeric(data["window_value"], errors="coerce").fillna(999)
    data = data.sort_values(["_order", "_value"])
    rows = []
    for row in data.itertuples(index=False):
        if row.window_type == "calendar":
            label = f"{int(float(row.window_value))} month" + ("s" if int(float(row.window_value)) != 1 else "")
        elif row.window_type == "cycle":
            label = f"{int(float(row.window_value))} complete cycles"
        else:
            label = "36-month full diary"
        rows.append(
            [
                label,
                f"{int(row.n_applicable):,}",
                f"{int(row.n_classifiable):,}",
                f"{int(row.positives):,}",
                f"{100*row.false_positive_rate_classifiable:.1f}% "
                f"({100*row.wilson95_low:.1f}%–{100*row.wilson95_high:.1f}%)",
                f"{100*row.positive_rate_all_attempted:.1f}%",
            ]
        )
    _add_table(
        doc,
        ["Window", "Applicable", "Classifiable", "Positive", "Classifiable rate (95% Wilson CI)", "All-attempted rate"],
        rows,
        font_size=7.2,
    )
    _table_note(
        doc,
        "Each window has 50,000 attempted heterogeneous-cohort participant-windows. "
        "Indeterminate reason counts are retained in tableS2_c3_window_sensitivity.csv.",
    )


def _add_minimum_sensitivity_table(doc: Document, data: pd.DataFrame) -> None:
    selected = data[
        (
            (data["window_type"] == "calendar")
            & pd.to_numeric(data["window_value"], errors="coerce").isin([4, 6, 12, 36])
        )
        | (data["window_type"] == "full")
    ].copy()
    selected["_value"] = pd.to_numeric(selected["window_value"], errors="coerce").fillna(36)
    selected = selected.sort_values(["cohort", "_value", "min_seizure_days"])
    rows = []
    for row in selected.itertuples(index=False):
        if row.window_type == "full":
            window = "36-month full diary"
        else:
            window = f"{int(float(row.window_value))} months"
        rows.append(
            [
                "Healthy" if row.cohort == "healthy_ovulatory" else "Heterogeneous",
                window,
                int(row.min_seizure_days),
                f"{int(row.n_classifiable):,}",
                f"{100*row.false_positive_rate_classifiable:.1f}%" if not pd.isna(row.false_positive_rate_classifiable) else "—",
                f"{100*row.positive_rate_all_attempted:.1f}%",
            ]
        )
    _add_table(
        doc,
        ["Cohort", "Window", "Minimum seizure days", "Classifiable", "Classifiable rate", "All-attempted rate"],
        rows,
        font_size=6.8,
    )


def _add_pattern_table(doc: Document) -> None:
    headers = [
        "Cohort / definition",
        "None",
        "Indeterminate",
        "C1 only",
        "C2 only",
        "C1+C2",
        "C3 only",
        "C3 with C1 and/or C2",
    ]
    rows = [
        ["Healthy / windowed Herzog", "88.1%", "0.8%", "5.6%", "3.4%", "2.1%", "N/A", "N/A"],
        ["Healthy / minimum-data", "86.8%", "3.3%", "5.1%", "2.9%", "1.9%", "N/A", "N/A"],
        ["Heterogeneous / windowed Herzog", "63.2%", "0.8%", "3.8%", "2.1%", "1.2%", "24.3%", "4.6%"],
        ["Heterogeneous / minimum-data", "62.3%", "3.3%", "3.5%", "1.7%", "1.0%", "24.0%", "4.2%"],
    ]
    _add_table(doc, headers, rows, font_size=7)
    _table_note(
        doc,
        "Supplementary Table S3. Mutually exclusive participant-window categories in 36-month "
        "strict-Herzog windows. Percentages use all attempted windows and sum to 100% within "
        "rounding after treating C3 as not applicable in the healthy cohort.",
    )


def _add_heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def _add_table(
    doc: Document,
    headers: Sequence[str],
    rows: Iterable[Sequence[object]],
    widths: Sequence[float] | None = None,
    font_size: float = 8,
) -> None:
    rows = list(rows)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = widths is None
    header = table.rows[0]
    _repeat_table_header(header)
    for index, text in enumerate(headers):
        cell = header.cells[index]
        cell.text = str(text)
        _shade_cell(cell, "D9EAF7")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths:
            cell.width = Inches(widths[index])
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(font_size)
    for row_values in rows:
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
        for index, value in enumerate(row_values):
            cell = row.cells[index]
            cell.text = str(value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths:
                cell.width = Inches(widths[index])
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.keep_together = True
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def _table_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Table Note")
    p.paragraph_format.keep_with_next = False
    p.add_run("Note. ").bold = True
    p.add_run(text)


def _add_figure(doc: Document, path: Path, caption: str, width: Inches) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    run = p.add_run(caption)
    run.bold = True
    run.font.size = Pt(9)
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=width)
    else:
        doc.add_paragraph(f"[Missing figure: {path}]")


def _repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _add_page_number(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def _set_core_properties(doc: Document, title: str) -> None:
    doc.core_properties.title = title
    doc.core_properties.subject = "False-positive catamenial epilepsy classification under simulated independence"
    doc.core_properties.author = "Daniel M. Goldenholz et al."
    doc.core_properties.keywords = "catamenial epilepsy, seizure diary, menstrual cycle, simulation"


def _fmt_threshold(value: float) -> str:
    return str(int(value)) if float(value).is_integer() else f"{value:.2f}"


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    main_doc, counts = build_main_document()
    main_doc.save(MAIN_OUT)
    appendix_doc = build_appendix_document()
    appendix_doc.save(APPENDIX_OUT)
    print(json.dumps({"main": str(MAIN_OUT), "appendix": str(APPENDIX_OUT), "counts": counts}, indent=2))


if __name__ == "__main__":
    main()
