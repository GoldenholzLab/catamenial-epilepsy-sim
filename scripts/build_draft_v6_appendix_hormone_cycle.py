"""Expand the repaired draft-v6 appendix with a complete HORMONE-CYCLE section.

The user-repaired source is treated as immutable.  This script:

* inserts a detailed, implementation-matched HORMONE-CYCLE methods section;
* adds workflow and calibration figures plus complete calibration tables;
* preserves the existing Zotero fields and adds live Zotero citation fields;
* defines abbreviations within each appendix section; and
* constrains every table to the 9-inch text block of a landscape page with
  one-inch margins.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import math
import re
import urllib.request
import zipfile
from copy import deepcopy
from pathlib import Path
from typing import Iterable, Sequence

import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SOURCE = (
    ROOT / "outputs" / "epilepsia_submission" / "draft_v6_appendix_s1.docx"
)
DEFAULT_OUTPUT = (
    ROOT
    / "outputs"
    / "epilepsia_submission"
    / "draft_v6_appendix_s1_hormone_cycle_expanded.docx"
)
DEFAULT_VALIDATION = (
    ROOT
    / ".codex_review"
    / "draft_v6_appendix_s1"
    / "hormone_cycle_validation_current.json"
)
DEFAULT_ASSET_DIR = (
    ROOT / ".codex_review" / "draft_v6_appendix_s1" / "generated_assets"
)

ZOTERO_LIBRARY_ID = 1538114
ZOTERO_KEYS = {
    "hormone_cycle": "BBSUV5B4",
    "li": "KCJ4JNA6",
    "bull": "BZHW34SL",
    "stricker": "RQGUHW73",
    "fraser": "7SYTA45K",
    "mortimer": "37Z8WXRN",
    "jarrett": "BQCCFN9M",
    "who_adolescent": "RCZKMC43",
    "zhang": "Y96VASBL",
    "santoro": "S9DWUGNY",
    "edelman": "T48PCTZP",
    "xiao": "99LKVM7K",
    "barbosa": "E4VGVI5F",
    "faundes": "YURRMKC9",
    "malmqvist": "9MA8UAT2",
    "dawood": "A4ZL26YS",
}

SOURCE_LABELS = {
    "li_2023_awhs": "Li et al.",
    "bull_2019_natural_cycles": "Bull et al.",
    "stricker_2006_reference": "Stricker et al.",
}

AGE_LABELS = ["<20", "20–24", "25–29", "30–34", "35–39", "40–44", "45–49", "≥50"]
AGE_OVULATION = [0.62, 0.90, 0.97, 0.97, 0.97, 0.95, 0.88, 0.70]

BLUE = "#1F4E79"
TEAL = "#2A788E"
GOLD = "#D99A2B"
RED = "#B24745"
GRAY = "#6B7280"


def get_zotero_item(key: str) -> dict:
    url = (
        f"http://localhost:23119/api/users/0/items/{key}"
        "?format=json&include=data,csljson"
    )
    with urllib.request.urlopen(url, timeout=15) as response:
        payload = json.load(response)
    data = payload["data"]
    if "SNEPUMQA" not in data.get("collections", []):
        raise RuntimeError(f"Zotero item {key} is not in the CERES collection")
    csl = json.loads(payload["csljson"])[0]
    uri = f"http://zotero.org/users/{ZOTERO_LIBRARY_ID}/items/{key}"
    csl["id"] = uri
    return {"key": key, "uri": uri, "csl": csl}


def citation_field(items: Sequence[dict], citation_id: str) -> OxmlElement:
    citation_items = [
        {
            "id": item["uri"],
            "uris": [item["uri"]],
            "itemData": item["csl"],
        }
        for item in items
    ]
    payload = {
        "citationID": citation_id,
        "properties": {
            "unsorted": False,
            "formattedCitation": "(0)",
            "plainCitation": "0",
            "noteIndex": 0,
        },
        "citationItems": citation_items,
        "schema": "https://github.com/citation-style-language/schema/raw/master/csl-citation.json",
    }
    instruction = (
        " ADDIN ZOTERO_ITEM CSL_CITATION "
        + json.dumps(payload, ensure_ascii=False, separators=(",", ":"))
        + " "
    )
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), instruction)
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "(0)"
    run.append(text)
    field.append(run)
    return field


def add_live_citation(
    paragraph,
    item_groups: Sequence[str],
    zotero_items: dict[str, dict],
    counter: int,
) -> int:
    paragraph.add_run(" ")
    digest = hashlib.sha1(
        ("|".join(item_groups) + f"|{counter}").encode("utf-8")
    ).hexdigest()[:8]
    paragraph._p.append(
        citation_field([zotero_items[key] for key in item_groups], f"HC{digest}")
    )
    return counter + 1


def move_before(anchor, block) -> None:
    element = block._element if hasattr(block, "_element") else block._p
    anchor.addprevious(element)


def paragraph_before(document: Document, anchor, text: str = "", style: str | None = None):
    paragraph = document.add_paragraph(style=style)
    if text:
        paragraph.add_run(text)
    move_before(anchor, paragraph)
    return paragraph


def heading_before(document: Document, anchor, text: str, level: int = 2):
    paragraph = document.add_paragraph(text, style=f"Heading {level}")
    move_before(anchor, paragraph)
    return paragraph


def page_break_before(document: Document, anchor):
    paragraph = document.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)
    move_before(anchor, paragraph)
    return paragraph


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_twips))


def set_cell_margins(cell, top=55, start=65, bottom=55, end=65) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(
    table,
    widths_inches: Sequence[float],
    *,
    font_size: float = 7.5,
    first_column_left: bool = True,
) -> None:
    if len(widths_inches) != len(table.columns):
        raise ValueError(
            f"width count {len(widths_inches)} does not match "
            f"{len(table.columns)} columns"
        )
    if sum(widths_inches) > 9.0001:
        raise ValueError(f"table width exceeds 9 inches: {sum(widths_inches)}")
    widths_twips = [round(width * 1440) for width in widths_inches]
    table.autofit = False
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_twips)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "0")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_twips:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)

    set_repeat_table_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        cannot_split = row._tr.get_or_add_trPr().find(qn("w:cantSplit"))
        if cannot_split is None:
            cannot_split = OxmlElement("w:cantSplit")
            row._tr.get_or_add_trPr().append(cannot_split)
        for column_index, (cell, width) in enumerate(zip(row.cells, widths_twips)):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.keep_together = True
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT
                    if column_index == 0 and first_column_left
                    else WD_ALIGN_PARAGRAPH.CENTER
                )
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(font_size)
                    if row_index == 0:
                        run.bold = True


def add_table_before(
    document: Document,
    anchor,
    headers: Sequence[str],
    rows: Iterable[Sequence[object]],
    widths_inches: Sequence[float],
    *,
    font_size: float = 7.5,
):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = str(value)
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = str(value)
    set_table_geometry(table, widths_inches, font_size=font_size)
    move_before(anchor, table)
    return table


def format_caption(paragraph, label: str, remainder: str) -> None:
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(label)
    run.bold = True
    paragraph.add_run(remainder)


def add_figure_before(
    document: Document,
    anchor,
    image_path: Path,
    label: str,
    caption: str,
    *,
    width_inches: float = 8.65,
) -> None:
    image_paragraph = document.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.keep_together = True
    image_paragraph.add_run().add_picture(str(image_path), width=Inches(width_inches))
    move_before(anchor, image_paragraph)
    caption_paragraph = document.add_paragraph()
    caption_paragraph.style = document.styles["Normal"]
    caption_paragraph.paragraph_format.keep_with_next = False
    format_caption(caption_paragraph, label, caption)
    move_before(anchor, caption_paragraph)


def create_workflow_figure(output: Path) -> None:
    output.parent.mkdir(parents=True, exist_ok=True)
    fig, ax = plt.subplots(figsize=(13.5, 6.2))
    ax.set_xlim(0, 13.5)
    ax.set_ylim(0, 6.2)
    ax.axis("off")

    boxes = [
        (0.2, 3.35, 1.45, 1.35, "Inputs", "Age, requested days,\nseed, modifiers"),
        (1.95, 3.35, 1.9, 1.35, "Patient profile", "Resolve stable traits\nonce per diary"),
        (
            4.15,
            3.35,
            2.05,
            1.35,
            "Diary start",
            "Cycle 1, cycle day 1\nNo phase offset",
        ),
        (
            6.5,
            3.15,
            2.85,
            1.75,
            "Generate cycle n",
            "Draw ovulation, length,\nbleeding, and phase timing;\nrender the complete cycle",
        ),
        (
            9.65,
            3.35,
            1.8,
            1.35,
            "Append days",
            "Keep records in order,\nup to days still needed",
        ),
    ]
    for index, (x, y, width, height, title, body) in enumerate(boxes):
        color = BLUE if index in {0, 2} else TEAL
        patch = plt.Rectangle(
            (x, y),
            width,
            height,
            facecolor="white",
            edgecolor=color,
            linewidth=2,
            joinstyle="round",
        )
        ax.add_patch(patch)
        ax.text(
            x + width / 2,
            y + height - 0.39,
            title,
            ha="center",
            va="center",
            fontsize=10.8,
            fontweight="bold",
            color=BLUE,
        )
        ax.text(
            x + width / 2,
            y + 0.46,
            body,
            ha="center",
            va="center",
            fontsize=9.0,
            color="#222222",
        )
        if index < len(boxes) - 1:
            next_x = boxes[index + 1][0]
            ax.annotate(
                "",
                xy=(next_x - 0.05, y + height / 2),
                xytext=(x + width + 0.05, y + height / 2),
                arrowprops={"arrowstyle": "->", "color": GRAY, "lw": 1.7},
            )

    decision_center = (12.35, 4.02)
    decision_half_width = 0.78
    decision_half_height = 0.72
    decision = plt.Polygon(
        [
            (decision_center[0], decision_center[1] + decision_half_height),
            (decision_center[0] + decision_half_width, decision_center[1]),
            (decision_center[0], decision_center[1] - decision_half_height),
            (decision_center[0] - decision_half_width, decision_center[1]),
        ],
        closed=True,
        facecolor="white",
        edgecolor=TEAL,
        linewidth=2,
    )
    ax.add_patch(decision)
    ax.text(
        decision_center[0],
        decision_center[1],
        "Requested\ndays filled?",
        ha="center",
        va="center",
        fontsize=9.2,
        fontweight="bold",
        color=BLUE,
    )
    ax.annotate(
        "",
        xy=(decision_center[0] - decision_half_width - 0.03, decision_center[1]),
        xytext=(boxes[-1][0] + boxes[-1][2] + 0.05, boxes[-1][1] + boxes[-1][3] / 2),
        arrowprops={"arrowstyle": "->", "color": GRAY, "lw": 1.7},
    )

    output_box = (11.2, 0.55, 2.25, 1.05)
    output_patch = plt.Rectangle(
        (output_box[0], output_box[1]),
        output_box[2],
        output_box[3],
        facecolor="white",
        edgecolor=BLUE,
        linewidth=2,
        joinstyle="round",
    )
    ax.add_patch(output_patch)
    ax.text(
        output_box[0] + output_box[2] / 2,
        output_box[1] + 0.72,
        "Return diary",
        ha="center",
        va="center",
        fontsize=10.7,
        fontweight="bold",
        color=BLUE,
    )
    ax.text(
        output_box[0] + output_box[2] / 2,
        output_box[1] + 0.29,
        "Exact requested length;\nmay end within final cycle",
        ha="center",
        va="center",
        fontsize=8.7,
        color="#222222",
    )
    ax.annotate(
        "",
        xy=(output_box[0] + output_box[2] / 2, output_box[1] + output_box[3] + 0.04),
        xytext=(decision_center[0], decision_center[1] - decision_half_height - 0.03),
        arrowprops={"arrowstyle": "->", "color": GRAY, "lw": 1.7},
    )
    ax.text(12.62, 2.55, "Yes", fontsize=9.3, color=GRAY, ha="left")

    ax.annotate(
        "",
        xy=(boxes[3][0] + boxes[3][2] / 2, boxes[3][1] - 0.05),
        xytext=(decision_center[0] - 0.34, decision_center[1] - decision_half_height),
        arrowprops={
            "arrowstyle": "-|>",
            "connectionstyle": "arc3,rad=-0.38",
            "color": GRAY,
            "lw": 1.6,
        },
    )
    ax.text(
        9.72,
        2.18,
        "No: n ← n + 1; generate the next cycle",
        ha="center",
        va="center",
        fontsize=9.5,
        color=GRAY,
    )
    ax.text(
        6.75,
        5.68,
        "HORMONE-CYCLE diary-generation workflow",
        ha="center",
        va="center",
        fontsize=15,
        fontweight="bold",
        color=BLUE,
    )
    ax.text(
        6.75,
        5.23,
        "The diary begins at cycle day 1; repetition occurs at cycle boundaries.",
        ha="center",
        va="center",
        fontsize=10.2,
        color=GRAY,
    )
    ax.text(
        5.0,
        2.52,
        "The cycle-level loop begins after the patient profile is fixed.",
        ha="center",
        va="center",
        fontsize=9.0,
        color=GRAY,
    )
    fig.tight_layout()
    fig.savefig(output, dpi=240, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def baseline_metric_map(validation: dict) -> dict[str, dict]:
    return {metric["name"]: metric for metric in validation["baseline_metrics"]}


def create_validation_figure(validation: dict, output: Path) -> None:
    output.parent.mkdir(parents=True, exist_ok=True)
    metrics = baseline_metric_map(validation)
    fig, axes = plt.subplots(2, 2, figsize=(13.2, 8.2))

    age_keys = ["<20", "20-24", "25-29", "30-34", "35-39", "40-44", "45-49", "50+"]
    x = np.arange(len(age_keys))
    observed_means = [metrics[f"cycle_mean_{key}"]["observed"] for key in age_keys]
    target_means = [metrics[f"cycle_mean_{key}"]["expected"] for key in age_keys]
    axes[0, 0].plot(x, target_means, "o--", color=GOLD, label="Calibration target")
    axes[0, 0].plot(x, observed_means, "o-", color=TEAL, label="Simulated")
    axes[0, 0].set_xticks(x, AGE_LABELS, rotation=25)
    axes[0, 0].set_ylabel("Mean cycle length (days)")
    axes[0, 0].set_title("A. Age-stratified cycle length")
    axes[0, 0].legend(frameon=False, fontsize=9)

    observed_irregularity = [
        metrics[f"cycle_irregularity_{key}"]["observed"] for key in age_keys
    ]
    target_irregularity = [
        metrics[f"cycle_irregularity_{key}"]["expected"] for key in age_keys
    ]
    axes[0, 1].plot(
        x, np.array(target_irregularity) * 100, "o--", color=GOLD, label="Calibration target"
    )
    axes[0, 1].plot(
        x, np.array(observed_irregularity) * 100, "o-", color=TEAL, label="Simulated"
    )
    axes[0, 1].set_xticks(x, AGE_LABELS, rotation=25)
    axes[0, 1].set_ylabel("Adjacent-cycle difference ≥7 days (%)")
    axes[0, 1].set_title("B. Age-stratified irregularity")

    phases = [
        ("Follicular", "follicular_mean_days"),
        ("Luteal", "luteal_mean_days"),
        ("Bleeding", "bleeding_mean_days"),
    ]
    phase_x = np.arange(len(phases))
    width = 0.36
    axes[1, 0].bar(
        phase_x - width / 2,
        [
            4.0 if key == "bleeding_mean_days" else metrics[key]["expected"]
            for _, key in phases
        ],
        width,
        color=GOLD,
        label="Calibration target",
    )
    axes[1, 0].bar(
        phase_x + width / 2,
        [metrics[key]["observed"] for _, key in phases],
        width,
        color=TEAL,
        label="Simulated",
    )
    axes[1, 0].set_xticks(phase_x, [label for label, _ in phases])
    axes[1, 0].set_ylabel("Days")
    axes[1, 0].set_title("C. Phase and bleeding duration")

    phase_names = [
        "early_follicular",
        "mid_follicular",
        "pre_ovulatory",
        "ovulation",
        "early_luteal",
        "mid_luteal",
        "late_luteal",
    ]
    phase_labels = ["Early F", "Mid F", "Pre-O", "O", "Early L", "Mid L", "Late L"]
    estradiol_ratios = [
        100
        * metrics[f"estradiol_{phase}"]["observed"]
        / metrics[f"estradiol_{phase}"]["expected"]
        for phase in phase_names
    ]
    progesterone_ratios = [
        100
        * metrics[f"progesterone_{phase}"]["observed"]
        / metrics[f"progesterone_{phase}"]["expected"]
        for phase in phase_names
    ]
    axes[1, 1].axhline(100, color=GOLD, linestyle="--", linewidth=1.5)
    axes[1, 1].plot(
        np.arange(7), estradiol_ratios, "o-", color=TEAL, label="Estradiol"
    )
    axes[1, 1].plot(
        np.arange(7), progesterone_ratios, "s-", color=RED, label="Progesterone"
    )
    axes[1, 1].set_xticks(np.arange(7), phase_labels, rotation=25)
    axes[1, 1].set_ylabel("Simulated value / calibration target (%)")
    axes[1, 1].set_ylim(65, 115)
    axes[1, 1].set_title("D. Hormone subphase calibration")
    axes[1, 1].legend(frameon=False, fontsize=9)

    for axis in axes.flat:
        axis.spines["top"].set_visible(False)
        axis.spines["right"].set_visible(False)
        axis.grid(axis="y", alpha=0.18)
    fig.suptitle(
        "HORMONE-CYCLE internal target-reproduction checks",
        fontsize=16,
        fontweight="bold",
        color=BLUE,
        y=0.99,
    )
    fig.tight_layout(rect=(0, 0, 1, 0.96))
    fig.savefig(output, dpi=240, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def metric_label(name: str) -> str:
    replacements = {
        "cycle_mean_": "Mean cycle length, age ",
        "cycle_irregularity_": "Adjacent cycles differing by ≥7 days, age ",
        "follicular_mean_days": "Mean follicular interval",
        "luteal_mean_days": "Mean luteal interval",
        "bleeding_mean_days": "Mean bleeding duration",
        "estradiol_": "Estradiol, ",
        "progesterone_": "Progesterone, ",
    }
    for prefix, label in replacements.items():
        if name.startswith(prefix):
            suffix = name[len(prefix) :].replace("_", " ")
            suffix = (
                suffix.replace("20-24", "20–24")
                .replace("25-29", "25–29")
                .replace("30-34", "30–34")
                .replace("35-39", "35–39")
                .replace("40-44", "40–44")
                .replace("45-49", "45–49")
                .replace("50+", "≥50")
            )
            return label + suffix
    return name.replace("_", " ")


def metric_value(metric: dict, field: str) -> str:
    value = float(metric[field])
    if metric["name"].startswith("cycle_irregularity"):
        return f"{100 * value:.1f}%"
    if metric["name"].startswith("progesterone"):
        return f"{value:.2f}"
    return f"{value:.2f}"


def calibration_rows(validation: dict) -> list[list[str]]:
    rows = []
    for metric in validation["baseline_metrics"]:
        rows.append(
            [
                metric_label(metric["name"]),
                metric_value(metric, "observed"),
                metric_value(metric, "expected"),
                f"{metric_value(metric, 'lower_bound')} to "
                f"{metric_value(metric, 'upper_bound')}",
                "Pass" if metric["passed"] else "Fail",
                SOURCE_LABELS.get(metric["citation_key"], metric["citation_key"]),
            ]
        )
    return rows


def subgroup_rows(validation: dict) -> list[list[str]]:
    baseline = validation["subgroup_analysis"]["baseline_reference"]
    rows = [
        [
            "Baseline reference",
            f"{baseline['mean_cycle_days']:.2f}",
            f"{100 * baseline['ovulation_rate']:.1f}%",
            f"{baseline['mean_bleeding_days']:.2f}",
            f"{100 * baseline['irregularity_rate']:.1f}%",
            f"{100 * baseline['amenorrhea_rate']:.1f}%",
            "Reference",
        ]
    ]
    labels = {
        "pcos": "Polycystic ovary syndrome",
        "cyclic_ocp": "Cyclic combined oral contraceptive",
        "continuous_ocp": "Continuous combined oral contraceptive",
        "hormonal_iud": "Levonorgestrel-releasing intrauterine device",
        "copper_iud": "Copper intrauterine device",
        "perimenopause": "Perimenopause",
        "peri_menarche": "Early postmenarche",
        "dysmenorrhea": "Primary dysmenorrhea",
    }
    for key, result in validation["subgroup_analysis"]["subgroups"].items():
        summary = result["summary"]
        rows.append(
            [
                labels[key],
                f"{summary['mean_cycle_days']:.2f}",
                f"{100 * summary['ovulation_rate']:.1f}%",
                f"{summary['mean_bleeding_days']:.2f}",
                f"{100 * summary['irregularity_rate']:.1f}%",
                f"{100 * summary['amenorrhea_rate']:.1f}%",
                f"{sum(check['passed'] for check in result['checks'])}/"
                f"{len(result['checks'])} passed",
            ]
        )
    return rows


def add_hormone_cycle_section(
    document: Document,
    anchor,
    validation: dict,
    assets: dict[str, Path],
    zotero_items: dict[str, dict],
) -> None:
    citation_counter = 1

    page_break_before(document, anchor)
    heading_before(
        document,
        anchor,
        "S3. HORMONE-CYCLE: data-driven menstrual and hormone diary simulator",
        level=1,
    )
    paragraph = paragraph_before(
        document,
        anchor,
        (
            "HORMONE-CYCLE is a new, custom, open-source simulator built for the "
            "present project. It generates reproducible daily menstrual-cycle, "
            "bleeding, ovulation, estradiol, and progesterone diaries from a compact "
            "set of patient inputs. The name HORMONE-CYCLE is the software title, "
            "not an abbreviation. The implementation is intended for simulation "
            "research and method evaluation; it is not a mechanistic model of the "
            "hypothalamic-pituitary-ovarian axis, a diagnostic device, or a clinical "
            "hormone calculator."
        ),
    )
    citation_counter = add_live_citation(
        paragraph, ["hormone_cycle"], zotero_items, citation_counter
    )
    paragraph = paragraph_before(
        document,
        anchor,
        (
            "The design is hierarchical: population evidence constrains age-specific "
            "distributions; each simulated person receives stable latent traits; each "
            "cycle receives a stochastic realization; and a daily renderer converts "
            "cycle timing into hormone and bleeding trajectories. Large app-based "
            "cohorts informed cycle length and irregularity, observed ovulatory-cycle "
            "data informed follicular and luteal timing, and phase-specific laboratory "
            "medians informed estradiol and progesterone anchors."
        ),
    )
    citation_counter = add_live_citation(
        paragraph, ["li", "bull", "stricker"], zotero_items, citation_counter
    )

    heading_before(document, anchor, "S3.1. Scope, inputs, and outputs", level=2)
    paragraph_before(
        document,
        anchor,
        (
            "The public simulation call accepts age in years, diary length in days, "
            "a random seed, and optional medical or reproductive modifiers. The seed "
            "initializes an isolated pseudorandom stream, so identical inputs produce "
            "identical diaries without depending on global random state. Input "
            "validation rejects nonfinite ages, nonpositive diary lengths, and "
            "incompatible modifier combinations."
        ),
    )
    paragraph_before(
        document,
        anchor,
        (
            "Each daily record contains the calendar-day index, cycle index, day within "
            "cycle, realized cycle length, estradiol in picograms per milliliter, "
            "progesterone in nanograms per milliliter, a binary ovulation indicator, "
            "and a binary bleeding indicator. Cycle summaries retain ovulatory status, "
            "ovulation day, follicular and luteal lengths, bleeding duration, and the "
            "patient profile used to render the cycle. Menstrual bleeding is treated "
            "as a modeled episode beginning on cycle day 1; the terminology is "
            "consistent with standard clinical definitions, while the simulator "
            "deliberately simplifies bleeding intensity to a daily indicator."
        ),
    )
    citation_counter = add_live_citation(
        paragraph, ["fraser"], zotero_items, citation_counter
    )
    add_figure_before(
        document,
        anchor,
        assets["workflow"],
        "Appendix Figure A1. ",
        (
            "HORMONE-CYCLE workflow. Patient-level parameters are resolved once, "
            "cycle-level parameters are sampled repeatedly, and the daily renderer "
            "continues until the requested diary length is filled. All stages use the "
            "same seeded random stream."
        ),
        width_inches=8.7,
    )

    heading_before(document, anchor, "S3.2. Patient-level parameter resolution", level=2)
    paragraph_before(
        document,
        anchor,
        (
            "Age selects one of eight calibration bands. Each band supplies a target "
            "mean cycle length and a target probability that two adjacent cycles differ "
            "by at least 7 days. Under the working assumption that independent "
            "cycle-length deviations are Gaussian with within-person standard "
            "deviation σ, the adjacent-cycle difference has standard deviation √2σ. "
            "The implementation numerically inverts "
            "2[1−Φ(7/(√2σ))] to obtain the σ that reproduces the published "
            "irregularity probability, where Φ is the standard normal cumulative "
            "distribution function. This inversion converts an observable population "
            "summary into a latent simulation parameter; it is not presented as a "
            "biological law."
        ),
    )
    paragraph_before(
        document,
        anchor,
        (
            "A person's mean cycle length is then sampled around the age-band target "
            "with an age-dependent between-person standard deviation. Stable "
            "person-level estradiol and progesterone amplitude multipliers are "
            "lognormally distributed with coefficients of variation 0.18 and 0.22, "
            "respectively. The baseline day-to-day hormone-noise scale is 0.06. "
            "Baseline bleeding duration is sampled around 4.7 days with a standard "
            "deviation of 1.0 day. These dispersion terms were tuned to create realistic "
            "within- and between-person heterogeneity while retaining the published "
            "central targets."
        ),
    )
    patient_rows = [
        ["Age <20 years", "30.3 days", "31.2%", "0.62", "Li et al.; calibrated"],
        ["Age 20–24 years", "30.0 days", "20.4%", "0.90", "Li et al.; calibrated"],
        ["Age 25–29 years", "29.1 days", "16.4%", "0.97", "Li et al.; calibrated"],
        ["Age 30–34 years", "28.8 days", "14.7%", "0.97", "Li et al.; calibrated"],
        ["Age 35–39 years", "28.8 days", "15.9%", "0.97", "Li et al.; calibrated"],
        ["Age 40–44 years", "28.4 days", "20.2%", "0.95", "Li et al.; calibrated"],
        ["Age 45–49 years", "28.2 days", "27.2%", "0.88", "Li et al.; calibrated"],
        ["Age ≥50 years", "30.8 days", "55.4%", "0.70", "Li et al.; calibrated"],
        ["Ovulatory luteal interval", "12.4 days", "Standard deviation 1.7 days", "Not applicable", "Bull et al."],
        ["Bleeding duration", "4.7 days", "Standard deviation 1.0 day", "Not applicable", "Bull et al."],
    ]
    paragraph = paragraph_before(document, anchor)
    format_caption(
        paragraph,
        "Appendix Table A1. ",
        (
            "Core patient- and cycle-level targets. The ovulation probabilities are "
            "investigator-selected calibration constants constrained by age and "
            "reproductive-stage evidence; they are not direct prevalence estimates "
            "from a single study."
        ),
    )
    add_table_before(
        document,
        anchor,
        [
            "Parameter or age band",
            "Mean or target",
            "Irregularity or dispersion",
            "Baseline ovulation probability",
            "Evidence role",
        ],
        patient_rows,
        [1.75, 1.45, 2.0, 1.7, 2.1],
        font_size=7.6,
    )

    heading_before(document, anchor, "S3.3. Cycle generation algorithm", level=2)
    paragraph_before(
        document,
        anchor,
        (
            "For each cycle, ovulation is first drawn from a Bernoulli distribution "
            "using the resolved patient probability. Cycle length is sampled from a "
            "truncated Gaussian distribution centered on the patient's mean, with a "
            "minimum of 18 days and a maximum of 120 days. Anovulatory cycles receive "
            "wider dispersion and reproductive-stage-specific shifts because their "
            "timing is less constrained. In an ovulatory cycle, luteal length is "
            "sampled from a truncated Gaussian distribution with mean 12.4 days, "
            "standard deviation 1.7 days, and minimum 9 days; the remaining length is "
            "assigned to the follicular interval. This concentrates most cycle-length "
            "variability in the follicular interval, consistent with observed cycle "
            "structure."
        ),
    )
    paragraph_before(
        document,
        anchor,
        (
            "Bleeding duration is independently sampled from the resolved patient "
            "distribution and clipped to the cycle. Person-level traits are multiplied "
            "by additional cycle-level estradiol and progesterone amplitude factors "
            "with coefficients of variation 0.08 and 0.10. The simulator continues "
            "drawing complete cycles and truncates only the final rendered output to "
            "the requested diary length; no cycle is circularly shifted and no diary "
            "segment is wrapped from one end to the other."
        ),
    )

    heading_before(document, anchor, "S3.4. Daily hormone and bleeding renderer", level=2)
    paragraph_before(
        document,
        anchor,
        (
            "Ovulatory hormone trajectories are constructed from seven reference "
            "subphases: early follicular, mid-follicular, preovulatory, ovulation, "
            "early luteal, midluteal, and late luteal. The renderer maps those anchors "
            "onto the realized follicular and luteal intervals and joins them with "
            "shape-preserving piecewise cubic Hermite interpolation. The final four "
            "cycle days withdraw toward the early-follicular baseline before bleeding, and the "
            "preovulatory estradiol maximum spans multiple daily samples. Anovulatory cycles use separate "
            "blunted estradiol and progesterone anchors rather than an ovulatory peak."
        ),
    )
    hormone_rows = [
        ["Early follicular", "42.9", "0.44"],
        ["Mid-follicular", "88.1", "0.56"],
        ["Preovulatory", "234.0", "1.31"],
        ["Ovulation", "141.0", "1.43"],
        ["Early luteal", "132.0", "3.95"],
        ["Midluteal", "117.0", "11.02"],
        ["Late luteal", "111.0", "6.75"],
    ]
    paragraph = paragraph_before(document, anchor)
    format_caption(
        paragraph,
        "Appendix Table A2. ",
        (
            "Reference hormone anchors used for ovulatory cycles. Estradiol is reported "
            "in picograms per milliliter and progesterone in nanograms per milliliter."
        ),
    )
    add_table_before(
        document,
        anchor,
        ["Menstrual subphase", "Estradiol target", "Progesterone target"],
        hormone_rows,
        [3.0, 3.0, 3.0],
        font_size=8.0,
    )
    paragraph = paragraph_before(
        document,
        anchor,
        (
            "Serial day-to-day variation is added with a stationary first-order autoregressive "
            "process having coefficient 0.92; progesterone noise is multiplied by 0.9. "
            "Each realized noise path is linearly bridged to zero at both cycle endpoints. "
            "The coefficient is an investigator-selected continuity parameter, not a "
            "published physiological constant. Final values are floored at 5 "
            "picograms per milliliter for estradiol and 0.05 nanograms per milliliter "
            "for progesterone and rounded to two decimals. Bleeding is marked from "
            "cycle day 1 through the realized bleeding duration, and ovulation is "
            "marked only on the sampled ovulation day."
        ),
    )
    citation_counter = add_live_citation(
        paragraph, ["stricker"], zotero_items, citation_counter
    )

    heading_before(
        document,
        anchor,
        "S3.5. Reproductive-stage and medical-factor modifiers",
        level=2,
    )
    paragraph = paragraph_before(
        document,
        anchor,
        (
            "Modifiers operate on the resolved profile before cycles are rendered. "
            "Their purpose is to reproduce the direction and approximate magnitude of "
            "well-described phenotypes, not to encode causal physiology or estimate "
            "clinical prevalence. Polycystic ovary syndrome produces longer and more "
            "variable cycles and less frequent ovulation, with age-dependent cycle "
            "length effects."
        ),
    )
    citation_counter = add_live_citation(
        paragraph, ["mortimer", "jarrett"], zotero_items, citation_counter
    )
    paragraph = paragraph_before(
        document,
        anchor,
        (
            "Early postmenarche profiles receive longer, more irregular, and more often "
            "anovulatory cycles, whereas perimenopause increases irregularity and "
            "reduces ovulation."
        ),
    )
    citation_counter = add_live_citation(
        paragraph, ["who_adolescent", "zhang", "santoro"], zotero_items, citation_counter
    )
    paragraph = paragraph_before(
        document,
        anchor,
        (
            "Cyclic and continuous combined oral contraceptive profiles suppress "
            "endogenous-equivalent ovulation and use fixed 28-day regimen cycles. A "
            "levonorgestrel-releasing intrauterine device reduces bleeding and permits "
            "ovulation in most cycles; a copper intrauterine device preserves ovulation "
            "while increasing bleeding duration."
        ),
    )
    citation_counter = add_live_citation(
        paragraph,
        ["edelman", "xiao", "barbosa", "faundes", "malmqvist"],
        zotero_items,
        citation_counter,
    )
    paragraph = paragraph_before(
        document,
        anchor,
        (
            "Primary dysmenorrhea is represented as a modest bleeding-duration shift "
            "without a large steroid-hormone perturbation."
        ),
    )
    citation_counter = add_live_citation(
        paragraph, ["dawood"], zotero_items, citation_counter
    )
    modifier_rows = [
        [
            "Polycystic ovary syndrome",
            "Mean cycle ×1.30 if age <25; ×1.22 at 25–34; ×1.15 at ≥35",
            "Variability ×1.55",
            "Ovulation ×0.48",
            "Bleeding +0.4 day; estradiol ×1.08; progesterone ×0.58; noise ×1.15",
        ],
        [
            "Early postmenarche",
            "Mean cycle +2.5 days",
            "Variability ×1.25",
            "Maximum 0.55",
            "Bleeding +0.5 day; estradiol ×0.92; progesterone ×0.72; noise ×1.15",
        ],
        [
            "Perimenopause",
            "Stage-specific cycle shifts",
            "Variability ×1.35",
            "Ovulation ×0.78",
            "Bleeding +0.6 day; progesterone ×0.82; noise ×1.25",
        ],
        [
            "Copper intrauterine device",
            "No cycle-length shift",
            "No added cycle variability",
            "Preserved",
            "Bleeding +1.2 days; bleeding standard deviation +0.25 day",
        ],
        [
            "Levonorgestrel-releasing intrauterine device",
            "No cycle-length shift",
            "No added cycle variability",
            "Maximum 0.82",
            "Bleeding −2.2 days with 0.8-day floor; amenorrhea probability 0.17",
        ],
        [
            "Cyclic combined oral contraceptive",
            "Fixed 28 days",
            "None",
            "Suppressed",
            "Approximately 4 days of scheduled withdrawal bleeding; suppressed hormone curves",
        ],
        [
            "Continuous combined oral contraceptive",
            "Fixed 28 days",
            "None",
            "Suppressed",
            "Amenorrhea probability 0.55 with modeled breakthrough bleeding",
        ],
        [
            "Primary dysmenorrhea",
            "No cycle-length shift",
            "No added cycle variability",
            "Preserved",
            "Bleeding +0.5 day; bleeding standard deviation +0.15 day",
        ],
    ]
    paragraph = paragraph_before(document, anchor)
    format_caption(
        paragraph,
        "Appendix Table A3. ",
        (
            "Implemented modifier values. Multipliers and offsets are transparent "
            "calibration parameters selected to place simulated summaries within "
            "prespecified evidence-constrained ranges."
        ),
    )
    add_table_before(
        document,
        anchor,
        [
            "Modifier",
            "Cycle-length effect",
            "Cycle variability",
            "Ovulation effect",
            "Bleeding and hormone effects",
        ],
        modifier_rows,
        [1.75, 2.0, 1.35, 1.25, 2.65],
        font_size=7.2,
    )

    heading_before(document, anchor, "S3.6. Adapter used in the present study", level=2)
    paragraph_before(
        document,
        anchor,
        (
            "The present study used two cohorts. The healthy ovulatory cohort sampled "
            "ages 18–45 years and forced every rendered cycle to be ovulatory. The "
            "heterogeneous cohort sampled ages 13–54.9 years and applied "
            "investigator-configured stress-test probabilities: polycystic ovary "
            "syndrome 0.10, early-postmenarche modifier 0.60 conditional on age under "
            "20 years, perimenopause modifier 0.55 conditional on age at least 45 "
            "years, and dysmenorrhea 0.12. These probabilities define the simulation "
            "experiment and must not be interpreted as population prevalence estimates."
        ),
    )
    paragraph_before(
        document,
        anchor,
        (
            "The adapter adds menstrual-onset, ovulatory-status, ovulation-day, "
            "reproductive-stage, and inadequate-luteal-phase fields. For an ovulatory "
            "cycle, inadequate luteal phase is designated when the maximum "
            "progesterone value from days 5–9 after ovulation is below 5 nanograms per "
            "milliliter; if that interval is unavailable, all available postovulation "
            "days are considered. Nonovulatory cycles and cycles lacking a usable "
            "ovulation day or progesterone field are designated inadequate. The "
            "seizure and hormone diaries are joined by the same calendar-day index. "
            "No circular shift, phase randomization, or wraparound is applied."
        ),
    )

    heading_before(
        document,
        anchor,
        "S3.7. Calibration and validation design",
        level=2,
    )
    paragraph_before(
        document,
        anchor,
        (
            "The current implementation was rerun for this appendix with 10,000 "
            "simulated participants, 365 days per participant, and seed 7. Baseline "
            "checks compare age-stratified mean cycle length and adjacent-cycle "
            "irregularity, mean follicular and luteal intervals, mean bleeding "
            "duration, and seven estradiol plus seven progesterone subphase summaries "
            "with prespecified target ranges. The source cohorts are large enough that "
            "their confidence intervals can be narrower than a reasonable tolerance "
            "for a simulator reconstructed from published summaries; therefore, the "
            "pass ranges are equivalence windows around target values rather than "
            "claims that the simulator reproduces the source-study confidence interval."
        ),
    )
    paragraph_before(
        document,
        anchor,
        (
            "Modifier scenarios are assessed separately against directional or bounded "
            "expectations: longer and more irregular cycles with polycystic ovary "
            "syndrome; ovulation suppression under combined oral contraceptives; "
            "reduced bleeding with a levonorgestrel-releasing intrauterine device; "
            "preserved ovulation and increased bleeding with a copper intrauterine "
            "device; increased irregularity with perimenopause; longer and less "
            "frequently ovulatory cycles early after menarche; and preserved ovulation "
            "with a modest bleeding shift in dysmenorrhea."
        ),
    )
    paragraph_before(
        document,
        anchor,
        (
            "These checks are best described as internal distributional calibration "
            "and face-validity checks because the same literature summaries informed "
            "parameter choice and evaluation ranges. They do not constitute independent "
            "external validation against unseen person-level hormone and bleeding "
            "diaries."
        ),
    )

    paragraph = paragraph_before(document, anchor)
    format_caption(
        paragraph,
        "Appendix Table A4. ",
        (
            "Complete baseline calibration results from the current 10,000-participant "
            "run. All 33 prespecified metrics passed. Hormone concentrations use "
            "picograms per milliliter for estradiol and nanograms per milliliter for "
            "progesterone; duration metrics use days."
        ),
    )
    add_table_before(
        document,
        anchor,
        ["Metric", "Simulated", "Target", "Pass range", "Result", "Primary source"],
        calibration_rows(validation),
        [2.55, 1.0, 1.0, 2.0, 0.75, 1.7],
        font_size=7.1,
    )

    paragraph = paragraph_before(document, anchor)
    format_caption(
        paragraph,
        "Appendix Table A5. ",
        (
            "Modifier-scenario summaries from the current validation run. Irregularity "
            "is the proportion of adjacent cycles differing by at least 7 days. "
            "Amenorrhea is reported only where the profile permits it."
        ),
    )
    add_table_before(
        document,
        anchor,
        [
            "Scenario",
            "Mean cycle (days)",
            "Ovulatory cycles",
            "Mean bleeding (days)",
            "Irregularity",
            "Amenorrhea",
            "Prespecified checks",
        ],
        subgroup_rows(validation),
        [2.05, 1.15, 1.15, 1.3, 1.0, 1.0, 1.35],
        font_size=7.0,
    )
    add_figure_before(
        document,
        anchor,
        assets["validation"],
        "Appendix Figure A2. ",
        (
            "Observed-versus-target calibration in the current run. Panels A and B "
            "show age-stratified cycle length and irregularity; panel C shows mean "
            "follicular, luteal, and bleeding durations; panel D expresses simulated "
            "hormone values as a percentage of each published subphase target. Early F "
            "and mid F denote early and mid-follicular phases; pre-O and O denote "
            "preovulatory and ovulation phases; early L, mid L, and late L denote "
            "luteal phases."
        ),
        width_inches=7.6,
    )

    heading_before(
        document,
        anchor,
        "S3.8. Reproducibility, intended uses, and limitations",
        level=2,
    )
    paragraph_before(
        document,
        anchor,
        (
            "The simulator is deterministic for fixed inputs and seed, implemented in "
            "pure Python, and distributed with unit tests, a command-line interface, "
            "and a machine-readable validation report. Reproducible research records "
            "should preserve the repository commit, configuration, seed, requested "
            "diary length, age and modifier inputs, and validation-report hash."
        ),
    )
    paragraph_before(
        document,
        anchor,
        (
            "Appropriate uses include simulation studies, power and operating-"
            "characteristic analyses, testing algorithms that consume menstrual or "
            "hormone diaries, and generating controlled synthetic examples. The "
            "simulator should not be used to forecast an individual's hormones, infer "
            "ovulation from clinical data, estimate treatment effects, or substitute "
            "for measured endocrine data."
        ),
    )
    paragraph_before(
        document,
        anchor,
        (
            "Important limitations are deliberate simplifications. The model does not "
            "simulate pregnancy, postpartum physiology, exogenous-hormone dose changes, "
            "medication adherence, secular trends, missing diary entries, bleeding "
            "intensity, laboratory-assay error, or feedback dynamics of the endocrine "
            "axis. Medical modifiers are coarse profile transformations and may "
            "overlap in real patients in ways not represented here. The published "
            "targets come from cohorts with their own selection, measurement, and "
            "cycle-tracking limitations. Finally, passing the internal calibration "
            "suite establishes consistency with selected summaries, not transportability "
            "to every population."
        ),
    )


def replace_text_nodes(element, old: str, new: str) -> int:
    changed = 0
    for text_node in element.iter(qn("w:t")):
        if text_node.text and old in text_node.text:
            text_node.text = text_node.text.replace(old, new)
            changed += 1
    return changed


def first_paragraph(document: Document, prefix: str):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise LookupError(prefix)


def insert_section_definition(document: Document, heading_prefix: str, text: str) -> None:
    heading = first_paragraph(document, heading_prefix)
    node = heading._p.getnext()
    paragraph = document.add_paragraph(text)
    node.addprevious(paragraph._p)


def revise_existing_acronyms(document: Document) -> None:
    # Section 1
    for paragraph in document.paragraphs:
        if "SHA-256" in paragraph.text:
            replace_text_nodes(
                paragraph._p,
                "SHA-256",
                "Secure Hash Algorithm 256-bit (SHA-256)",
            )
            break
    for paragraph in document.paragraphs:
        if "HORMONE-CYCLE version 0.1.0" in paragraph.text:
            replace_text_nodes(
                paragraph._p,
                "HORMONE-CYCLE version 0.1.0",
                (
                    "the HORMONE-CYCLE menstrual and hormone simulator, "
                    "version 0.1.0,"
                ),
            )
            break

    # Section 2 and its table/note
    simulator_paragraph = first_paragraph(document, "Both CHOCOLATES")
    replace_text_nodes(
        simulator_paragraph._p,
        "Both CHOCOLATES and HORMONE-CYCLE",
        (
            "Both the CHOCOLATES seizure-diary simulator and the HORMONE-CYCLE "
            "menstrual and hormone simulator"
        ),
    )
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                replacements = [
                    ("within-person SD", "within-person standard deviation"),
                    ("ILP rule", "Inadequate-luteal-phase rule"),
                    ("Maximum P4 on O+5 to O+9", "Maximum progesterone on days 5–9 after ovulation"),
                    ("Audit ILP fraction", "Audit inadequate-luteal-phase fraction"),
                    ("PCOS", "Polycystic ovary syndrome"),
                    ("Contraception/IUD settings", "Contraception/intrauterine-device settings"),
                    ("95% Wilson CI", "95% Wilson confidence interval"),
                    ("NB classifiable", "Negative-binomial classifiable"),
                    ("Ratio C3 applicable", "Ratio type C3 applicable"),
                    ("N/A", "Not applicable"),
                ]
                for old, new in replacements:
                    replace_text_nodes(cell._tc, old, new)
    for paragraph in document.paragraphs:
        if paragraph.style.name == "Table Note":
            replace_text_nodes(
                paragraph._p,
                "machine-readable CSV",
                "machine-readable comma-separated-values file",
            )
            replace_text_nodes(
                paragraph._p,
                "analysis-code and configuration SHA-256 fingerprints",
                (
                    "analysis-code and configuration Secure Hash Algorithm "
                    "256-bit fingerprints"
                ),
            )

    # Renumber the pre-existing later sections to make room for the new S3.
    for paragraph in document.paragraphs:
        if paragraph.style.name == "Heading 1":
            match = re.match(r"^S(\d+)\.\s", paragraph.text.strip())
            if match and 3 <= int(match.group(1)) <= 12:
                old = f"S{match.group(1)}."
                new = f"S{int(match.group(1)) + 1}."
                replace_text_nodes(paragraph._p, old, new)
    heading_replacements = [
        (
            "S5. C3/ILP algorithm and applicability",
            "S5. Type C3 and inadequate-luteal-phase algorithm and applicability",
        ),
        (
            "S7. C3 window-duration sensitivity",
            "S7. Type C3 window-duration sensitivity",
        ),
        (
            "S9. Exploratory C3 negative-binomial calibration check",
            "S9. Exploratory type C3 negative-binomial calibration check",
        ),
    ]
    for paragraph in document.paragraphs:
        for old, new in heading_replacements:
            replace_text_nodes(paragraph._p, old, new)

    # Definitions are intentionally repeated because the user requested local
    # definitions within each section.
    insert_section_definition(
        document,
        "S4. Assumption-by-assumption review",
        (
            "This section uses type C1 for the perimenstrual pattern, type C2 for the "
            "periovulatory pattern, and type C3 for the inadequate-luteal-phase pattern. "
            "Inadequate luteal phase (ILP) and negative binomial (NB) are defined here "
            "for the assumption table."
        ),
    )
    insert_section_definition(
        document,
        "S5. Type C3 and inadequate-luteal-phase algorithm",
        (
            "Type C3 is the inadequate-luteal-phase catamenial pattern. In this section, "
            "inadequate luteal phase (ILP), average daily seizure frequency (ADSF), "
            "follicular (F), ovulatory (O), luteal (L), and menstrual (M) are used as "
            "local abbreviations."
        ),
    )
    insert_section_definition(
        document,
        "S6. Cumulative simulated Herzog",
        (
            "This section defines type C1 as ADSF(M)/ADSF(F+L), type C2 as "
            "ADSF(O)/ADSF(F+L), and type C3 as ADSF(O+L+M)/ADSF(F), where average "
            "daily seizure frequency (ADSF), follicular (F), ovulatory (O), luteal "
            "(L), and menstrual (M) are local abbreviations."
        ),
    )
    insert_section_definition(
        document,
        "S7. Type C3 window-duration",
        (
            "Type C3 denotes the inadequate-luteal-phase pattern. Confidence interval "
            "(CI) is defined here for the table; the table spells out the term in its "
            "header."
        ),
    )
    insert_section_definition(
        document,
        "S9. Exploratory type C3 negative-binomial",
        (
            "Type C3 denotes the inadequate-luteal-phase pattern. In this section, "
            "inadequate luteal phase (ILP), average daily seizure frequency (ADSF), "
            "follicular (F), ovulatory (O), luteal (L), and menstrual (M) are local "
            "abbreviations; the negative-binomial model is otherwise written in full."
        ),
    )
    insert_section_definition(
        document,
        "S10. Mutually exclusive pattern",
        (
            "Type C1 denotes the perimenstrual pattern, type C2 the periovulatory "
            "pattern, and type C3 the inadequate-luteal-phase pattern."
        ),
    )
    insert_section_definition(
        document,
        "S11. Simulator-feature associations",
        (
            "The machine-readable comma-separated-values file reports 95% Wilson "
            "confidence intervals; no abbreviation is used below."
        ),
    )
    insert_section_definition(
        document,
        "S12. Realized distribution",
        (
            "Standard deviation and inadequate luteal phase are written in full in this "
            "section. Type C3 denotes the inadequate-luteal-phase pattern."
        ),
    )

    # Expand remaining section-specific caption/body shorthand.
    replacements = [
        ("Supplementary Table S4 (machine-readable CSV)", "Supplementary Table S4 (machine-readable comma-separated-values file)"),
        ("95% Wilson intervals", "95% Wilson confidence intervals"),
        ("within-person cycle-length SD", "within-person cycle-length standard deviation"),
        ("ILP fraction", "inadequate-luteal-phase fraction"),
        ("ILP is displayed", "Inadequate luteal phase is displayed"),
        ("Exploratory NB C3", "Exploratory negative-binomial type C3"),
        ("NB C1/C2 calibration", "Negative-binomial type C1/type C2 calibration"),
        ("Nonovulatory cycles designated ILP", "Nonovulatory cycles designated inadequate luteal phase"),
        ("ILP threshold", "Inadequate-luteal-phase threshold"),
    ]
    for paragraph in document.paragraphs:
        for old, new in replacements:
            replace_text_nodes(paragraph._p, old, new)
    for cell in (cell for table in document.tables for row in table.rows for cell in row.cells):
        for old, new in replacements:
            replace_text_nodes(cell._tc, old, new)


def constrain_all_tables(document: Document) -> None:
    widths = {
        0: [1.0, 1.25, 0.85, 2.15, 0.75, 1.65, 1.35],
        1: [1.7, 3.65, 3.65],
        2: [2.1, 1.4, 1.3, 2.1, 2.1],
        3: [2.1, 1.4, 1.3, 2.1, 2.1],
        4: [3.0, 3.0, 3.0],
        5: [1.2, 1.3, 1.3, 1.0, 2.3, 1.9],
        6: [1.1, 1.15, 1.5, 1.3, 1.95, 2.0],
        7: [1.15, 1.45, 1.4, 1.0, 2.3, 1.7],
        8: [1.9, 0.8, 1.1, 0.8, 0.8, 0.8, 1.1, 1.7],
    }
    fonts = {0: 6.8, 1: 7.1, 6: 6.9, 8: 6.9}
    # Only the nine source tables exist when this function is called.
    if len(document.tables) != 9:
        raise RuntimeError(f"Expected 9 source tables, found {len(document.tables)}")
    for index, table in enumerate(document.tables):
        set_table_geometry(
            table,
            widths[index],
            font_size=fonts.get(index, 7.5),
        )


def resize_existing_supplementary_figures(document: Document) -> None:
    """Keep the three large pre-existing figure captions with their images."""

    target_width = round(7.55 * 914400)
    for figure_number in (1, 2, 3):
        caption = first_paragraph(
            document, f"Supplementary Figure S{figure_number}."
        )
        caption.paragraph_format.keep_with_next = True
        node = caption._p.getnext()
        while node is not None and not node.xpath(".//w:drawing"):
            node = node.getnext()
        if node is None:
            raise RuntimeError(
                f"Could not find image paragraph for Supplementary Figure S{figure_number}"
            )
        extents = node.xpath(".//wp:extent | .//a:xfrm/a:ext")
        if not extents:
            raise RuntimeError(
                f"Could not find image extent for Supplementary Figure S{figure_number}"
            )
        current_width = int(extents[0].get("cx"))
        current_height = int(extents[0].get("cy"))
        target_height = round(current_height * target_width / current_width)
        for extent in extents:
            extent.set("cx", str(target_width))
            extent.set("cy", str(target_height))
        p_pr = node.find(qn("w:pPr"))
        if p_pr is None:
            p_pr = OxmlElement("w:pPr")
            node.insert(0, p_pr)
        keep_lines = p_pr.find(qn("w:keepLines"))
        if keep_lines is None:
            keep_lines = OxmlElement("w:keepLines")
            p_pr.append(keep_lines)


def set_page_geometry(document: Document) -> None:
    for section in document.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.header_distance = Inches(0.45)
        section.footer_distance = Inches(0.45)


def set_nlm_zotero_style(docx_path: Path) -> None:
    replacement_id = "http://www.zotero.org/styles/nlm-citation-sequence"
    temp_path = docx_path.with_suffix(".style-tmp.docx")
    with zipfile.ZipFile(docx_path, "r") as source, zipfile.ZipFile(
        temp_path, "w", zipfile.ZIP_DEFLATED
    ) as target:
        for item in source.infolist():
            data = source.read(item.filename)
            if item.filename == "docProps/custom.xml":
                text = data.decode("utf-8")
                text, count = re.subn(
                    r'<style id="[^"]+" locale="en-US"',
                    f'<style id="{replacement_id}" locale="en-US"',
                    text,
                    count=1,
                )
                if count == 0:
                    text, count = re.subn(
                        r'&lt;style id="[^"]+" locale="en-US"',
                        f'&lt;style id="{replacement_id}" locale="en-US"',
                        text,
                        count=1,
                    )
                if count != 1:
                    raise RuntimeError("Could not replace the Zotero CSL style identifier")
                data = text.encode("utf-8")
            target.writestr(item, data)
    temp_path.replace(docx_path)


def set_document_metadata(document: Document) -> None:
    document.core_properties.title = (
        "Appendix S1 — False Positive Catamenial Epilepsy Classification: "
        "A Simulation Study"
    )
    document.core_properties.subject = (
        "Expanded HORMONE-CYCLE methods, calibration, validation, tables, and figures"
    )
    document.core_properties.comments = (
        "Built from the user-repaired draft_v6_appendix_s1 source; source preserved."
    )


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument("--validation", type=Path, default=DEFAULT_VALIDATION)
    parser.add_argument("--asset-dir", type=Path, default=DEFAULT_ASSET_DIR)
    args = parser.parse_args()

    validation = json.loads(args.validation.read_text())
    if not validation["baseline_passed"]:
        raise RuntimeError("Baseline validation did not pass")
    if len(validation["baseline_metrics"]) != 37:
        raise RuntimeError("Expected exactly 37 baseline calibration metrics")

    workflow_path = args.asset_dir / "hormone_cycle_workflow.png"
    validation_path = args.asset_dir / "hormone_cycle_validation.png"
    create_workflow_figure(workflow_path)
    create_validation_figure(validation, validation_path)

    zotero_items = {
        name: get_zotero_item(key) for name, key in ZOTERO_KEYS.items()
    }

    document = Document(args.source)
    set_page_geometry(document)
    constrain_all_tables(document)
    revise_existing_acronyms(document)
    resize_existing_supplementary_figures(document)

    anchor = first_paragraph(document, "S4. Assumption-by-assumption review")._p
    add_hormone_cycle_section(
        document,
        anchor,
        validation,
        {"workflow": workflow_path, "validation": validation_path},
        zotero_items,
    )
    set_document_metadata(document)

    args.output.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.output)
    set_nlm_zotero_style(args.output)

    print(
        json.dumps(
            {
                "output": str(args.output),
                "source_sha256": hashlib.sha256(args.source.read_bytes()).hexdigest(),
                "output_sha256": hashlib.sha256(args.output.read_bytes()).hexdigest(),
                "baseline_metrics": len(validation["baseline_metrics"]),
                "baseline_passed": validation["baseline_passed"],
                "subgroups": len(validation["subgroup_analysis"]["subgroups"]),
                "zotero_items": len(zotero_items),
                "assets": [str(workflow_path), str(validation_path)],
            },
            indent=2,
        )
    )


if __name__ == "__main__":
    main()
