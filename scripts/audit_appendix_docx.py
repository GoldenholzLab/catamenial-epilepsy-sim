"""Print the block order and page/table geometry of an appendix DOCX."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    args = parser.parse_args()

    document = Document(args.docx)
    print(
        f"sections={len(document.sections)} "
        f"paragraphs={len(document.paragraphs)} tables={len(document.tables)}"
    )
    for index, section in enumerate(document.sections):
        print(
            f"SECTION {index}: "
            f"page={section.page_width.inches:.2f}x{section.page_height.inches:.2f}in "
            f"margins=L{section.left_margin.inches:.2f}/R{section.right_margin.inches:.2f}/"
            f"T{section.top_margin.inches:.2f}/B{section.bottom_margin.inches:.2f}in"
        )

    paragraph_index = 0
    table_index = 0
    for child in document.element.body.iterchildren():
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            paragraph = Paragraph(child, document._body)
            text = " ".join(paragraph.text.split())
            if text:
                print(
                    f"P{paragraph_index:03d} [{paragraph.style.name}] "
                    f"{text[:260]}"
                )
            paragraph_index += 1
        elif tag == "tbl":
            table = Table(child, document._body)
            grid = child.find(qn("w:tblGrid"))
            grid_widths = []
            if grid is not None:
                grid_widths = [
                    int(column.get(qn("w:w"), "0")) for column in grid.iterchildren()
                ]
            header = " | ".join(
                " ".join(cell.text.split())[:75] for cell in table.rows[0].cells
            )
            print(
                f"T{table_index:02d} {len(table.rows)}x{len(table.columns)} "
                f"grid={sum(grid_widths) / 1440:.2f}in :: {header}"
            )
            table_index += 1


if __name__ == "__main__":
    main()
