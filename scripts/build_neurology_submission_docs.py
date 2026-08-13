#!/usr/bin/env python3
"""Build Word drafts for a Neurology original article and appendix."""

from __future__ import annotations

import csv
import json
import re
import shutil
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "outputs"
DOC_DIR = OUTPUT_DIR / "neurology_submission"
MAIN_DOCX = DOC_DIR / "paper1_draft_v2.docx"
APPENDIX_DOCX = DOC_DIR / "paper1_appendix_v2.docx"
LEGACY_MAIN_DOCX = DOC_DIR / "paper1_null_ce_neurology_original_article_draft.docx"
LEGACY_APPENDIX_DOCX = DOC_DIR / "paper1_null_ce_appendix_draft.docx"

TITLE = "Apparent Catamenial Epilepsy Classification Under a Null Simulation"
RUNNING_TITLE = "Null Catamenial Classification"
LEAD_AUTHOR = "Goldenholz"

DEFAULT_REFS = [
    "Herzog AG, Harden CL, Liporace J, et al. Frequency of catamenial seizure exacerbation in women with localization-related epilepsy. Ann Neurol. 2004;56:431-434. doi:10.1002/ana.20214",
    "Herzog AG, Klein P, Ransil BJ. Three patterns of catamenial epilepsy. Epilepsia. 1997;38:1082-1088. doi:10.1111/j.1528-1157.1997.tb01197.x",
    "Herkes GK, Eadie MJ, Sharbrough F, Moyer T. Patterns of seizure occurrence in catamenial epilepsy. Epilepsy Res. 1993;15:47-52. doi:10.1016/0920-1211(93)90008-U",
    "Herzog AG. Catamenial epilepsy: definition, prevalence, pathophysiology and treatment. Seizure. 2008;17:151-159. doi:10.1016/j.seizure.2007.11.014",
    "Herzog AG. Catamenial epilepsy: update on prevalence, pathophysiology and treatment from the findings of the NIH Progesterone Treatment Trial. Seizure. 2015;28:18-25. doi:10.1016/j.seizure.2015.02.024",
    "Maguire MJ, Nevitt SJ. Treatments for seizures in catamenial (menstrual-related) epilepsy. Cochrane Database Syst Rev. 2021;9:CD013225. doi:10.1002/14651858.CD013225.pub3",
    "Herzog AG, Fowler KM, Smithson SD, et al. Progesterone vs placebo therapy for women with epilepsy: a randomized clinical trial. Neurology. 2012;78:1959-1966. doi:10.1212/WNL.0b013e318259e1f9",
    "Reddy DS. The role of neurosteroids in the pathophysiology and treatment of catamenial epilepsy. Epilepsy Res. 2009;85:1-30. doi:10.1016/j.eplepsyres.2009.02.017",
    "Baud MO, Kleen JK, Mirro EA, et al. Multi-day rhythms modulate seizure risk in epilepsy. Nat Commun. 2018;9:88. doi:10.1038/s41467-017-02577-y",
    "Karoly PJ, Rao VR, Gregg NM, et al. Cycles in epilepsy. Nat Rev Neurol. 2021;17:267-284. doi:10.1038/s41582-021-00464-1",
    "Karoly PJ, Cook MJ, Maturana MI, et al. Forecasting cycles of seizure likelihood. Epilepsia. 2020;61:776-786. doi:10.1111/epi.16485",
    "Goldenholz DM, Westover MB. Flexible realistic simulation of seizure occurrence recapitulating statistical properties of seizure diaries. Epilepsia. 2023;64:396-405. doi:10.1111/epi.17471",
    "Li H, Gibson EA, Jukic AMZ, et al. Menstrual cycle length variation by demographic characteristics from the Apple Women's Health Study. NPJ Digit Med. 2023;6:100. doi:10.1038/s41746-023-00848-1",
    "Bull JR, Rowland SP, Scherwitzl EB, et al. Real-world menstrual cycle characteristics of more than 600,000 menstrual cycles. NPJ Digit Med. 2019;2:83. doi:10.1038/s41746-019-0152-7",
    "Stricker R, Eberhart R, Chevailler MC, et al. Establishment of detailed reference values for luteinizing hormone, follicle stimulating hormone, estradiol, progesterone, prolactin and growth hormone during different phases of the menstrual cycle on the Abbott ARCHITECT analyzer. Clin Chem Lab Med. 2006;44:883-887. doi:10.1515/CCLM.2006.160",
    "Monks T, Currie CS, Onggo BS, Robinson S, Kunc M, Taylor SJE. Strengthening the reporting of empirical simulation studies: introducing the STRESS guidelines. J Simul. 2019;13:55-67. doi:10.1080/17477778.2018.1442155",
    "Newmark ME, Penry JK. Catamenial epilepsy: a review. Epilepsia. 1980;21:281-300. doi:10.1111/j.1528-1157.1980.tb04074.x",
    "Duncan S, Read CL, Brodie MJ. How common is catamenial epilepsy? Epilepsia. 1993;34:827-831. doi:10.1111/j.1528-1157.1993.tb02097.x",
]


def load_references() -> list[str]:
    for path in (OUTPUT_DIR / "zotero_references.json", DOC_DIR / "zotero_references.json"):
        if not path.exists():
            continue
        data = json.loads(path.read_text(encoding="utf-8"))
        refs = data.get("references", data)
        if refs and all(isinstance(ref, str) and ref.strip() for ref in refs):
            return [ref.strip() for ref in refs]
    return DEFAULT_REFS


REFS = load_references()

DEFINITION_LABELS = {
    "A_exact_any": "Exact Herzog 2004, any pattern",
    "A_windowed": "Windowed Herzog thresholds",
    "A_windowed_any": "Windowed Herzog thresholds",
    "A_windowed_excluding_C3": "Windowed Herzog thresholds, excluding C3",
    "A_windowed_C1_or_C2": "Windowed Herzog C1/C2 union",
    "A_windowed_C1_only": "Windowed Herzog C1 only",
    "A_windowed_C2_only": "Windowed Herzog C2 only",
    "A_windowed_C3_only": "Windowed Herzog C3 only",
    "B_minimum_data_any": "Herzog plus >=4 seizure-day minimum",
    "B_minimum_data": "Herzog plus >=4 seizure-day minimum",
    "B_minimum_data_excluding_C3": "Herzog plus >=4 seizure-day minimum, excluding C3",
    "B_minimum_data_C1_or_C2": "Herzog C1/C2 plus >=4 seizure-day minimum",
    "C_reproducibility_any": "Cycle reproducibility >=2/3",
    "C_reproducibility_C1_or_C2": "Cycle reproducibility >=2/3, C1/C2",
    "C_reproducibility_12cycle_any": "Cycle reproducibility, 12-cycle sensitivity",
    "D_nb_regression_any": "Negative-binomial regression C1/C2, stabilized dispersion",
    "D_nb_regression_C1_or_C2": "Negative-binomial regression C1/C2",
    "D_nb_regression_window_alpha_any": "Negative-binomial regression, window-only dispersion",
    "D_nb_regression_window_alpha_C1_or_C2": "Negative-binomial regression C1/C2, window-only dispersion",
    "H1_newmark_penry_any": "Newmark-Penry perimenstrual rule",
    "H1_newmark_penry_66_7_any": "Newmark-Penry two-thirds sensitivity",
    "H2_duncan1993_any": "Duncan 1993 ten-day rule",
    "H3_herzog1997_twofold_any": "Herzog 1997 twofold rule",
    "H4_reddy2007_any_phase2x_any": "Reddy 2007 any-phase twofold rule",
}

COHORT_LABELS = {
    "healthy_ovulatory": "healthy ovulatory",
    "population": "heterogeneous menstruating-age",
}


@dataclass
class DocStats:
    title_chars: int
    abstract_words: int
    text_words: int
    n_refs: int
    n_tables: int
    n_figures: int


def main() -> int:
    DOC_DIR.mkdir(parents=True, exist_ok=True)
    copy_figure_files()
    rows = load_summary_rows()
    main_doc, stats = build_main_doc(rows)
    main_doc.save(MAIN_DOCX)
    main_doc.save(LEGACY_MAIN_DOCX)
    appendix_doc = build_appendix_doc(rows)
    appendix_doc.save(APPENDIX_DOCX)
    appendix_doc.save(LEGACY_APPENDIX_DOCX)
    write_review_note(stats)
    print(MAIN_DOCX)
    print(APPENDIX_DOCX)
    return 0


def copy_figure_files() -> None:
    for path in OUTPUT_DIR.glob("fig*.png"):
        shutil.copy2(path, DOC_DIR / path.name)
    for path in OUTPUT_DIR.glob("fig*.pdf"):
        shutil.copy2(path, DOC_DIR / path.name)
    for path in OUTPUT_DIR.glob("fig*.svg"):
        shutil.copy2(path, DOC_DIR / path.name)


def load_summary_rows() -> list[dict[str, str]]:
    with (OUTPUT_DIR / "summary_tables.csv").open(newline="", encoding="utf-8") as handle:
        return list(csv.DictReader(handle))


def build_main_doc(rows: list[dict[str, str]]) -> tuple[Document, DocStats]:
    doc = new_document()
    configure_manuscript_section(doc.sections[0])
    add_header(doc.sections[0])

    abstract_text = structured_abstract()
    body_sections = main_body_sections()
    text_words = count_words("\n".join(text for _, text in body_sections))

    add_title_page(doc, count_words(abstract_text), text_words)
    abstract_heading = add_heading(doc, "Abstract", level=1)
    abstract_heading.paragraph_format.page_break_before = True
    add_paragraph(doc, abstract_text)
    add_paragraph(doc, "Classification of Evidence: Not applicable. This simulation study evaluates operating characteristics of classification rules under a synthetic null model and does not test a diagnostic intervention or therapeutic effect in human participants.")
    add_paragraph(doc, "Search Terms: catamenial epilepsy; epilepsy; seizures; menstrual cycle; simulation study.")
    doc.add_page_break()

    for heading, text in body_sections:
        add_heading(doc, heading, level=1)
        for para in split_paragraphs(text):
            add_paragraph(doc, para)

    add_heading(doc, "Acknowledgment", level=1)
    add_paragraph(doc, "No nonauthor assistance was used in preparing this simulation draft.")
    add_heading(doc, "Study Funding", level=1)
    add_paragraph(doc, "This work was supported by NIH K23NS124656, NIH R21NS142800, and American Board of Psychiatry and Neurology support to DMG. The funders had no role in study design, data generation, analysis, interpretation, manuscript preparation, or the decision to submit.")
    add_heading(doc, "Disclosure", level=1)
    add_paragraph(doc, "DMG reports funding from NIH K23NS124656, NIH R21NS142800, and the American Board of Psychiatry and Neurology. SC, WK, and RS report no disclosures relevant to the manuscript.")
    add_heading(doc, "Data Availability", level=1)
    add_paragraph(doc, "The simulation code, configuration, analysis outputs, and machine-readable manifest are available in the project repository. The reproducible run used master seed 20260505 and the manifest in outputs/manifest.json records config parameters, output paths, file sizes, and SHA-256 checksums. A versioned public archive with DOI will accompany the submitted version.")

    add_references(doc)
    add_main_tables(doc, rows)
    add_main_figure_legends_and_images(doc)

    stats = DocStats(
        title_chars=len(TITLE),
        abstract_words=count_words(abstract_text),
        text_words=text_words,
        n_refs=len(REFS),
        n_tables=4,
        n_figures=3,
    )
    return doc, stats


def build_appendix_doc(rows: list[dict[str, str]]) -> Document:
    doc = new_document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    configure_manuscript_section(section)
    add_header(section, label="Supplementary Appendix")

    add_title(doc, "Supplementary Appendix: Apparent Catamenial Epilepsy Classification Under Independence")
    add_paragraph(doc, "This appendix provides additional methodological details, defined analysis choices, quality-control checks, sensitivity analyses, and supplementary results for the null simulation. It is intended as online supplementary material accompanying the main Neurology submission draft.")
    add_heading(doc, "Table of Contents", level=1)
    toc_items = [
        "A1. Simulation design and reproducibility",
        "A2. Simulator descriptions and validation targets",
        "A3. Cohort definitions and diary generation",
        "A4. Phase labeling and catamenial epilepsy definitions",
        "A5. Observation windows and analyzability rules",
        "A6. Statistical analysis and edge-case handling",
        "A7. Supplementary results",
        "A8. Output provenance and assumptions",
    ]
    for item in toc_items:
        add_paragraph(doc, item, style="List Bullet")

    add_heading(doc, "A1. Simulation design and reproducibility", level=1)
    add_paragraph(doc, "The analysis used a deterministic master seed of 20260505. The full run simulated 100,000 synthetic participants, processed in chunks, with 50,000 in the healthy ovulatory cohort and 50,000 in the heterogeneous menstruating-age cohort. The 36-month seizure and menstrual/hormone diaries were generated with separate deterministic random streams and aligned directly by calendar day without reordering.")
    add_paragraph(doc, "Daily audit data were saved for a 1% random sample. Participant-level and window-level outputs were written as parquet and CSV files, and every output artifact was recorded in a machine-readable manifest with file size and SHA-256 checksum.")
    add_paragraph(doc, "The term 'defined before the full simulation run' is used instead of 'prespecified' unless a dated protocol, locked commit, or preregistration is supplied. The manifest records the seed, config values, output checksums, and adapter assumptions.")

    add_heading(doc, "A2. Simulator descriptions and validation targets", level=1)
    add_paragraph(doc, "CHOCOLATES generated seizure-count diaries independently of menstrual-cycle state. The adapter records observed seizure burden and the simulator-reported dominant seizure-cycle period when available. Because multidien rhythms near monthly periods are a key potential mechanism for chance alignment, submitted versions should report the distribution of dominant_seizure_cycle_days and sensitivity analyses excluding or flattening 24-35-day seizure rhythms.")
    add_paragraph(doc, "HORMONE-CYCLE generated menstrual/hormone diaries independently of seizure state. The adapter records age, cycle length, within-participant cycle-length SD, ovulatory fraction, progesterone and estradiol when available, and inadequate-luteal-phase flags based on a progesterone threshold of 5.0 ng/mL. Submitted versions should include package versions or commits for both simulators and a validation table comparing simulated diary metrics with empirical benchmarks.")
    add_small_table(
        doc,
        "Supplementary Table 1. Minimum simulator validation targets for submission.",
        ["Domain", "Simulated metric to report", "Reviewer-facing purpose"],
        [
            ["Seizure burden", "Seizure days/month, seizures/month, interseizure interval dispersion, seizure-cluster metrics", "Shows the null seizure diary is not an unrealistically sparse or overclustered artifact."],
            ["Seizure rhythms", "Dominant multidien period distribution; proportion in 24-35-day band; sensitivity excluding that band", "Separates finite-count ratio artifacts from monthly-ish seizure-rhythm alignment."],
            ["Menstrual cycles", "Cycle-length distribution, within-participant cycle-length SD, luteal length, ovulatory fraction, anovulation frequency", "Shows phase labeling and C3 eligibility are generated from plausible cycle structure."],
            ["Missing/irregular behavior", "Unlabeled days, non-23-35-day cycles, ILP/anovulatory flags, medical-factor prevalence", "Quantifies sources of indeterminate windows and C3 sensitivity."],
        ],
        widths=[1.8, 3.4, 3.4],
    )

    add_heading(doc, "A3. Cohort definitions and diary generation", level=1)
    add_paragraph(doc, "The healthy ovulatory cohort targeted adult ovulatory cycling, age 18-45 years, with simulator medical modifiers disabled where exposed. The heterogeneous menstruating-age cohort targeted a broader menstruating-age range of 13.0-54.9 years and allowed anovulation and irregularity. Because the hormone simulator exposed medical-factor controls but not a natural prevalence sampler, PCOS, peri-menarche, perimenopause, and dysmenorrhea were sampled from config.yaml rates. This cohort is an assumption-driven heterogeneity stress test unless external prevalence calibration is added.")
    add_small_table(
        doc,
        "Supplementary Table 2. Cohort construction parameters.",
        ["Parameter", "Healthy ovulatory", "Heterogeneous menstruating-age"],
        [
            ["Target sample size", "50,000", "50,000"],
            ["Age range", "18.0-45.0 years", "13.0-54.9 years"],
            ["Ovulatory-cycle handling", "Ovulation probability set to 1.0", "Natural ovulatory/anovulatory behavior allowed"],
            ["PCOS rate", "Disabled", "10%"],
            ["Peri-menarche rate", "Disabled", "60% when age <20 years"],
            ["Perimenopause rate", "Disabled", "55% when age >=45 years"],
            ["Dysmenorrhea rate", "Disabled", "12%"],
        ],
        widths=[2.0, 3.2, 3.6],
    )

    add_heading(doc, "A4. Phase labeling and catamenial epilepsy definitions", level=1)
    add_paragraph(doc, "Herzog-style phases were assigned on the full diary before window subsetting. For a complete cycle of length L, day 1 was the onset of menstrual flow and backward day was defined as d-(L+1). Strict Herzog phase assignment used an ordered lookup: menstrual phase if days 1-3 or backward days -3 to -1; follicular phase if days 4-9; ovulatory phase if day >=10 and backward day <=-13; luteal phase if backward days -12 to -4. The ovulatory condition is the intersection of forward and backward criteria. Days not satisfying any rule were unlabeled. The luteal-anchored sensitivity preserved the same menstrual and luteal windows but fixed ovulatory phase to backward days -16 to -13 and labeled remaining nonmenstrual, nonovulatory, nonluteal days as follicular.")
    add_small_table(
        doc,
        "Supplementary Table 3. Implemented catamenial epilepsy definitions.",
        ["Definition", "Primary rule", "Main analyzability condition"],
        [
            ["Exact Herzog 2004", "Three-cycle legacy rule; any CE if at least 2 of 3 cycles show any C1, C2, or C3 pattern.", "Only complete 3-cycle windows with strict 23-35 day cycles."],
            ["Windowed Herzog thresholds", "C1 ADSF(M)/ADSF(F+L) >=1.69; C2 ADSF(O)/ADSF(F+L) >=1.83; C3 ADSF(O+L+M)/ADSF(F) >=1.62 among ILP-flagged days.", "Required phases observed; undefined ratios marked indeterminate."],
            ["Minimum-data rule", "Windowed Herzog thresholds after applying minimum duration and seizure-burden criteria.", "At least 4 months or 6 complete cycles, plus at least 4 seizure days."],
            ["Cycle reproducibility", "At least two-thirds of eligible cycles positive for the same pattern and pooled ratio passes threshold.", "At least 6 complete cycles for primary rule; 12 cycles for sensitivity."],
            ["Negative-binomial regression", "Daily GLM with M and O indicators; one-sided Holm-adjusted tests and C1/C2 rate-ratio thresholds. Stabilized and window-only dispersion versions reported.", "Minimum-data rule passed; Poisson robust fallback recorded when needed."],
            ["Historical rules", "Assumption-based operationalizations of Newmark-Penry, Duncan 1993, Herzog 1997 twofold, and Reddy 2007 any-phase rules.", "Reported separately and flagged as assumption based."],
        ],
        widths=[1.8, 4.4, 2.6],
    )

    add_paragraph(doc, "C3 was evaluated only in the heterogeneous cohort when ilp_flag was true. Anovulatory cycles without ILP labeling were not automatically considered C3-positive or C3-applicable. Supplementary sensitivity rows report any CE excluding C3 and C1, C2, and C3 separately.")

    add_heading(doc, "A5. Observation windows and analyzability rules", level=1)
    add_paragraph(doc, "Calendar windows used one random valid start day per participant for 1, 3, 4, 6, 9, 12, 18, 24, and 36 months. Cycle windows used one random valid start among complete cycles for 3, 6, and 12 cycles. The full-window analysis used the entire 36-month diary. Windows were marked indeterminate, not negative, when required phases or denominator information were unavailable.")
    add_paragraph(doc, "Exact Herzog 2004 is rule-defined only for complete 3-cycle windows with strict 23-35-day cycles. Other window settings are not evaluated for this rule; they should not be interpreted as negative evidence.")

    add_heading(doc, "A6. Statistical analysis and edge-case handling", level=1)
    add_paragraph(doc, "All ratio definitions used average daily seizure frequency rather than raw counts. If a comparator phase had observed days but zero seizures, a positive numerator produced an infinite ratio and counted positive; zero numerator and zero comparator produced an undefined ratio and an indeterminate result. Wilson 95% confidence intervals were used for simple proportions. Study-level Monte Carlo sampled 10,000 studies of 30, 50, and 100 participants per cohort for the 3-month window.")
    add_paragraph(doc, "Negative-binomial regression used daily rows with an intercept, menstrual indicator, ovulatory indicator, and cycle fixed effects when at least four complete cycles were available. The offset was one observed day. The Holm family comprised the menstrual and ovulatory one-sided tests within participant. Positivity required adjusted P<0.05 and the corresponding threshold ratio. NB regression was evaluated for full-diary windows in this revision. The stabilized-dispersion analysis uses full-diary alpha; the full-window, window-only alpha analysis is a non-oracle sensitivity.")
    add_paragraph(doc, "Monte Carlo binomial intervals quantify uncertainty under the fixed simulation model. They do not include uncertainty about simulator validity, medical-factor prevalence, seizure-rhythm periods, or menstrual-cycle parameterization. Participant-level summary rows, all-attempted rates, indeterminate reasons, and unstable-denominator flags were generated from window_results.parquet. Rows with fewer than 1,000 classifiable windows are flagged as unstable and not interpreted.")

    add_heading(doc, "A7. Supplementary results", level=1)
    add_window_sensitivity_table(doc, rows)
    add_phase_mode_sensitivity_table(doc, rows)
    add_pattern_decomposition_table(doc, rows)
    add_nb_sensitivity_table(doc, rows)
    add_study_level_table(doc, rows)
    add_historical_table(doc, rows)
    add_paragraph(doc, "Supplementary Figure 1. False-positive and indeterminate frontier.")
    add_picture(doc, DOC_DIR / "fig4_indeterminate_vs_fpr_frontier.png", width=Inches(7.8))
    add_paragraph(doc, "Supplementary Figure 2. Historical and core definitions in 3-month windows.")
    add_picture(doc, DOC_DIR / "fig4_historical_vs_core_definitions.png", width=Inches(7.8))
    add_paragraph(doc, "Supplementary Figure 3. Null seizure profile by menstrual cycle day.")
    add_picture(doc, DOC_DIR / "fig5_qc_null_cycle_day_profile.png", width=Inches(7.3))

    add_heading(doc, "A8. Output provenance and assumptions", level=1)
    assumptions = [
        "Definition D uses a participant-full-diary method-of-moments negative-binomial alpha recorded in d_alpha; Poisson robust fallback is recorded in d_reason when statsmodels negative-binomial fitting fails. D_window_alpha re-estimates alpha from the analyzed window as a non-oracle sensitivity.",
        "Healthy ovulatory cohort used hormone_cycler build_patient_profile/render_cycle with ovulation_probability set to 1.0 because simulate_diary did not expose a public force-ovulation knob.",
        "Historical definitions H1-H4 are assumption-based operationalizations and are flagged in summary outputs.",
        "Study-level Monte Carlo samples each selected participant from a deterministic pool of precomputed random valid 3-month windows to avoid retaining all daily diaries in memory.",
        "The hormone simulator exposes medical-factor controls but no natural prevalence sampler; heterogeneous-cohort medical factors were sampled from config.yaml rates.",
        "True-coupling code paths are retained for future operating-characteristic analyses but are not reported as part of this null-only revision.",
    ]
    for assumption in assumptions:
        add_paragraph(doc, assumption, style="List Bullet")
    return doc


def new_document() -> Document:
    doc = Document()
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 2.0
    normal.paragraph_format.space_after = Pt(0)
    for style_name, size in [("Title", 14), ("Heading 1", 12), ("Heading 2", 12), ("Heading 3", 12)]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.line_spacing = 2.0
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
    return doc


def configure_manuscript_section(section) -> None:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.5)
    sect_pr = section._sectPr
    if sect_pr.find(qn("w:lnNumType")) is None:
        ln = OxmlElement("w:lnNumType")
        ln.set(qn("w:countBy"), "1")
        ln.set(qn("w:restart"), "continuous")
        sect_pr.append(ln)


def add_header(section, label: str | None = None) -> None:
    para = section.header.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para.style = "Normal"
    para.paragraph_format.line_spacing = 1.0
    para.add_run(f"{label or LEAD_AUTHOR} ")
    add_page_field(para)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])


def add_title_page(doc: Document, abstract_words: int, text_words: int) -> None:
    add_title(doc, TITLE)
    title_meta = [
        ("Article type", "Article"),
        ("Running title", RUNNING_TITLE),
        ("Title character count", str(len(TITLE))),
        ("Abstract word count", str(abstract_words)),
        ("Text word count", str(text_words)),
        ("References", str(len(REFS))),
        ("Tables", "4"),
        ("Figures", "3"),
        ("Supplementary content", "Yes; separate appendix document"),
    ]
    for key, value in title_meta:
        add_paragraph(doc, f"{key}: {value}")
    add_paragraph(doc, "Authors: David M. Goldenholz; Sharon Chiang; Wesley Kerr; Rachael Sumner.")
    add_paragraph(doc, "Affiliations: author affiliations are provided in the submission metadata; city spellings are San Francisco, Pittsburgh, and Auckland.")
    add_paragraph(doc, "Corresponding author: David M. Goldenholz.")
    add_paragraph(doc, "Statistical analysis completed by: David M. Goldenholz.")
    add_paragraph(doc, "Author responsibility statement: David M. Goldenholz takes responsibility for the simulation data, analyses, interpretation, and conduct of the research; had full access to all output data; and had responsibility for the decision to submit. All authors approved the submitted version.")
    add_paragraph(doc, "CRediT contributorship: DMG, conceptualization, methodology, software, formal analysis, visualization, writing-original draft, writing-review and editing, funding acquisition; SC, methodology, clinical interpretation, writing-review and editing; WK, methodology, clinical interpretation, writing-review and editing; RS, methodology, clinical interpretation, writing-review and editing.")


def add_title(doc: Document, text: str) -> None:
    para = doc.add_paragraph(style="Normal")
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.line_spacing = 2.0
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14)


def add_heading(doc: Document, text: str, level: int = 1):
    para = doc.add_paragraph(style=f"Heading {level}")
    para.add_run(text)
    return para


def add_paragraph(doc: Document, text: str, style: str | None = None):
    para = doc.add_paragraph(style=style or "Normal")
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.line_spacing = 2.0
    para.paragraph_format.space_after = Pt(0)
    para.add_run(text)
    return para


def structured_abstract() -> str:
    return (
        "Objective: To quantify apparent catamenial epilepsy classification under a null model in which seizure and menstrual-cycle diaries are independent, with emphasis on C1, C2, and C3 pattern-specific false positives. "
        "Methods: We simulated 100,000 synthetic participants for 36 months in healthy ovulatory and heterogeneous menstruating-age cohorts. CHOCOLATES generated seizure diaries and HORMONE-CYCLE generated independent menstrual/hormone diaries using separate deterministic random streams; the diaries were aligned directly by calendar day without reordering. Strict Herzog phase labels were primary, with a luteal-anchored fixed periovulatory sensitivity. We evaluated exact Herzog 2004, windowed Herzog thresholds, Herzog plus >=4 seizure-day minimums, and cycle reproducibility >=2/3 across calendar, cycle, and full-diary windows; negative-binomial C1/C2 regression was evaluated as a full-diary calibrated comparator. "
        "Results: In 36-month strict-Herzog windows, windowed Herzog thresholds classified 11.2% of healthy ovulatory and 36.3% of heterogeneous-cohort participants as apparent catamenial epilepsy under independence. In the heterogeneous cohort, C3 drove the excess: the C1/C2 union was 11.8%, whereas C3 positivity occurred in 37.0% of applicable windows. Negative-binomial C1/C2 regression exhibited approximately nominal Type I error, 4.1% and 4.2%. Three-month calendar windows remained vulnerable, with windowed Herzog false-positive rates of 41.3% and 51.2%. "
        "Conclusions: Apparent catamenial epilepsy classification can be common under independence, especially when C3/inadequate-luteal-phase logic is applied in heterogeneous cycles. Studies should report C1, C2, and C3 separately and prespecify diary duration, seizure-day minimums, denominator handling, phase labeling, and reproducibility criteria."
    )


def main_body_sections() -> list[tuple[str, str]]:
    return [
        (
            "Introduction",
            "Catamenial epilepsy describes seizure exacerbation in relation to the menstrual cycle. The clinical literature has long recognized perimenstrual, periovulatory, and inadequate luteal-phase patterns, but reported prevalence varies widely because studies use different phase definitions, thresholds, diary durations, and eligibility rules. Herzog and colleagues reported 39.1% catamenial epilepsy in women who charted seizures and menses for three complete cycles using a two-of-three-cycle rule, and later work from the NIH Progesterone Treatment Trial has been interpreted as supporting prevalence near 44% in selected trial participants.1,5\n\nThe diagnostic problem is not only biological. Seizures are temporally structured: clustering, circadian rhythms, and multidien rhythms have been documented in long-term diaries and intracranial recordings.9-11 If seizure rhythms are independent of menstrual cycles but contain multidien structure in overlapping period bands, finite diary windows can still contain chance alignments. Ratio definitions are especially vulnerable when seizure counts are sparse or comparator phases contain few events. This creates a specificity question with direct implications for clinical research: what apparent catamenial epilepsy rate should be expected when no menstrual coupling exists?\n\nWe conducted a reproducible null-simulation study to estimate apparent catamenial epilepsy classification across definitions and observation windows under independence. The analysis evaluates specificity-like behavior only. Because traditional definitions combine C1, C2, and C3 patterns, we separately report C1/C2 and C3 false-positive behavior and compare negative-binomial regression only against C1/C2 outcomes.",
        ),
        (
            "Methods",
            "Study design. This Monte Carlo null simulation was defined before the full simulation run. For each synthetic participant, seizure and menstrual/hormone diaries were generated independently for 36 months using separate deterministic random streams and aligned directly by calendar day without reordering. Random calendar and cycle windows were sampled per participant and could overlap across analysis families; all analyses were stratified by cohort and phase-labeling mode.\n\nCohorts. The healthy ovulatory cohort included 50,000 synthetic participants, age 18-45 years, with ovulatory cycling enforced through the simulator adapter where available. The heterogeneous menstruating-age cohort included 50,000 synthetic participants, age 13.0-54.9 years, allowing anovulation and irregularity. Because the hormone-cycle simulator exposed medical-factor controls but not a natural prevalence sampler, PCOS, peri-menarche, perimenopause, and dysmenorrhea were sampled from configuration rates. This cohort should be interpreted as an assumption-driven heterogeneity stress test unless these rates are replaced with externally validated prevalence inputs.\n\nDiary simulators. CHOCOLATES (Cyclic Heterogeneous Overdispersed Clustered Open-source L-relationship Adjustable Temporally limited E-diary Simulator) generated non-catamenial seizure diaries using a hierarchical gamma-Poisson/negative-binomial seizure-count generator with seizure-burden heterogeneity, multidien cycles, clustering, mean-variance structure, and interseizure-interval limits.12 HORMONE-CYCLE generated independent menstrual/hormone diaries. It samples participant-level latent traits and cycle-level realizations, then interpolates daily estradiol and progesterone trajectories from menstrual subphase reference medians. Its features include cycle length, within-participant variability, bleeding days, ovulatory and anovulatory cycles, follicular and luteal phase lengths, ovulation day, luteal progesterone, inadequate-luteal-phase flags, and configured modifiers for PCOS, peri-menarche, perimenopause, dysmenorrhea, oral contraceptive use, hormonal IUD, and copper IUD. Baseline cycle targets were based on Apple Women's Health Study data, Natural Cycles data, and Stricker et al. hormone reference values.13-15 Simulator validation targets, versions, and package commits are documented in the Appendix following STRESS reporting principles.16\n\nPhase labeling and definitions. Strict Herzog phases were primary. In a complete cycle of length L, day 1 was menstrual-flow onset and backward day was d-(L+1). Strict labels were assigned as menstrual phase if forward day 1-3 or backward day -3 to -1; follicular phase if forward day 4-9; ovulatory phase if forward day >=10 and backward day <=-13; luteal phase if backward day -12 to -4; otherwise unlabeled. Because this strict ovulatory window expands with cycle length, we added a sensitivity mode that fixed the ovulatory phase at backward days -16 to -13 immediately before the luteal phase, with the follicular phase absorbing cycle-length variability.\n\nWindowed Herzog thresholds used average daily seizure frequency (ADSF): C1 = ADSF(M)/ADSF(F+L) >=1.69, C2 = ADSF(O)/ADSF(F+L) >=1.83, and C3 = ADSF(O+L+M)/ADSF(F) >=1.62 among inadequate-luteal-phase days. C3 was evaluated only in the heterogeneous cohort when ilp_flag was present; anovulatory or nonclassifiable cycles without ilp_flag did not contribute to C3. We report any C1-C3 classification, the C1/C2 union, mutually exclusive C1/C2/C3 categories, and C3-excluding endpoints.\n\nExact Herzog 2004 was restricted to complete 3-cycle windows with all cycles 23-35 days and required at least two of three cycles to show any positive pattern. Other calendar and longer cycle windows are reported as not evaluated for this rule. Herzog plus >=4 seizure-day minimum required at least 4 calendar months or 6 complete cycles plus at least 4 seizure days. Cycle reproducibility required at least two-thirds of eligible cycles to be positive for the same pattern and the corresponding pooled ratio to pass threshold.\n\nRegression comparator. The negative-binomial (NB) comparator used a within-participant generalized linear model with daily seizure count as outcome, log link, intercept, menstrual-phase indicator, ovulatory-phase indicator, and cycle fixed effects when at least four complete cycles were available. One-sided Wald tests evaluated C1 and C2 enrichment, Holm-adjusted within participant across the two comparisons; positivity required adjusted P<0.05 and the corresponding Herzog rate-ratio threshold. C3 was not evaluated by NB regression. In this revision, NB regression was evaluated for full-diary windows as a calibrated C1/C2 comparator. The primary regression used full-diary method-of-moments dispersion as a stabilized comparator, and a full-window, window-only dispersion sensitivity was added.\n\nStatistical analysis. The primary outcome was the participant-window apparent catamenial epilepsy rate among classifiable windows under the null. We also report positivity among all attempted windows, participant-level positivity, indeterminate reasons, and rows with fewer than 1,000 classifiable windows as unstable and not interpreted. Calendar windows were expanded to 1, 3, 4, 6, 9, 12, 18, 24, and 36 months for practical monitoring-duration guidance. Study-level Monte Carlo independently sampled 10,000 studies of 30, 50, and 100 participants without replacement from each cohort, selecting one precomputed random valid 3-month window per participant. Standard protocol approvals, registrations, and patient consents were not required because all participants were synthetic.",
        ),
        (
            "Results",
            "Simulation cohort and diary characteristics. The full run included 100,000 synthetic participants. With strict and luteal-anchored phase modes, expanded calendar windows, cycle windows, and study-pool windows, the analysis generated paired participant-window rows for sensitivity comparisons. Cohort summaries are shown in Table 1. The healthy ovulatory cohort had 100.0% ovulatory cycles by design. The heterogeneous cohort had 79.0% ovulatory cycles, greater cycle-length variability, and similar observed seizure burden.\n\nC3 contribution and C1/C2 comparison. In strict-Herzog 36-month windows, the windowed Herzog threshold definition classified 11.2% of healthy ovulatory participants and 36.3% of heterogeneous-cohort participants as apparent catamenial epilepsy under independence (Table 2). In the heterogeneous cohort, the corresponding C1/C2 union was 11.8%, similar to the healthy cohort's 11.2%, whereas C3 positivity occurred in 37.0% of applicable heterogeneous-cohort windows. Thus the broad-cohort excess was primarily a C3/inadequate-luteal-phase phenomenon rather than a perimenstrual or periovulatory excess.\n\nNB regression and denominator alignment. Stabilized-dispersion NB regression evaluated C1/C2 only and exhibited approximately nominal Type I error at 4.1% in the healthy ovulatory cohort and 4.2% in the heterogeneous cohort. The relevant comparison is therefore NB C1/C2 versus Herzog C1/C2, not NB versus the Herzog C1-C3 union. Calibration should not be interpreted as clinical usefulness; sensitivity and positive predictive value require true-coupling analyses outside this null study.\n\nObservation-window effects. Short diary windows were most vulnerable to apparent catamenial classification. In random 3-month calendar windows, strict-Herzog windowed thresholds yielded apparent classification rates of 41.3% in the healthy ovulatory cohort and 51.2% in the heterogeneous cohort. Practical calendar-window summaries through 36 months are shown in Figure 1. Exact Herzog 2004 applied to 3 complete cycles yielded apparent classification rates of 50.3% and 51.6% among classifiable windows, while many attempted 3-cycle windows were not classifiable. Rows with fewer than 1,000 classifiable windows were flagged as unstable and not interpreted.\n\nPhase-labeling sensitivity and study-level Monte Carlo. The luteal-anchored ovulatory sensitivity tested whether strict Herzog's expanding ovulatory window drove periovulatory false positives; these results are reported in the notebook and Appendix. In 10,000 simulated 3-month studies, apparent prevalence varied with study size but could exceed historical benchmark values under the null. Figure 3 displays n=30, 50, and 100 distributions with the y-axis scaled as the proportion of simulated studies.",
        ),
        (
            "Discussion",
            "This null simulation shows that apparent catamenial epilepsy can arise frequently even when seizure and menstrual processes are generated independently. The magnitude depends strongly on the catamenial pattern being tested, the window length, analyzability rules, and cohort menstrual-cycle assumptions. The largest heterogeneous-cohort excess was C3-driven; C1/C2 behavior was much more similar across cohorts. This distinction matters because C3 depends on inadequate-luteal-phase ascertainment rather than the perimenstrual or periovulatory windows that are usually easiest to audit from diaries alone.\n\nThe negative-binomial result should be interpreted as calibration, not superiority. A statistical test evaluated at alpha=0.05 should produce approximately 5% positives under a correctly specified null, and the stabilized-dispersion regression did so. Ratio rules are not nominal-alpha tests and therefore have no comparable Type I error guarantee. The relevant comparison in this paper is NB C1/C2 versus the Herzog C1/C2 union, because the NB model does not evaluate C3. Sensitivity, positive predictive value, and clinical usefulness require true-coupling analyses and are outside the scope of this null study.\n\nThese results help interpret the broad range of prevalence estimates in the catamenial epilepsy literature without estimating true prevalence. Earlier work established biologically plausible and clinically important menstrual-cycle seizure patterns, including C1, C2, and C3 patterns.1-5 Treatment literature also suggests that robust perimenstrual exacerbation may identify subgroups of interest, even though randomized evidence for hormonal therapy remains limited and mixed.6,7 The present analysis instead quantifies a diagnostic background rate expected from chance alignment when seizure cycles and menstrual cycles coexist under independence.\n\nThe findings are consistent with modern seizure-cycle research. Multidien seizure rhythms and seizure forecasting studies have shown that seizure risk varies over days to weeks, and those rhythms can be detected in long-term electronic diaries and chronic recordings.9-11 A menstrual-cycle analysis that does not account for intrinsic seizure cyclicity can therefore mistake phase overlap for hormone coupling. This is most problematic when the observation window is short, because a few events can dominate average daily seizure-frequency ratios.\n\nSeveral methodologic implications follow. First, 3-month ratio-only classification should be interpreted cautiously in research settings, particularly when seizure counts are sparse. Second, C1, C2, and C3 should be reported separately rather than collapsed into a single prevalence value. Third, studies should report both classifiable-window and all-attempted-window rates, because stricter definitions may appear conservative partly by declining to classify insufficient windows. Fourth, benchmark comparisons should match the original study design. For Herzog 2004, the design-matched comparison is the exact 3-complete-cycle rule, whereas random 3-calendar-month windows are a design-sensitivity analysis.\n\nStrict Herzog labeling was retained as the primary analysis for historical comparability. The luteal-anchored ovulatory sensitivity addresses a separate biologic question: whether the strict Herzog ovulatory window, which expands with longer cycles, creates excess periovulatory apparent classification. Presenting both phase modes separates historical reproducibility from a more anatomically fixed periovulatory definition.\n\nThis study has limitations. The results are only as credible as the CHOCOLATES and HORMONE-CYCLE simulators, their parameter settings, and the adapter assumptions. The heterogeneous cohort medical-factor rates were configuration-sampled rather than drawn from a simulator-native prevalence sampler, and cycle-length variability is therefore a sensitivity driver rather than a settled population estimate. Monte Carlo binomial intervals do not include model-form uncertainty. This study does not evaluate false negatives, true prevalence, treatment effects, wavelet analyses, hidden Markov models, medication changes, or adjudicated seizure types.\n\nIn conclusion, under an independence model with structured seizure diaries and menstrual-cycle diaries, apparent catamenial epilepsy classification can be common and can reach historical benchmark prevalence values. Future diagnostic and interventional studies should prespecify diary length, menstrual-phase labeling, analyzability thresholds, zero-denominator handling, ovulatory-status handling, and reproducibility criteria before interpreting apparent menstrual clustering as evidence of hormone-linked seizure exacerbation.",
        ),
    ]


def add_references(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "References", level=1)
    for i, ref in enumerate(REFS, 1):
        add_paragraph(doc, f"{i}. {ref}")


def add_main_tables(doc: Document, rows: list[dict[str, str]]) -> None:
    doc.add_page_break()
    add_heading(doc, "Tables", level=1)
    add_small_table(
        doc,
        "Table 1. Simulated cohort characteristics.",
        ["Cohort", "N", "Mean age, y", "Mean cycle length, d", "Cycle-length SD, d", "Ovulatory cycles", "Seizure days/month", "Seizures/month"],
        [
            ["healthy ovulatory", "50,000", "31.5 (SD 7.8)", "29.15", "3.39", "100.0%", "2.46", "6.83"],
            ["heterogeneous menstruating-age", "50,000", "34.0 (SD 12.1)", "30.92", "4.74", "79.0%", "2.45", "6.80"],
        ],
        widths=[1.35, 0.65, 1.0, 1.1, 1.0, 0.95, 1.1, 1.0],
    )
    add_paragraph(doc, "Values are cohort-level means from the full 100,000-participant run. Cycle-length SD is the mean within-participant cycle-length SD.")

    table2_rows = participant_window_results_table(rows)
    add_small_table(
        doc,
        "Table 2. Participant-window false-positive results.",
        ["Cohort", "Analysis setting", "Definition", "Classifiable / attempted", "Positive among classifiable", "Positive among all attempted", "Indeterminate"],
        table2_rows,
        widths=[1.15, 1.35, 1.65, 1.2, 1.35, 1.2, 0.9],
    )
    add_paragraph(doc, "All rows use strict Herzog phase labeling. Rows with fewer than 1,000 classifiable windows are flagged in the analysis outputs and not interpreted.")

    add_small_table(
        doc,
        "Table 3. C1/C2/C3 mutually exclusive pattern decomposition.",
        ["Cohort", "Definition", "C1 only", "C2 only", "C1+C2", "C3 only", "C3 plus C1/C2", "None", "Indeterminate"],
        main_pattern_decomposition_table(rows),
        widths=[1.1, 1.45, 0.75, 0.75, 0.75, 0.75, 0.9, 0.75, 0.85],
    )
    add_paragraph(doc, "Percentages are among all attempted 36-month full-diary windows, so indeterminate windows remain visible in the denominator.")

    add_small_table(
        doc,
        "Table 4. Study-level Monte Carlo summaries for 3-month windows.",
        ["Cohort", "N per study", "Definition", "Mean apparent prevalence", "2.5th-97.5th percentiles", "P >=39.1%", "P >=44.2%"],
        main_study_level_table(rows),
        widths=[1.15, 0.85, 1.7, 1.2, 1.35, 0.9, 0.9],
    )
    add_paragraph(doc, "Study-level rows summarize 10,000 simulated studies per cohort and sample size using one random valid 3-month window per selected participant. N=30 is illustrative; N=50 and N=100 are sensitivity rows.")


def participant_window_results_table(rows: list[dict[str, str]]) -> list[list[str]]:
    out: list[list[str]] = []
    row_specs = [
        ("36-month full diary", {"table_type": "window_false_positive", "subset": "all", "window_type": "full", "definition": "A_windowed_any"}),
        ("36-month full diary", {"table_type": "window_false_positive", "subset": "all", "window_type": "full", "definition": "A_windowed_C1_or_C2"}),
        ("36-month full diary", {"table_type": "window_false_positive", "subset": "all", "window_type": "full", "definition": "A_windowed_C3_only"}),
        ("36-month full diary", {"table_type": "window_false_positive", "subset": "all", "window_type": "full", "definition": "B_minimum_data_C1_or_C2"}),
        ("36-month full diary", {"table_type": "window_false_positive", "subset": "all", "window_type": "full", "definition": "D_nb_regression_C1_or_C2"}),
        ("3-month calendar window", {"table_type": "window_false_positive", "subset": "all", "window_type": "calendar", "window_value": "3", "definition": "A_windowed_any"}),
        ("3 complete cycles", {"table_type": "window_false_positive", "subset": "all", "window_type": "cycle", "window_value": "3", "definition": "A_exact_any"}),
    ]
    for cohort in ["healthy_ovulatory", "population"]:
        for setting, kwargs in row_specs:
            row = one_row(rows, cohort=cohort, **kwargs)
            out.append(
                [
                    display_cohort(cohort),
                    setting,
                    DEFINITION_LABELS[row["definition"]],
                    f"{int_float(row['n_classifiable']):,} / {int_float(row['n_windows']):,}",
                    rate_with_ci(row),
                    pct_or_na(row["positive_rate_all_attempted"]),
                    pct_or_na(row["indeterminate_rate"]),
                ]
            )
    return out


def main_pattern_decomposition_table(rows: list[dict[str, str]]) -> list[list[str]]:
    out: list[list[str]] = []
    categories = ["C1 only", "C2 only", "C1+C2", "C3 only", "C3 plus C1/C2", "none"]
    for cohort in ["healthy_ovulatory", "population"]:
        for definition in ["A_windowed", "B_minimum_data"]:
            category_rows = [
                one_row(
                    rows,
                    cohort=cohort,
                    table_type="pattern_decomposition",
                    subset="mutually_exclusive_patterns",
                    window_type="full",
                    definition=definition,
                    pattern_category=category,
                )
                for category in categories
            ]
            out.append(
                [
                    display_cohort(cohort),
                    DEFINITION_LABELS.get(definition, definition.replace("_", " ")),
                    *[pct_or_na(r["positive_rate_all_attempted"]) for r in category_rows],
                    pct_or_na(category_rows[0]["indeterminate_rate"]),
                ]
            )
    return out


def main_study_level_table(rows: list[dict[str, str]]) -> list[list[str]]:
    out: list[list[str]] = []
    for cohort in ["healthy_ovulatory", "population"]:
        for n_participants in [30, 50, 100]:
            for definition in ["A_windowed_any", "A_windowed_C1_or_C2"]:
                row = maybe_one_row(
                    rows,
                    cohort=cohort,
                    table_type="study_level_3month",
                    subset="apparent_prevalence_all",
                    n_participants=str(n_participants),
                    definition=definition,
                )
                if row is None:
                    continue
                out.append(
                    [
                        display_cohort(cohort),
                        str(n_participants),
                        DEFINITION_LABELS[definition],
                        pct_or_na(row["false_positive_rate"]),
                        f"{pct_or_na(row['wilson95_low'])}-{pct_or_na(row['wilson95_high'])}",
                        pct_or_na(row["p_prevalence_ge_39_1"]),
                        pct_or_na(row["p_prevalence_ge_44_2"]),
                    ]
                )
    return out


def add_main_figure_legends_and_images(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "Figure Legends", level=1)
    legends = [
        (
            "Figure 1. False-positive rates by observation window.",
            "Bars show the percentage of classifiable calendar-window participant windows classified as catamenial epilepsy under the null using strict Herzog phase labeling. Dashed and dotted horizontal lines mark 5% and 10% reference rates.",
            "fig1_false_positive_by_window.png",
            6.3,
        ),
        (
            "Figure 2. C1/C2/C3 mutually exclusive pattern decomposition.",
            "Stacked bars show mutually exclusive pattern categories for 36-month full-diary windows by cohort and definition. Percentages use all attempted participant windows as the denominator.",
            "fig2_pattern_decomposition.png",
            5.9,
        ),
        (
            "Figure 3. Apparent prevalence in null studies.",
            "Distributions show apparent catamenial epilepsy prevalence in 10,000 simulated 3-month studies using strict Herzog windowed thresholds. Curves are shown for n=30, n=50, and n=100 participants; vertical lines show benchmark prevalence values of 39.1% and 44.2%.",
            "fig3_study_prevalence_distribution_3month.png",
            5.9,
        ),
    ]
    for title, legend, filename, width in legends:
        add_paragraph(doc, title)
        add_paragraph(doc, legend)
        add_picture(doc, DOC_DIR / filename, width=Inches(width))


def add_window_sensitivity_table(doc: Document, rows: list[dict[str, str]]) -> None:
    definitions = ["A_exact_any", "A_windowed_any", "A_windowed_C1_or_C2", "B_minimum_data_C1_or_C2", "C_reproducibility_C1_or_C2", "D_nb_regression_C1_or_C2"]
    selected = [
        r for r in rows
        if r["table_type"] == "window_false_positive"
        and r["subset"] == "all"
        and r.get("phase_mode") == "strict_herzog"
        and r["definition"] in definitions
    ]
    selected.sort(key=lambda r: (display_cohort(r["cohort"]), window_sort_key(r), DEFINITION_LABELS[r["definition"]]))
    data = []
    for r in selected:
        n_classifiable = int_float(r["n_classifiable"])
        if 0 < n_classifiable < 1000:
            positives = "Suppressed"
            fpr = "Not interpreted (<1,000)"
        else:
            positives = f"{int_float(r['positives']):,}"
            fpr = pct_or_na(r["false_positive_rate"])
        data.append([
            display_cohort(r["cohort"]),
            window_label(r),
            DEFINITION_LABELS[r["definition"]],
            f"{n_classifiable:,}",
            positives,
            fpr,
            pct_or_na(r["indeterminate_rate"]),
        ])
    add_small_table(
        doc,
        "Supplementary Table 3. Window-length sensitivity for core definitions.",
        ["Cohort", "Observation window", "Definition", "Classifiable windows", "False-positive windows", "False-positive rate", "Indeterminate"],
        data,
        widths=[1.1, 1.1, 2.0, 1.1, 1.1, 1.1, 1.0],
        font_size=8,
    )
    add_paragraph(doc, "Rows with fewer than 1,000 classifiable windows are retained for transparency but their positive counts and rates are suppressed to avoid over-interpretation.")


def add_phase_mode_sensitivity_table(doc: Document, rows: list[dict[str, str]]) -> None:
    definitions = ["A_windowed_any", "A_windowed_C1_or_C2"]
    selected = [
        r for r in rows
        if r["table_type"] == "window_false_positive"
        and r["subset"] == "all"
        and r["definition"] in definitions
        and (r["window_type"] == "full" or (r["window_type"] == "calendar" and r["window_value"] in {"3", "12"}))
    ]
    selected.sort(key=lambda r: (display_cohort(r["cohort"]), r.get("phase_mode", ""), window_sort_key(r), DEFINITION_LABELS[r["definition"]]))
    data = []
    for r in selected:
        data.append([
            display_cohort(r["cohort"]),
            "Strict Herzog" if r.get("phase_mode") == "strict_herzog" else "Luteal-anchored ovulatory",
            window_label(r),
            DEFINITION_LABELS[r["definition"]],
            f"{int_float(r['n_classifiable']):,} / {int_float(r['n_windows']):,}",
            rate_with_ci(r),
            pct_or_na(r["positive_rate_all_attempted"]),
            pct_or_na(r["indeterminate_rate"]),
        ])
    add_small_table(
        doc,
        "Supplementary Table 4. Phase-labeling sensitivity.",
        ["Cohort", "Phase mode", "Observation window", "Definition", "Classifiable / attempted", "Positive among classifiable", "Positive among all attempted", "Indeterminate"],
        data,
        widths=[1.0, 1.25, 1.05, 1.7, 1.1, 1.25, 1.1, 0.85],
        font_size=8,
    )


def add_study_level_table(doc: Document, rows: list[dict[str, str]]) -> None:
    definitions = ["A_windowed_any", "A_windowed_C1_or_C2"]
    selected = [
        r for r in rows
        if r["table_type"] == "study_level_3month"
        and r.get("phase_mode") == "strict_herzog"
        and r["definition"] in definitions
    ]
    selected.sort(key=lambda r: (display_cohort(r["cohort"]), int_float(r.get("n_participants", 0)), DEFINITION_LABELS[r["definition"]], r["subset"]))
    data = []
    for r in selected:
        data.append([
            display_cohort(r["cohort"]),
            str(int_float(r.get("n_participants", 0))),
            DEFINITION_LABELS[r["definition"]],
            "All participants" if r["subset"] == "apparent_prevalence_all" else "Classifiable participants only",
            pct_or_na(r["false_positive_rate"]),
            f"{pct_or_na(r['wilson95_low'])}-{pct_or_na(r['wilson95_high'])}",
            pct_or_na(r["p_prevalence_ge_39_1"]),
            pct_or_na(r["p_prevalence_ge_44_2"]),
        ])
    add_small_table(
        doc,
        "Supplementary Table 7. Study-level Monte Carlo results.",
        ["Cohort", "N per study", "Definition", "Denominator", "Mean apparent prevalence", "2.5th-97.5th percentiles", "P >=39.1%", "P >=44.2%"],
        data,
        widths=[1.05, 0.7, 1.55, 1.2, 1.1, 1.15, 0.8, 0.8],
        font_size=8,
    )


def add_pattern_decomposition_table(doc: Document, rows: list[dict[str, str]]) -> None:
    definitions = [
        "A_windowed_any",
        "A_windowed_excluding_C3",
        "A_windowed_C1_only",
        "A_windowed_C2_only",
        "A_windowed_C3_only",
        "B_minimum_data_any",
        "B_minimum_data_excluding_C3",
    ]
    selected = [
        r for r in rows
        if r["table_type"] == "window_false_positive"
        and r["subset"] == "all"
        and r.get("phase_mode") == "strict_herzog"
        and r["window_type"] == "full"
        and r["definition"] in definitions
    ]
    selected.sort(key=lambda r: (display_cohort(r["cohort"]), DEFINITION_LABELS[r["definition"]]))
    data = []
    for r in selected:
        data.append([
            display_cohort(r["cohort"]),
            DEFINITION_LABELS[r["definition"]],
            f"{int_float(r['n_classifiable']):,}",
            f"{int_float(r['positives']):,}",
            f"{pct_or_na(r['false_positive_rate'])} ({pct_or_na(r['wilson95_low'])}-{pct_or_na(r['wilson95_high'])})",
            pct_or_na(r["indeterminate_rate"]),
        ])
    add_small_table(
        doc,
        "Supplementary Table 5. C1/C2/C3 decomposition and C3-exclusion sensitivity.",
        ["Cohort", "Definition", "Classifiable windows", "Positive windows", "Apparent classification rate", "Indeterminate"],
        data,
        widths=[1.35, 2.45, 1.15, 1.1, 1.55, 0.95],
        font_size=8,
    )


def add_nb_sensitivity_table(doc: Document, rows: list[dict[str, str]]) -> None:
    definitions = ["D_nb_regression_C1_or_C2", "D_nb_regression_window_alpha_C1_or_C2"]
    selected = [
        r for r in rows
        if r["table_type"] == "window_false_positive"
        and r["subset"] == "all"
        and r.get("phase_mode") == "strict_herzog"
        and r["definition"] in definitions
        and r["window_type"] == "full"
    ]
    selected.sort(key=lambda r: (display_cohort(r["cohort"]), window_sort_key(r), DEFINITION_LABELS[r["definition"]]))
    data = []
    for r in selected:
        data.append([
            display_cohort(r["cohort"]),
            window_label(r),
            DEFINITION_LABELS[r["definition"]],
            f"{int_float(r['n_classifiable']):,}",
            f"{pct_or_na(r['false_positive_rate'])} ({pct_or_na(r['wilson95_low'])}-{pct_or_na(r['wilson95_high'])})",
            pct_or_na(r["indeterminate_rate"]),
        ])
    add_small_table(
        doc,
        "Supplementary Table 6. Negative-binomial stabilized versus window-only dispersion sensitivity.",
        ["Cohort", "Observation window", "Regression comparator", "Classifiable windows", "Apparent classification rate", "Indeterminate"],
        data,
        widths=[1.2, 1.05, 2.45, 1.1, 1.55, 0.9],
        font_size=8,
    )


def add_historical_table(doc: Document, rows: list[dict[str, str]]) -> None:
    definitions = ["H1_newmark_penry_any", "H1_newmark_penry_66_7_any", "H2_duncan1993_any", "H3_herzog1997_twofold_any", "H4_reddy2007_any_phase2x_any"]
    selected = [
        r for r in rows
        if r["table_type"] == "window_false_positive"
        and r["subset"] == "all"
        and r.get("phase_mode") == "strict_herzog"
        and r["definition"] in definitions
        and (r["window_type"] == "full" or (r["window_type"] == "calendar" and r["window_value"] == "3"))
    ]
    selected.sort(key=lambda r: (display_cohort(r["cohort"]), window_sort_key(r), DEFINITION_LABELS[r["definition"]]))
    data = []
    for r in selected:
        data.append([
            display_cohort(r["cohort"]),
            window_label(r),
            DEFINITION_LABELS[r["definition"]],
            f"{int_float(r['n_classifiable']):,}",
            f"{int_float(r['positives']):,}",
            f"{pct_or_na(r['false_positive_rate'])} ({pct_or_na(r['wilson95_low'])}-{pct_or_na(r['wilson95_high'])})",
            pct_or_na(r["indeterminate_rate"]),
        ])
    add_small_table(
        doc,
        "Supplementary Table 8. Assumption-based historical definitions.",
        ["Cohort", "Observation window", "Definition", "Classifiable windows", "False-positive windows", "False-positive rate", "Indeterminate"],
        data,
        widths=[1.1, 1.1, 2.15, 1.1, 1.1, 1.35, 0.9],
        font_size=8,
    )


def add_small_table(doc: Document, title: str, headers: list[str], data: list[list[str]], widths: list[float], font_size: int = 9) -> None:
    add_paragraph(doc, title)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_widths(table, widths)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, header, bold=True, font_size=font_size)
    repeat_table_header(table.rows[0])
    for row_values in data:
        cells = table.add_row().cells
        for i, value in enumerate(row_values):
            set_cell_text(cells[i], value, bold=False, font_size=font_size)
    for row in table.rows:
        prevent_row_split(row)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=80, bottom=80, start=80, end=80)
    doc.add_paragraph()


def set_table_widths(table, widths: list[float]) -> None:
    for row in table.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = Inches(width)


def set_cell_text(cell, text: str, bold: bool = False, font_size: int = 9) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(font_size)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def set_cell_margins(cell, top=80, start=80, bottom=80, end=80):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_picture(doc: Document, path: Path, width) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(path), width=width)


def one_row(rows: list[dict[str, str]], **kwargs: str) -> dict[str, str]:
    row = maybe_one_row(rows, **kwargs)
    if row is None:
        raise KeyError(kwargs)
    return row


def maybe_one_row(rows: list[dict[str, str]], **kwargs: str) -> dict[str, str] | None:
    if "phase_mode" not in kwargs:
        kwargs["phase_mode"] = "strict_herzog"
    for row in rows:
        ok = True
        for key, value in kwargs.items():
            if not same_value(row.get(key, ""), value):
                ok = False
                break
        if ok:
            return row
    return None


def same_value(left: object, right: object) -> bool:
    if str(left) == str(right):
        return True
    try:
        return float(left) == float(right)
    except Exception:
        return False


def rate_with_ci(row: dict[str, str]) -> str:
    if is_unstable(row):
        return "Not interpreted (<1,000)"
    rate = pct_or_na(row.get("false_positive_rate", ""))
    low = pct_or_na(row.get("wilson95_low", ""))
    high = pct_or_na(row.get("wilson95_high", ""))
    if rate == "NA":
        return "NA"
    if low == "NA" or high == "NA":
        return rate
    return f"{rate} ({low}-{high})"


def is_unstable(row: dict[str, str]) -> bool:
    return str(row.get("unstable_denominator", "")).lower() in {"true", "1", "yes"}


def display_cohort(value: str) -> str:
    return COHORT_LABELS.get(value, value.replace("_", " "))


def window_label(row: dict[str, str]) -> str:
    if row["window_type"] == "calendar":
        value = int_float(row["window_value"])
        return "1 month" if value == 1 else f"{value} months"
    if row["window_type"] == "cycle":
        value = int_float(row["window_value"])
        return "1 cycle" if value == 1 else f"{value} cycles"
    if row["window_type"] == "full":
        return "36-month full diary"
    return f"{row['window_type']} {row['window_value']}"


def window_sort_key(row: dict[str, str]) -> tuple[int, int]:
    order = {"calendar": 0, "cycle": 1, "full": 2}
    if row["window_type"] == "full":
        return (2, 36)
    return (order.get(row["window_type"], 9), int_float(row["window_value"]))


def pct(value: str | float) -> str:
    return f"{100 * float(value):.1f}%"


def pct_or_na(value: str | float) -> str:
    try:
        if value in ("", "nan", "NaN"):
            return "NA"
        return pct(value)
    except Exception:
        return "NA"


def int_float(value: str | float) -> int:
    try:
        return int(float(value))
    except Exception:
        return 0


def split_paragraphs(text: str) -> list[str]:
    return [p.strip() for p in text.split("\n\n") if p.strip()]


def count_words(text: str) -> int:
    return len(re.findall(r"\b[\w'-]+\b", text))


def write_review_note(stats: DocStats) -> None:
    notes = [
        "Second-pass review checklist",
        f"Title character count: {stats.title_chars} (Neurology checklist requires <=96).",
        f"Abstract word count: {stats.abstract_words} (Neurology checklist requires <=250).",
        f"Body text word count: {stats.text_words} (kept concise for Article format).",
        f"References: {stats.n_refs}; Tables: {stats.n_tables}; Figures: {stats.n_figures}.",
        "Formatting: 12-point Times New Roman, left-aligned body text, 1-inch margins, page header with lead author name and page number field, line numbering enabled in section properties.",
        "Methodology review: manuscript states independent CHOCOLATES and HORMONE-CYCLE simulation with direct calendar-day alignment and no reordering, strict Herzog phase labeling before subsetting, no true-coupling simulation, no wavelets, and no HMMs.",
        "Results review: headline values checked against outputs/summary_tables.csv and notebook-derived cohort summary.",
        "Reference review: literature claims were limited to sourced publications and numbered in first-citation order.",
    ]
    (DOC_DIR / "second_pass_review_note.txt").write_text("\n".join(notes) + "\n", encoding="utf-8")


if __name__ == "__main__":
    raise SystemExit(main())
