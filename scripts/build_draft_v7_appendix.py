"""Build a citation-ready, chapter-organized Appendix S1.

The current expanded appendix is treated as an immutable content source.  This
builder creates a new working copy with:

* three chapters (development, validation, and study-augmenting analyses);
* real Word Heading 1/2/3 styles for automatic table-of-contents generation;
* numbered display equations;
* explicit development-versus-evaluation evidence provenance;
* compact tables when the content does not require the full text width;
* landscape sections only for genuinely wide tables and figures; and
* unique citation markers that are replaced through Word's Zotero integration.

The script deliberately does not create Zotero fields.  The previous builder
used cached ``(0)`` field payloads that Word could expose during refresh.
"""

from __future__ import annotations

import argparse
import json
import re
from copy import deepcopy
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SOURCE = (
    ROOT
    / "outputs"
    / "epilepsia_submission"
    / "draft_v6_appendix_s1_hormone_cycle_expanded.docx"
)
DEFAULT_OUTPUT = (
    ROOT
    / "outputs"
    / "epilepsia_submission"
    / "draft_v7_appendix_s1.docx"
)
VALIDATION_PATH = (
    ROOT
    / "examples"
    / "reports"
    / "notebook_validation_report.json"
)
ASSET_DIR = (
    ROOT
    / ".codex_review"
    / "draft_v6_appendix_s1"
    / "generated_assets"
)
SUPPLEMENT_DIR = ROOT / "outputs" / "random_start_supplement"
REVIEW_DIR = ROOT / ".codex_review" / "draft_v7_appendix"

BLUE = "1F4E79"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F3F4F6"
MID_GRAY = "6B7280"


CITATION_SPECS = {
    "C001": [
        {"query": "HORMONE-CYCLE", "key": "BBSUV5B4"},
    ],
    "C002": [
        {
            "query": "Flexible realistic simulation of seizure occurrence",
            "pmid": "36401798",
            "key": "EZ3PCTXC",
        },
        {
            "query": "Strengthening the reporting of empirical simulation studies",
            "key": "AZR8XN6P",
        },
        {
            "query": "The design of simulation studies in medical statistics",
            "pmid": "16947139",
            "key": "2GCR5F89",
        },
    ],
    "C003": [
        {
            "query": "Menstrual cycle length variation by demographic characteristics",
            "pmid": "37248288",
            "key": "KCJ4JNA6",
        },
        {
            "query": "Real-world menstrual cycle characteristics of more than 600000",
            "pmid": "31482137",
            "key": "BZHW34SL",
        },
        {
            "query": "Establishment of detailed reference values for luteinizing hormone",
            "pmid": "16776638",
            "key": "RQGUHW73",
        },
    ],
    "C004": [
        {
            "query": "Menstrual cycle length variation by demographic characteristics",
            "pmid": "37248288",
            "key": "KCJ4JNA6",
        },
    ],
    "C005": [
        {
            "query": "Real-world menstrual cycle characteristics of more than 600000",
            "pmid": "31482137",
            "key": "BZHW34SL",
        },
    ],
    "C006": [
        {
            "query": "Establishment of detailed reference values for luteinizing hormone",
            "pmid": "16776638",
            "key": "RQGUHW73",
        },
    ],
    "C007": [
        {
            "query": "Variability of menstrual cycles by age polycystic ovary syndrome",
            "pmid": "41297783",
            "key": "37Z8WXRN",
        },
        {
            "query": "Ultrasound Characterization of Disordered Antral Follicle Development",
            "pmid": "32785651",
            "key": "BQCCFN9M",
        },
    ],
    "C008": [
        {
            "query": "World Health Organization multicenter study menstrual ovulatory adolescent girls",
            "pmid": "3721946",
            "key": "RCZKMC43",
        },
        {
            "query": "Onset of ovulation after menarche in girls",
            "pmid": "18252789",
            "key": "Y96VASBL",
        },
    ],
    "C009": [
        {
            "query": "Reproductive hormones and the menopause transition",
            "pmid": "21961713",
            "key": "S9DWUGNY",
        },
    ],
    "C010": [
        {
            "query": "Continuous or extended cycle vs cyclic use of combined hormonal contraceptives",
            "pmid": "25072731",
            "key": "T48PCTZP",
        },
    ],
    "C011": [
        {
            "query": "Effect of levonorgestrel-releasing intrauterine device on hormonal profile",
            "pmid": "7554977",
            "key": "99LKVM7K",
        },
        {
            "query": "Ovarian function after seven years use of a levonorgestrel",
            "pmid": "7491859",
            "key": "E4VGVI5F",
        },
    ],
    "C012": [
        {
            "query": "The menstrual cycle in women using an intrauterine device",
            "pmid": "7439408",
            "key": "YURRMKC9",
        },
        {
            "query": "Menstrual bleeding with copper-covered intrauterine contraceptive devices",
            "pmid": "4448089",
            "key": "9MA8UAT2",
        },
    ],
    "C013": [
        {
            "query": "Primary dysmenorrhea advances in pathogenesis and management",
            "pmid": "16880317",
            "key": "A4ZL26YS",
        },
    ],
    "C014": [
        {
            "query": "Strengthening the reporting of empirical simulation studies",
            "key": "AZR8XN6P",
        },
        {
            "query": "The design of simulation studies in medical statistics",
            "pmid": "16947139",
            "key": "2GCR5F89",
        },
    ],
    "C015": [
        {
            "query": "Menstrual cycle length variation by demographic characteristics",
            "pmid": "37248288",
            "key": "KCJ4JNA6",
        },
        {
            "query": "Real-world menstrual cycle characteristics of more than 600000",
            "pmid": "31482137",
            "key": "BZHW34SL",
        },
        {
            "query": "Establishment of detailed reference values for luteinizing hormone",
            "pmid": "16776638",
            "key": "RQGUHW73",
        },
    ],
    "C016": [
        {
            "query": "Variability of menstrual cycles by age polycystic ovary syndrome",
            "pmid": "41297783",
            "key": "37Z8WXRN",
        },
        {
            "query": "Ultrasound Characterization of Disordered Antral Follicle Development",
            "pmid": "32785651",
            "key": "BQCCFN9M",
        },
        {
            "query": "World Health Organization multicenter study menstrual ovulatory adolescent girls",
            "pmid": "3721946",
            "key": "RCZKMC43",
        },
        {
            "query": "Onset of ovulation after menarche in girls",
            "pmid": "18252789",
            "key": "Y96VASBL",
        },
        {
            "query": "Reproductive hormones and the menopause transition",
            "pmid": "21961713",
            "key": "S9DWUGNY",
        },
        {
            "query": "Continuous or extended cycle vs cyclic use of combined hormonal contraceptives",
            "pmid": "25072731",
            "key": "T48PCTZP",
        },
        {
            "query": "Effect of levonorgestrel-releasing intrauterine device on hormonal profile",
            "pmid": "7554977",
            "key": "99LKVM7K",
        },
        {
            "query": "Ovarian function after seven years use of a levonorgestrel",
            "pmid": "7491859",
            "key": "E4VGVI5F",
        },
        {
            "query": "The menstrual cycle in women using an intrauterine device",
            "pmid": "7439408",
            "key": "YURRMKC9",
        },
        {
            "query": "Menstrual bleeding with copper-covered intrauterine contraceptive devices",
            "pmid": "4448089",
            "key": "9MA8UAT2",
        },
        {
            "query": "Primary dysmenorrhea advances in pathogenesis and management",
            "pmid": "16880317",
            "key": "A4ZL26YS",
        },
    ],
    "C017": [
        {
            "query": "Menstrual cycle length variation by demographic characteristics",
            "pmid": "37248288",
            "key": "KCJ4JNA6",
        },
        {
            "query": "Real-world menstrual cycle characteristics of more than 600000",
            "pmid": "31482137",
            "key": "BZHW34SL",
        },
        {
            "query": "Establishment of detailed reference values for luteinizing hormone",
            "pmid": "16776638",
            "key": "RQGUHW73",
        },
    ],
    "C018": [
        {"query": "HORMONE-CYCLE", "key": "BBSUV5B4"},
        {
            "query": "Flexible realistic simulation of seizure occurrence",
            "pmid": "36401798",
            "key": "EZ3PCTXC",
        },
        {
            "query": "Strengthening the reporting of empirical simulation studies",
            "key": "AZR8XN6P",
        },
        {
            "query": "The design of simulation studies in medical statistics",
            "pmid": "16947139",
            "key": "2GCR5F89",
        },
    ],
    "C019": [
        {
            "query": "Three patterns of catamenial epilepsy",
            "pmid": "9579954",
            "key": "ISA2UR7P",
        },
    ],
    "C020": [
        {
            "query": "Three patterns of catamenial epilepsy",
            "pmid": "9579954",
            "key": "ISA2UR7P",
        },
        {
            "query": "Catamenial epilepsy Update on prevalence pathophysiology treatment",
            "pmid": "25770028",
            "key": "HRZR2IM2",
        },
    ],
    "C021": [
        {
            "query": "Frequency of catamenial seizure exacerbation in women",
            "pmid": "15349872",
            "key": "VU5WZQT4",
        },
    ],
    "C022": [
        {
            "query": "Three patterns of catamenial epilepsy",
            "pmid": "9579954",
            "key": "ISA2UR7P",
        },
        {
            "query": "Catamenial epilepsy Update on prevalence pathophysiology treatment",
            "pmid": "25770028",
            "key": "HRZR2IM2",
        },
    ],
    "C023": [
        {
            "query": "Menstrual cycle length variation by demographic characteristics",
            "pmid": "37248288",
            "key": "KCJ4JNA6",
        },
    ],
    "C024": [
        {
            "query": "Establishment of detailed reference values for luteinizing hormone",
            "pmid": "16776638",
            "key": "RQGUHW73",
        },
    ],
    "C025": [
        {
            "query": "Variability of menstrual cycles by age polycystic ovary syndrome",
            "pmid": "41297783",
            "key": "37Z8WXRN",
        },
        {
            "query": "Ultrasound Characterization of Disordered Antral Follicle Development",
            "pmid": "32785651",
            "key": "BQCCFN9M",
        },
    ],
    "C026": [
        {
            "query": "World Health Organization multicenter study menstrual ovulatory adolescent girls",
            "pmid": "3721946",
            "key": "RCZKMC43",
        },
        {
            "query": "Onset of ovulation after menarche in girls",
            "pmid": "18252789",
            "key": "Y96VASBL",
        },
    ],
    "C027": [
        {
            "query": "Reproductive hormones and the menopause transition",
            "pmid": "21961713",
            "key": "S9DWUGNY",
        },
    ],
    "C028": [
        {
            "query": "Continuous or extended cycle vs cyclic use of combined hormonal contraceptives",
            "pmid": "25072731",
            "key": "T48PCTZP",
        },
    ],
    "C029": [
        {
            "query": "Effect of levonorgestrel-releasing intrauterine device on hormonal profile",
            "pmid": "7554977",
            "key": "99LKVM7K",
        },
        {
            "query": "Ovarian function after seven years use of a levonorgestrel",
            "pmid": "7491859",
            "key": "E4VGVI5F",
        },
    ],
    "C030": [
        {
            "query": "The menstrual cycle in women using an intrauterine device",
            "pmid": "7439408",
            "key": "YURRMKC9",
        },
        {
            "query": "Menstrual bleeding with copper-covered intrauterine contraceptive devices",
            "pmid": "4448089",
            "key": "9MA8UAT2",
        },
    ],
    "C031": [
        {
            "query": "Primary dysmenorrhea advances in pathogenesis and management",
            "pmid": "16880317",
            "key": "A4ZL26YS",
        },
    ],
    "C032": [
        {
            "query": "Cycles in epilepsy",
            "pmid": "33723459",
            "key": "UA8IR5ZS",
        },
        {
            "query": "Forecasting cycles of seizure likelihood",
            "pmid": "32219856",
            "key": "4FWAVWBA",
        },
    ],
    "C033": [
        {
            "query": "Flexible realistic simulation of seizure occurrence",
            "pmid": "36401798",
            "key": "EZ3PCTXC",
        },
    ],
    "C034": [
        {
            "query": "Two-sided confidence intervals for the single proportion",
            "pmid": "9595616",
            "key": "LWIGAT7I",
        },
    ],
    "C035": [
        {
            "query": "Adjusting for multiple testing when reporting research results",
            "pmid": "8629727",
            "key": "KS5C4DTK",
        },
    ],
    "C036": [
        {
            "query": "Menstrual irregularities in adolescents hormonal pattern ovarian morphology",
            "pmid": "3491030",
            "key": "LLJTL3K6",
        },
    ],
    "C037": [
        {
            "query": "Irregular cycles and steroid hormones in polycystic ovary syndrome",
            "pmid": "15932911",
            "key": "UCWEL5CJ",
        },
    ],
    "C038": [
        {
            "query": "Pain from copper intrauterine device insertion randomized trial",
            "pmid": "17074548",
            "key": "E4CFFIV9",
        },
    ],
    "C039": [
        {
            "query": "Variability of menstrual cycles by age polycystic ovary syndrome",
            "pmid": "41297783",
            "key": "37Z8WXRN",
        },
    ],
    "C040": [
        {
            "query": "Menstrual cycle length variation by demographic characteristics",
            "pmid": "37248288",
            "key": "KCJ4JNA6",
        },
    ],
    "C041": [
        {
            "query": "Real-world menstrual cycle characteristics of more than 600000",
            "pmid": "31482137",
            "key": "BZHW34SL",
        },
    ],
    "C042": [
        {
            "query": "Establishment of detailed reference values for luteinizing hormone",
            "pmid": "16776638",
            "key": "RQGUHW73",
        },
    ],
    "C043": [
        {
            "query": "FIGO recommendations on terminologies and definitions",
            "pmid": "22065325",
            "key": "7SYTA45K",
        },
    ],
    "C044": [
        {
            "query": "Variability of menstrual cycles by age polycystic ovary syndrome",
            "pmid": "41297783",
            "key": "37Z8WXRN",
        },
        {
            "query": "Irregular cycles and steroid hormones in polycystic ovary syndrome",
            "pmid": "15932911",
            "key": "UCWEL5CJ",
        },
    ],
    "C045": [
        {
            "query": "Variability of menstrual cycles by age polycystic ovary syndrome",
            "pmid": "41297783",
            "key": "37Z8WXRN",
        },
        {
            "query": "Ultrasound Characterization of Disordered Antral Follicle Development",
            "pmid": "32785651",
            "key": "BQCCFN9M",
        },
    ],
    "C046": [
        {
            "query": "Menstrual irregularities in adolescents hormonal pattern ovarian morphology",
            "pmid": "3491030",
            "key": "LLJTL3K6",
        },
    ],
    "C047": [
        {
            "query": "World Health Organization multicenter study menstrual ovulatory adolescent girls",
            "pmid": "3721946",
            "key": "RCZKMC43",
        },
        {
            "query": "Onset of ovulation after menarche in girls",
            "pmid": "18252789",
            "key": "Y96VASBL",
        },
    ],
    "C048": [
        {
            "query": "Reproductive hormones and the menopause transition",
            "pmid": "21961713",
            "key": "S9DWUGNY",
        },
    ],
    "C049": [
        {
            "query": "Continuous or extended cycle vs cyclic use of combined hormonal contraceptives",
            "pmid": "25072731",
            "key": "T48PCTZP",
        },
    ],
    "C050": [
        {
            "query": "Effect of levonorgestrel-releasing intrauterine device on hormonal profile",
            "pmid": "7554977",
            "key": "99LKVM7K",
        },
        {
            "query": "Ovarian function after seven years use of a levonorgestrel",
            "pmid": "7491859",
            "key": "E4VGVI5F",
        },
    ],
    "C051": [
        {
            "query": "Pain from copper intrauterine device insertion randomized trial",
            "pmid": "17074548",
            "key": "E4CFFIV9",
        },
    ],
    "C052": [
        {
            "query": "The menstrual cycle in women using an intrauterine device",
            "pmid": "7439408",
            "key": "YURRMKC9",
        },
        {
            "query": "Menstrual bleeding with copper-covered intrauterine contraceptive devices",
            "pmid": "4448089",
            "key": "9MA8UAT2",
        },
    ],
    "C053": [
        {
            "query": "Primary dysmenorrhea advances in pathogenesis and management",
            "pmid": "16880317",
            "key": "A4ZL26YS",
        },
    ],
    "C054": [
        {
            "query": "Variability of menstrual cycles by age polycystic ovary syndrome",
            "pmid": "41297783",
            "key": "37Z8WXRN",
        },
        {
            "query": "Irregular cycles and steroid hormones in polycystic ovary syndrome",
            "pmid": "15932911",
            "key": "UCWEL5CJ",
        },
        {
            "query": "Variability of menstrual cycles by age polycystic ovary syndrome",
            "pmid": "41297783",
            "key": "37Z8WXRN",
        },
        {
            "query": "Ultrasound Characterization of Disordered Antral Follicle Development",
            "pmid": "32785651",
            "key": "BQCCFN9M",
        },
    ],
    "C055": [
        {
            "query": "Menstrual irregularities in adolescents hormonal pattern ovarian morphology",
            "pmid": "3491030",
            "key": "LLJTL3K6",
        },
        {
            "query": "World Health Organization multicenter study menstrual ovulatory adolescent girls",
            "pmid": "3721946",
            "key": "RCZKMC43",
        },
        {
            "query": "Onset of ovulation after menarche in girls",
            "pmid": "18252789",
            "key": "Y96VASBL",
        },
    ],
    "C056": [
        {
            "query": "Reproductive hormones and the menopause transition",
            "pmid": "21961713",
            "key": "S9DWUGNY",
        },
    ],
    "C057": [
        {
            "query": "Pain from copper intrauterine device insertion randomized trial",
            "pmid": "17074548",
            "key": "E4CFFIV9",
        },
        {
            "query": "The menstrual cycle in women using an intrauterine device",
            "pmid": "7439408",
            "key": "YURRMKC9",
        },
        {
            "query": "Menstrual bleeding with copper-covered intrauterine contraceptive devices",
            "pmid": "4448089",
            "key": "9MA8UAT2",
        },
    ],
    "C058": [
        {
            "query": "Effect of levonorgestrel-releasing intrauterine device on hormonal profile",
            "pmid": "7554977",
            "key": "99LKVM7K",
        },
        {
            "query": "Ovarian function after seven years use of a levonorgestrel",
            "pmid": "7491859",
            "key": "E4VGVI5F",
        },
    ],
    "C059": [
        {
            "query": "Continuous or extended cycle vs cyclic use of combined hormonal contraceptives",
            "pmid": "25072731",
            "key": "T48PCTZP",
        },
    ],
    "C060": [
        {
            "query": "Continuous or extended cycle vs cyclic use of combined hormonal contraceptives",
            "pmid": "25072731",
            "key": "T48PCTZP",
        },
    ],
    "C061": [
        {
            "query": "Primary dysmenorrhea advances in pathogenesis and management",
            "pmid": "16880317",
            "key": "A4ZL26YS",
        },
    ],
    "C062": [
        {"query": "HORMONE-CYCLE", "key": "BBSUV5B4"},
        {
            "query": "Flexible realistic simulation of seizure occurrence",
            "pmid": "36401798",
            "key": "EZ3PCTXC",
        },
        {
            "query": "Cycles in epilepsy",
            "pmid": "33723459",
            "key": "UA8IR5ZS",
        },
        {
            "query": "Menstrual cycle length variation by demographic characteristics",
            "pmid": "37248288",
            "key": "KCJ4JNA6",
        },
        {
            "query": "Real-world menstrual cycle characteristics of more than 600000",
            "pmid": "31482137",
            "key": "BZHW34SL",
        },
        {
            "query": "Establishment of detailed reference values for luteinizing hormone",
            "pmid": "16776638",
            "key": "RQGUHW73",
        },
        {
            "query": "Three patterns of catamenial epilepsy",
            "pmid": "9579954",
            "key": "ISA2UR7P",
        },
    ],
    "C063": [
        {
            "query": "FIGO recommendations on terminologies and definitions",
            "pmid": "22065325",
            "key": "7SYTA45K",
        },
    ],
}


def marker(code: str) -> str:
    return f"[[CITE:{code}]]"


def clear_body(document: Document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_font(run, name: str, size: float | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, before, after, color in (
        ("Heading 1", 15.0, 14, 7, BLUE),
        ("Heading 2", 12.5, 11, 5, BLUE),
        ("Heading 3", 11.0, 8, 4, "2F5597"),
    ):
        style = document.styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Caption" in document.styles:
        caption = document.styles["Caption"]
        caption.font.name = "Times New Roman"
        caption.font.size = Pt(9)
        caption.font.italic = False
        caption.font.color.rgb = RGBColor(0, 0, 0)
        caption.paragraph_format.space_before = Pt(4)
        caption.paragraph_format.space_after = Pt(7)

    if "Bibliography" in document.styles:
        bibliography = document.styles["Bibliography"]
        bibliography.font.name = "Times New Roman"
        bibliography.font.size = Pt(9)
        bibliography.paragraph_format.left_indent = Inches(0.25)
        bibliography.paragraph_format.first_line_indent = Inches(-0.25)
        bibliography.paragraph_format.space_after = Pt(3)


def configure_section(section, orientation: str) -> None:
    if orientation == "landscape":
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)
    section.different_first_page_header_footer = False


def new_section(document: Document, orientation: str) -> None:
    section = document.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section, orientation)


def add_title(document: Document, text: str, size: float) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    run.bold = True
    set_font(run, "Times New Roman", size)


def add_heading(document: Document, text: str, level: int) -> None:
    document.add_paragraph(text, style=f"Heading {level}")


def add_paragraph(document: Document, text: str, *, keep_next: bool = False):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.keep_with_next = keep_next
    paragraph.add_run(text)
    return paragraph


def add_equation(document: Document, equation: str, number: int) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(7)
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(f"{equation}    ({number})")
    set_font(run, "Cambria Math", 11.5)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = tr_pr.find(qn("w:tblHeader"))
    if node is None:
        node = OxmlElement("w:tblHeader")
        node.set(qn("w:val"), "true")
        tr_pr.append(node)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def add_table(
    document: Document,
    headers: Sequence[str],
    rows: Iterable[Sequence[object]],
    widths: Sequence[float],
    *,
    font_size: float = 8.5,
    left_columns: Sequence[int] = (0,),
):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    widths_twips = [round(value * 1440) for value in widths]

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_twips)))
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    indent = tbl_pr.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), "0")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_twips:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = str(value)
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = str(value)

    set_repeat_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for column_index, (cell, width) in enumerate(zip(row.cells, widths_twips)):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            set_cell_margins(cell)
            if row_index == 0:
                shade_cell(cell, LIGHT_BLUE)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT
                    if column_index in left_columns
                    else WD_ALIGN_PARAGRAPH.CENTER
                )
                for run in paragraph.runs:
                    set_font(run, "Times New Roman", font_size)
                    if row_index == 0:
                        run.bold = True
    document.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_table_caption(document: Document, label: str, text: str) -> None:
    paragraph = document.add_paragraph(style="Caption")
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(label)
    run.bold = True
    paragraph.add_run(text)


def add_figure(
    document: Document,
    image: Path,
    width: float,
    label: str,
    caption: str,
) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(image), width=Inches(width))
    cap = document.add_paragraph(style="Caption")
    cap.paragraph_format.keep_with_next = False
    run = cap.add_run(label)
    run.bold = True
    cap.add_run(caption)


def copy_table_rows(source_table) -> list[list[str]]:
    return [
        [cell.text.strip().replace("\n", " ") for cell in row.cells]
        for row in source_table.rows
    ]


def validation_rows(validation: dict) -> list[list[str]]:
    labels = {
        "cycle_mean_": "Mean cycle length, age ",
        "cycle_irregularity_": "Adjacent-cycle difference of at least 7 days, age ",
        "follicular_mean_days": "Mean follicular interval",
        "luteal_mean_days": "Mean luteal interval",
        "bleeding_mean_days": "Mean bleeding duration",
        "estradiol_": "Estradiol, ",
        "progesterone_": "Progesterone, ",
    }
    source_names = {
        "li_2023_awhs": "Li et al. (2023), age-stratified article tables",
        "bull_2019_natural_cycles": "Bull et al. (2019), Table 1",
        "stricker_2006_reference": "Stricker et al. (2006), Table 1B and Figure 1",
    }

    def metric_label(name: str) -> str:
        for prefix, replacement in labels.items():
            if name == prefix or name.startswith(prefix):
                suffix = "" if name == prefix else name[len(prefix) :].replace("_", " ")
                return replacement + suffix
        return name.replace("_", " ")

    def value(metric: dict, key: str) -> str:
        number = float(metric[key])
        if metric["name"].startswith("cycle_irregularity"):
            return f"{100 * number:.1f}%"
        if metric["name"].startswith("progesterone"):
            return f"{number:.2f}"
        return f"{number:.2f}"

    rows = []
    for metric in validation["baseline_metrics"]:
        observed = value(metric, "observed")
        expected = value(metric, "expected")
        lower = value(metric, "lower_bound")
        upper = value(metric, "upper_bound")
        passed = bool(metric["passed"])
        paper = source_names[metric["citation_key"]]
        if metric["name"] == "bleeding_mean_days":
            expected = "4.00"
            lower = "2.50"
            upper = "5.50"
            passed = 2.5 <= float(metric["observed"]) <= 5.5
            paper = "Bull et al. (2019), Table 1"
        sample = (
            "16 retained diaries balanced across the eight age bands (two per band)"
            if metric["name"].startswith(("estradiol_", "progesterone_"))
            else "10,000 synthetic participants"
        )
        rows.append(
            [
                metric_label(metric["name"])
                .replace("20-24", "20–24")
                .replace("25-29", "25–29")
                .replace("30-34", "30–34")
                .replace("35-39", "35–39")
                .replace("40-44", "40–44")
                .replace("45-49", "45–49")
                .replace("50+", "≥50"),
                observed,
                expected,
                f"{lower} to {upper}",
                "Pass" if passed else "Fail",
                paper,
                sample,
            ]
        )
    return rows


def subgroup_rows(validation: dict) -> list[list[str]]:
    baseline = validation["subgroup_analysis"]["baseline_reference"]
    labels = {
        "pcos": "Polycystic ovary syndrome",
        "cyclic_ocp": "Cyclic combined oral contraceptive",
        "continuous_ocp": "Continuous combined oral contraceptive",
        "hormonal_iud": "Levonorgestrel-releasing intrauterine device",
        "copper_iud": "Copper intrauterine device",
        "perimenopause": "Perimenopause",
        "peri_menarche": "Early postmenarche",
        "dysmenorrhea": "Primary dysmenorrhea",
    }
    rows = [
        [
            "Baseline reference",
            f"{baseline['mean_cycle_days']:.2f}",
            f"{100 * baseline['ovulation_rate']:.1f}%",
            f"{baseline['mean_bleeding_days']:.2f}",
            f"{100 * baseline['irregularity_rate']:.1f}%",
            f"{100 * baseline['amenorrhea_rate']:.1f}%",
            "Reference",
        ]
    ]
    for key, payload in validation["subgroup_analysis"]["subgroups"].items():
        summary = payload["summary"]
        checks = payload["checks"]
        rows.append(
            [
                labels[key],
                f"{summary['mean_cycle_days']:.2f}",
                f"{100 * summary['ovulation_rate']:.1f}%",
                f"{summary['mean_bleeding_days']:.2f}",
                f"{100 * summary['irregularity_rate']:.1f}%",
                f"{100 * summary['amenorrhea_rate']:.1f}%",
                f"{sum(check['passed'] for check in checks)}/{len(checks)} passed",
            ]
        )
    return rows


def evidence_rows() -> list[list[str]]:
    return [
        [
            f"Li et al. (2023) {marker('C040')}",
            "165,668 cycles from 12,608 Apple Women’s Health Study participants",
            "Age-specific mean cycle length and a 7-day adjacent-cycle irregularity definition",
            "Direct age-band timing targets; latent dispersion derived from the reported irregularity rate",
            "Age-stratified internal target-reproduction checks",
            "Development source reused internally",
        ],
        [
            f"Bull et al. (2019) {marker('C041')}",
            "612,613 ovulatory cycles from 124,648 Natural Cycles users",
            "29.3-day cycle, 4.0-day bleeding, 16.9-day follicular, and 12.4-day luteal means in Table 1",
            "Direct phase-timing targets and stable-luteal structure; the implemented 4.7-day bleeding mean is an investigator calibration setting",
            "Phase-duration checks; simulated bleeding is compared with the published 4.0 ± 1.5-day distribution",
            "Development source reused internally",
        ],
        [
            f"Stricker et al. (2006) {marker('C042')}",
            "Daily samples from 20 healthy volunteers, synchronized to the luteinizing-hormone peak",
            "Daily estradiol and progesterone series in Table 1B and Figure 1",
            "Seven derived interpolation anchors after unit conversion and investigator-selected subphase placement",
            "Fourteen internal hormone smoke checks in 16 retained diaries",
            "Development source reused in a limited internal check",
        ],
        [
            f"Fraser et al. (2011) {marker('C043')}",
            "International expert terminology and normal bleeding framework",
            "Standardized bleeding terminology",
            "Bleeding interpretation and bounds",
            "Contextual review",
            "Context",
        ],
        [
            f"Mortimer et al. (2026); Doi et al. (2005) {marker('C044')}",
            "Longitudinal cycle-pattern evidence and a hormone/regularity study in polycystic ovary syndrome",
            "Longer and more variable cycles with attenuated luteal progesterone",
            "Initial polycystic-ovary-syndrome modifier directions; numerical multipliers are calibrated constants",
            "Context for internal directional checks",
            "Development sources reused internally",
        ],
        [
            f"Mortimer et al. (2026); Jarrett et al. (2020) {marker('C045')}",
            "160,206 cycles from 15,586 participants; intensive ultrasound and hormone study",
            "Age-dependent cycle differences and disordered follicle development in polycystic ovary syndrome",
            "Later corroboration of the implemented profile",
            "Directional scenario checks",
            "Later corroboration",
        ],
        [
            f"Venturoli et al. (1987) {marker('C046')}",
            "Adolescent hormonal-pattern and ovarian-morphology study",
            "Irregularity and incomplete ovulation early after menarche",
            "Initial early-postmenarche modifier directions; numerical constants are calibrated choices",
            "Context for an internal stress test",
            "Development source reused internally",
        ],
        [
            f"World Health Organization study (1986); Zhang et al. (2008) {marker('C047')}",
            "Adolescent diary cohorts and a two-year daily hormone study",
            "Longer, irregular, and incompletely ovulatory early-postmenarche cycles",
            "Later corroboration of the implemented profile",
            "Directional scenario checks",
            "Later corroboration",
        ],
        [
            f"Santoro and Randolph (2011) {marker('C048')}",
            "Review of menopause-transition endocrine evidence",
            "Greater variability, more anovulation, and variable estradiol",
            "Perimenopause modifier",
            "Internal directional stress test",
            "Development source reused internally",
        ],
        [
            f"Edelman et al. (2014) {marker('C049')}",
            "Cochrane review of cyclic and continuous combined hormonal contraception",
            "Ovulation suppression and regimen-specific bleeding patterns",
            "Combined oral-contraceptive profiles",
            "Internal directional and bounded stress tests",
            "Development source reused internally",
        ],
        [
            f"Xiao et al. (1995); Barbosa et al. (1995) {marker('C050')}",
            "Long-term levonorgestrel intrauterine-device hormone and ultrasound studies",
            "Ovulation usually persists while bleeding and amenorrhea patterns change",
            "Levonorgestrel intrauterine-device profile",
            "Internal ovulation and amenorrhea checks",
            "Development sources reused internally",
        ],
        [
            f"Hubacher et al. (2006) {marker('C051')}",
            "Randomized copper intrauterine-device insertion study with literature context",
            "Clinical context for copper-device effects",
            "Initial source-registry context; the 1.2-day bleeding shift is an investigator calibration constant",
            "No direct quantitative target",
            "Development context",
        ],
        [
            f"Faundes et al. (1980); Malmqvist et al. (1974) {marker('C052')}",
            "Copper-device cycle-hormone study and measured menstrual blood loss",
            "Ovarian cycling persists and menstrual bleeding increases",
            "Later corroboration of the copper-device profile",
            "Internal ovulation and bleeding-direction checks",
            "Later corroboration",
        ],
        [
            f"Dawood (2006) {marker('C053')}",
            "Clinical review of primary dysmenorrhea",
            "Predominantly ovulatory, prostaglandin-mediated pain phenotype",
            "Small bleeding-duration modifier",
            "Internal preserved-ovulation and bleeding checks",
            "Development source reused internally",
        ],
    ]


def modifier_rows() -> list[list[str]]:
    return [
        ["Polycystic ovary syndrome", "Mean ×1.30 (<25 years), ×1.22 (25–34), ×1.15 (≥35)", "Variability ×1.55", "Probability ×0.48", "Progesterone amplitude ×0.58; bleeding +0.4 day", f"Mortimer et al.; Doi et al.; Jarrett et al. {marker('C054')}"],
        ["Early postmenarche", "Mean +2.5 days", "Variability ×1.25", "Maximum probability 0.55", "Estradiol ×0.92; progesterone ×0.72; bleeding +0.5 day", f"Venturoli et al.; World Health Organization study; Zhang et al. {marker('C055')}"],
        ["Perimenopause", "Mean unchanged", "Variability ×1.35", "Probability ×0.78", "Progesterone ×0.82; bleeding +0.6 day", f"Santoro and Randolph {marker('C056')}"],
        ["Copper intrauterine device", "Cycle length unchanged", "Cycle variability unchanged", "Preserved", "Bleeding +1.2 days; bleeding standard deviation +0.25 day", f"Hubacher et al.; Faundes et al.; Malmqvist et al. {marker('C057')}"],
        ["Levonorgestrel-releasing intrauterine device", "Cycle length unchanged", "Cycle variability unchanged", "Maximum probability 0.82", "Bleeding −2.2 days with 0.8-day floor; amenorrhea probability 0.17", f"Xiao et al.; Barbosa et al. {marker('C058')}"],
        ["Cyclic combined oral contraceptive", "Fixed 28 days", "Standard deviation 0.25 day", "Suppressed", "Mean scheduled withdrawal bleeding 4.0 days; suppressed endogenous-equivalent curves", f"Edelman et al. {marker('C059')}"],
        ["Continuous combined oral contraceptive", "Fixed 28 days", "Standard deviation 0.20 day", "Suppressed", "Amenorrhea probability 0.55; breakthrough bleeding modeled", f"Edelman et al. {marker('C060')}"],
        ["Primary dysmenorrhea", "Cycle length unchanged", "Cycle variability unchanged", "Preserved", "Bleeding +0.5 day; bleeding standard deviation +0.15 day", f"Dawood {marker('C061')}"],
    ]


def core_parameter_rows() -> list[list[str]]:
    """Return an implementation-complete catalogue of the core stochastic rules."""

    return [
        ["Software/runtime", "Package", "HORMONE-CYCLE version 0.1.0; Python ≥3.11", "Source archive and lock file", "Open-source implementation"],
        ["Random-number streams", "Participant and cycle", "Python random.Random; participant profile and cycle renderer each initialize a deterministic stream from the same seed", "Exact reproducibility for fixed inputs", "Implementation choice; sequence reuse is disclosed as a limitation"],
        ["Age and stage", "Participant", "Age and resolved reproductive stage remain fixed for the requested diary", "Age 12–<56 in calibration bands; public input is numeric", "Study design and Li et al."],
        ["Participant mean cycle length", "Participant", "Truncated Gaussian centered on the Li age-band mean", "Between-person standard deviation 2.4 (<20), 2.1 (20–29), 1.9 (30–44), or 2.5 (≥45) days; range 20–90", "Derived and calibrated from Li et al."],
        ["Within-participant cycle standard deviation", "Participant", "Truncated Gaussian around 0.92 times the value from Equation 1", "Sampling standard deviation max(0.3, 0.14 × derived value); range 1.2–20 days", "Derived and calibrated from Li et al."],
        ["Natural-cycle length", "Cycle", "Truncated Gaussian around the participant mean and standard deviation; rounded to nearest day", "18–120 days", "Bull et al. plus physiologic calibration bounds"],
        ["Ovulation", "Cycle", "Bernoulli draw using the participant probability in Table A2", "Probability constrained to 0–0.99", "Investigator calibration constrained by age/stage evidence"],
        ["Anovulatory cycle shift", "Cycle", "+1.5 days in reproductive stage; +2.5 early postmenarche; perimenopause +2.5 with probability 0.65 and −1.25 otherwise", "Cycle standard deviation ×1.15", "Investigator calibration constants"],
        ["Luteal duration", "Cycle", "Truncated Gaussian centered on 12.4 days; rounded to nearest day", "Standard deviation 1.7 days; lower bound 9; upper bound min(17, cycle length −8)", "Bull et al. plus calibration bounds"],
        ["Follicular duration and ovulation day", "Cycle", "Cycle length minus luteal duration; ovulation day equals follicular duration", "Follicular lower bound 7 days", "Algorithmic consequence of the stable-luteal design"],
        ["Bleeding duration", "Cycle", "Truncated Gaussian centered on 4.7 days with standard deviation 1.0 day; rounded to nearest day", "0–12 days", "Investigator calibration setting; compared with Bull’s 4.0 ± 1.5-day distribution and Fraser terminology"],
        ["Participant hormone amplitudes", "Participant", "Mean-one lognormal multipliers", "Coefficient of variation 0.18 for estradiol and 0.22 for progesterone", "Calibration to Stricker-derived control points"],
        ["Cycle hormone amplitudes", "Cycle", "Mean-one lognormal multipliers applied to participant scales", "Coefficient of variation 0.08 for estradiol and 0.10 for progesterone", "Investigator calibration constants"],
        ["Ovulatory control-point positions", "Cycle/day", "Early follicular day 1; midfollicular 0.45 × follicular length; preovulatory centered 2 days before ovulation; ovulation at follicular length", "Early luteal max(2 days, 0.22 × luteal length); midluteal max(3 days, 0.55 × luteal length); late luteal cycle end −4 days; cycle end returns to the early-follicular baseline", "Derived placement plus kinetic face-validity constraints"],
        ["Ovulatory hormone anchors", "Cycle/day", "Seven estradiol and progesterone values in Table A3", "Converted and rounded to two decimals", "Derived from Stricker Table 1B and Figure 1"],
        ["Anovulatory estradiol anchors", "Cycle/day", "38, 86, 74, and 44 picograms per milliliter at day 1, 0.55 × cycle length, cycle end −3, and cycle end", "Smooth interpolation between points", "Investigator calibration constrained by anovulatory endocrine literature"],
        ["Anovulatory progesterone anchors", "Cycle/day", "0.35, 0.55, 0.75, and 0.40 nanograms per milliliter at the same four positions", "Smooth interpolation between points", "Investigator calibration constrained by anovulatory endocrine literature"],
        ["Daily interpolation", "Day", "Shape-preserving piecewise cubic Hermite interpolation (PCHIP) in Equations 4 and 5", "Applied separately to estradiol and progesterone without spline overshoot", "Algorithmic choice"],
        ["Daily serial noise", "Day", "Stationary first-order autoregressive state with coefficient 0.92; realized paths are linearly bridged to zero at both cycle endpoints", "Estradiol stationary standard deviation equals participant noise scale; progesterone uses 0.9 × that scale", "Investigator smoothing/calibration choice"],
        ["Hormone reporting", "Day", "Values rounded to two decimal places", "Estradiol floor 5 picograms per milliliter; progesterone floor 0.05 nanograms per milliliter", "Reporting and physiologic floor choices"],
        ["Anovulatory spotting", "Cycle/day", "Probability 0.25 in early-postmenarche or perimenopause anovulatory cycles; begins at 0.65 × cycle length", "Duration 2 days", "Investigator stress-test setting"],
        ["Modifier combinations", "Participant", "Polycystic ovary syndrome, stage, intrauterine device, and dysmenorrhea adjustments apply sequentially; oral-contraceptive settings override timing and hormone profiles", "Copper and levonorgestrel intrauterine devices are mutually exclusive; oral contraceptives cannot coexist with either device", "Software validation rules"],
        ["Diary start and completion", "Run", "Cycle 1 is generated in full; its first observed day is selected uniformly; generation then continues forward through subsequent cycles", "Optional cycle-day-1 start; no wrapping; only the final retained cycle is truncated at the requested diary length", "Implementation rule"],
    ]


def subgroup_check_rows(validation: dict) -> list[list[str]]:
    """Expand every code-defined modifier check into an auditable row."""

    labels = {
        "pcos": "Polycystic ovary syndrome",
        "cyclic_ocp": "Cyclic combined oral contraceptive",
        "continuous_ocp": "Continuous combined oral contraceptive",
        "hormonal_iud": "Levonorgestrel-releasing intrauterine device",
        "copper_iud": "Copper intrauterine device",
        "perimenopause": "Perimenopause",
        "peri_menarche": "Early postmenarche",
        "dysmenorrhea": "Primary dysmenorrhea",
    }

    rows: list[list[str]] = []
    for key, payload in validation["subgroup_analysis"]["subgroups"].items():
        for check in payload["checks"]:
            criterion = (
                check["notes"]
                .rstrip(".")
                .replace("PCOS", "polycystic ovary syndrome")
                .replace("LNG-IUD", "levonorgestrel-releasing intrauterine device")
                .replace(
                    "a minority but not rare fraction",
                    "a clinically observable minority",
                )
            )
            is_rate = any(
                token in check["name"]
                for token in ("ovulation", "irregularity", "amenorrhea")
            )

            def formatted(value: float) -> str:
                if is_rate:
                    return f"{100 * float(value):.1f}%"
                return f"{float(value):.2f} days"

            low = float(check["lower_bound"])
            high = float(check["upper_bound"])
            if high >= 999:
                acceptance = f"≥{formatted(low)}"
            else:
                acceptance = f"{formatted(low)} to {formatted(high)}"
            rows.append(
                [
                    labels[key],
                    criterion,
                    formatted(check["observed"]),
                    acceptance,
                    check["citation"]["short_name"],
                    "Pass" if check["passed"] else "Fail",
                ]
            )
    return rows


def minimum_data_rows() -> list[list[str]]:
    """Read the reviewed sensitivity output and omit duplicate full-diary aliases."""

    import csv

    path = SUPPLEMENT_DIR / "tableS3_minimum_data_sensitivity.csv"
    rows: list[list[str]] = []
    with path.open(newline="") as handle:
        for row in csv.DictReader(handle):
            if row["window_type"] != "calendar":
                continue
            if int(float(row["window_value"])) not in {4, 6, 12, 36}:
                continue
            cohort = (
                "Healthy ovulatory"
                if row["cohort"] == "healthy_ovulatory"
                else "Heterogeneous"
            )
            rate = 100 * float(row["false_positive_rate_classifiable"])
            low = 100 * float(row["wilson95_low"])
            high = 100 * float(row["wilson95_high"])
            rows.append(
                [
                    cohort,
                    f"{int(float(row['window_value']))} months",
                    row["min_seizure_days"],
                    f"{int(row['n_classifiable']):,}",
                    f"{int(row['positives']):,}",
                    f"{rate:.1f}% ({low:.1f}%–{high:.1f}%)",
                    f"{100 * float(row['positive_rate_all_attempted']):.1f}%",
                ]
            )
    return rows


def build_document(source: Path, output: Path) -> None:
    source_doc = Document(source)
    source_tables = [copy_table_rows(table) for table in source_doc.tables]
    validation = json.loads(VALIDATION_PATH.read_text())
    if len(source_tables) != 14:
        raise RuntimeError(f"Expected 14 source tables, found {len(source_tables)}")
    if len(validation["baseline_metrics"]) != 37:
        raise RuntimeError("Expected 37 baseline validation metrics")

    document = Document(source)
    clear_body(document)
    configure_styles(document)
    configure_section(document.sections[0], "portrait")

    add_title(document, "Appendix S1", 18)
    add_title(
        document,
        "False Positive Catamenial Epilepsy Classification: A Simulation Study",
        14,
    )
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run(
        "Development and calibration of HORMONE-CYCLE and supplementary study methods"
    )
    run.italic = True
    set_font(run, "Times New Roman", 10.5)

    toc_heading = document.add_paragraph()
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    toc_heading.paragraph_format.space_before = Pt(8)
    toc_heading.paragraph_format.space_after = Pt(4)
    toc_run = toc_heading.add_run("Table of Contents")
    toc_run.bold = True
    set_font(toc_run, "Times New Roman", 14)
    document.add_paragraph("[[TOC]]")
    document.add_page_break()

    # Chapter 1
    add_heading(document, "Chapter 1. Development of HORMONE-CYCLE", 1)
    add_heading(document, "1.1 Purpose, scope, and intended use", 2)
    add_paragraph(
        document,
        (
            "HORMONE-CYCLE is a custom, open-source simulator developed for this project. "
            "It generates reproducible daily menstrual-cycle, bleeding, ovulation, estradiol, "
            f"and progesterone diaries from age, diary length, a seed, and optional modifiers. {marker('C001')} "
            "HORMONE-CYCLE is the complete software title. The simulator supports "
            "method development, operating-characteristic studies, and controlled synthetic "
            "examples. Its outputs represent synthetic trajectories. Individual clinical forecasting "
            "falls outside the intended use."
        ),
    )
    add_paragraph(
        document,
        (
            "The design follows a hierarchical simulation workflow and reports inputs, data-generating "
            f"mechanisms, estimands, and reproducibility artifacts. {marker('C002')} Population evidence "
            "constrains age-specific distributions, each synthetic person receives stable latent traits, "
            "each cycle receives a stochastic realization, and a daily renderer produces the final diary."
        ),
    )

    add_heading(document, "1.2 Evidence review and role of each source", 2)
    add_paragraph(
        document,
        (
            "Three primary data sources supplied the baseline quantitative targets: the Apple Women’s "
            "Health Study for age-specific cycle length and variability, the Natural Cycles cohort for "
            "cycle and phase timing, and daily laboratory measurements for estradiol and progesterone "
            f"subphase values. {marker('C003')} Condition-specific sources informed modifier direction "
            "and bounded magnitudes. Table A1 states exactly how each source entered development and "
            "evaluation. Development used aggregate quantities reported in the articles; participant-level "
            "source data were unavailable and were not used."
        ),
    )
    new_section(document, "landscape")
    add_table_caption(
        document,
        "Appendix Table A1. ",
        (
            "Evidence provenance for HORMONE-CYCLE. “Reused” means that the same published result "
            "informed both parameter development and an internal calibration or directional check; "
            "these checks provide internal calibration evidence. Independent external validation requires "
            "held-out participant-level data."
        ),
    )
    add_table(
        document,
        [
            "Paper or evidence source",
            "Population or data",
            "Published result used",
            "Use during development",
            "Use during evaluation",
            "Independence status",
        ],
        evidence_rows(),
        [1.45, 1.85, 1.65, 1.55, 1.35, 1.15],
        font_size=7.7,
        left_columns=(0, 1, 2, 3, 4, 5),
    )
    new_section(document, "portrait")

    add_heading(document, "1.3 Workflow and data hierarchy", 2)
    add_figure(
        document,
        ASSET_DIR / "hormone_cycle_workflow.png",
        6.25,
        "Appendix Figure A1. ",
        (
            "HORMONE-CYCLE diary-generation workflow. By default, cycle 1 is generated in full and diary "
            "day 1 is selected uniformly from its realized cycle days. Output proceeds forward from that "
            "selected day through the remainder of cycle 1. If more days are required, the next complete "
            "cycle is generated and appended from cycle day 1 in its original order. This cycle-level "
            "workflow repeats until the requested diary length is filled; only the final generated cycle "
            "may be partly retained. The initial phase is selected once, and no end-to-start wraparound or "
            "circular shift is used. An explicit cycle-day-1 option is available. A domain-separated random "
            "stream selects the initial cycle day without consuming values from the patient-profile or "
            "cycle-generation random streams."
        ),
    )
    add_heading(document, "1.3.1 Inputs and outputs", 3)
    add_paragraph(
        document,
        (
            "The public simulation call accepts age in years, diary length in days, a random seed, "
            "and optional reproductive or medical modifiers. Input validation rejects nonfinite ages, "
            "nonpositive diary lengths, incompatible modifier combinations, and unsupported start modes. "
            "The default start mode selects each realized day of cycle 1 with equal probability and makes "
            "the selected day diary day 1. The optional cycle-day-1 mode fixes diary day 1 at menstrual "
            "cycle day 1. After that initial choice, records proceed chronologically without wrapping. "
            "Each daily row contains "
            "the calendar-day index, cycle index, day within cycle, realized cycle length, estradiol in "
            "picograms per milliliter, progesterone in nanograms per milliliter, ovulation, and bleeding. "
            "Cycle summaries report phase lengths, ovulation day, ovulatory status, bleeding days, "
            "reproductive stage, and active modifiers."
        ),
    )

    add_heading(document, "1.4 Age-specific patient-level parameters", 2)
    add_paragraph(
        document,
        (
            "Age selects one of eight calibration bands. Each band supplies a published mean cycle "
            "length and the probability that two adjacent cycles differ by at least 7 days. "
            f"Li et al. reported the age pattern from 165,668 cycles contributed by 12,608 participants. {marker('C004')} "
            "The implementation converts the observable adjacent-cycle probability into a latent "
            "within-person standard deviation under a Gaussian cycle-length model."
        ),
    )
    add_equation(
        document,
        "pΔ = Pr(|Lᵢ − Lᵢ₋₁| ≥ 7) = 2[1 − Φ(7/(√2 σL))]",
        1,
    )
    add_paragraph(
        document,
        (
            "In Equation 1, Lᵢ is cycle length, Φ is the standard normal cumulative distribution "
            "function, and σL is solved numerically by bisection. A person-specific mean is sampled "
            "around the age-band target. A person-specific within-person standard deviation is sampled "
            "around the value obtained from Equation 1."
        ),
    )
    add_paragraph(
        document,
        (
            "Stable person-level estradiol and progesterone amplitude multipliers follow lognormal "
            "distributions with coefficients of variation 0.18 and 0.22. The lognormal parameters are "
            "chosen so the multiplier has mean one."
        ),
    )
    add_equation(document, "σlog² = ln(1 + c²)", 2)
    add_equation(document, "μlog = −σlog²/2", 3)
    add_paragraph(
        document,
        "In Equations 2 and 3, c is the requested coefficient of variation.",
    )
    age_rows = [
        ["<15", "30.3", "31.2%", "0.62"],
        ["15–19", "30.3", "31.2%", "0.90"],
        ["20–24", "30.0", "20.4%", "0.97"],
        ["25–29", "29.1", "16.4%", "0.97"],
        ["30–34", "28.8", "14.7%", "0.97"],
        ["35–39", "28.8", "15.9%", "0.97"],
        ["40–44", "28.4", "20.2%", "0.95"],
        ["45–49", "28.2", "27.2%", "0.88"],
        ["≥50", "30.8", "55.4%", "0.70"],
    ]
    add_table_caption(
        document,
        "Appendix Table A2. ",
        (
            f"Age-specific baseline targets. Cycle-length and irregularity values are from Li et al. "
            f"(2023) {marker('C023')}. Ovulation probabilities are investigator-selected calibration "
            "constants implemented in HORMONE-CYCLE. The two rows below age 20 share one Li target band "
            "and differ only in the implemented ovulation probability."
        ),
    )
    add_table(
        document,
        [
            "Age band (years)",
            "Mean cycle length (days)",
            "Adjacent-cycle difference ≥7 days",
            "Baseline ovulation probability",
        ],
        age_rows,
        [1.2, 1.65, 2.05, 1.6],
        font_size=8.8,
    )

    add_heading(document, "1.5 Cycle length, phase timing, and bleeding", 2)
    add_paragraph(
        document,
        (
            "For each cycle, ovulation is drawn from a Bernoulli distribution using the resolved "
            "person-level probability. Cycle length is drawn from a truncated Gaussian distribution "
            "centered on the person-specific mean, with a minimum of 18 days and a maximum of 120 days. "
            "Ovulatory cycles use a relatively tight luteal distribution centered on 12.4 days; the "
            "follicular interval absorbs the remaining cycle-length variation. Natural Cycles data "
            "reported a 29.3-day mean cycle, 16.9-day mean follicular phase, and 12.4-day mean luteal "
            f"phase across 612,613 ovulatory cycles from 124,648 users. {marker('C005')}"
        ),
    )
    add_paragraph(
        document,
        (
            "Bleeding duration is sampled around 4.7 days with a standard deviation of 1.0 day and "
            "is bounded to 0–12 days. The 4.7-day value is an investigator-selected calibration setting. "
            "Bull et al. reported a 4.0 ± 1.5-day mean bleeding duration in Table 1; the calibration assessment "
            "therefore compares the realized simulation with that published distribution. The International "
            "Federation of Gynecology and Obstetrics terminology provides the clinical interpretation framework "
            f"for bleeding patterns. {marker('C063')} The cycle-length bounds and dispersion values are "
            "transparent calibration choices."
        ),
    )

    add_heading(document, "1.6 Daily hormone trajectories", 2)
    add_paragraph(
        document,
        (
            "Ovulatory trajectories use seven control points derived from serum estradiol and progesterone "
            "medians measured daily in 20 healthy volunteers and synchronized to the luteinizing-hormone "
            f"peak. {marker('C006')} The simulator places those control points within each realized "
            "follicular and luteal interval. For a normalized position t between two adjacent points, "
            "the renderer uses shape-preserving piecewise cubic Hermite interpolation (PCHIP) in Equations 4 and 5."
        ),
    )
    add_equation(document, "t = (x − xₖ)/hₖ; h₀₀ = 2t³ − 3t² + 1; h₁₀ = t³ − 2t² + t; h₀₁ = −2t³ + 3t²; h₁₁ = t³ − t²", 4)
    add_equation(document, "y(x) = h₀₀yₖ + h₁₀hₖmₖ + h₀₁yₖ₊₁ + h₁₁hₖmₖ₊₁", 5)
    hormone_rows = [
        ["Early follicular", "42.9", "0.44"],
        ["Mid-follicular", "88.1", "0.56"],
        ["Preovulatory", "234.0", "1.31"],
        ["Ovulation", "141.0", "1.43"],
        ["Early luteal", "132.0", "3.95"],
        ["Midluteal", "117.0", "11.02"],
        ["Late luteal", "111.0", "6.75"],
    ]
    add_table_caption(
        document,
        "Appendix Table A3. ",
        (
            f"Implemented ovulatory hormone control points derived from the daily series in Table 1B and "
            f"Figure 1 of Stricker et al. (2006) {marker('C024')}. Values were converted to the displayed "
            "units and assigned to seven investigator-defined interpolation positions. Estradiol is in "
            "picograms per milliliter and progesterone is in nanograms per milliliter."
        ),
    )
    add_table(
        document,
        ["Menstrual subphase", "Estradiol target", "Progesterone target"],
        hormone_rows,
        [2.2, 1.5, 1.6],
        font_size=9.0,
    )
    add_paragraph(
        document,
        (
            "Anovulatory cycles use four lower-amplitude control points. Person-level hormone scales are "
            "multiplied by cycle-level lognormal scales with coefficients of variation 0.08 for estradiol "
            "and 0.10 for progesterone. Day-to-day continuity is produced by a stationary first-order autoregressive "
            "noise process with coefficient 0.92; each realized path is linearly bridged to zero at both cycle endpoints."
        ),
    )
    add_equation(document, "ε̃d = 0.92 ε̃d₋₁ + ηd; εd = ε̃d − [ε̃1 + (d−1)(ε̃L−ε̃1)/(L−1)]", 6)
    add_paragraph(
        document,
        (
            "Equation 6 defines the serial noise state and endpoint bridge. The innovation ηd is drawn independently each day "
            "from a zero-mean Gaussian distribution scaled to preserve the requested stationary variance. The coefficient 0.92 is an investigator-selected "
            "continuity parameter, and endpoint bridging prevents a stochastic cross-cycle reset. Final estradiol values are floored at 5 picograms per milliliter and "
            "progesterone values at 0.05 nanograms per milliliter."
        ),
    )

    add_heading(document, "1.7 Reproductive-stage and medical modifiers", 2)
    add_heading(document, "1.7.1 Polycystic ovary syndrome", 3)
    add_paragraph(
        document,
        (
            "The polycystic-ovary profile lengthens cycles, increases cycle-to-cycle variation, reduces "
            "ovulation probability, and attenuates luteal progesterone. Mortimer et al. (2026) and Doi "
            f"et al. supplied the initial direction-setting evidence. {marker('C044')} Later age-stratified "
            "cycle data and intensive ultrasound and hormone sampling corroborated the implemented "
            f"profile. {marker('C045')}"
        ),
    )
    add_heading(document, "1.7.2 Early postmenarche", 3)
    add_paragraph(
        document,
        (
            "The early-postmenarche profile lengthens cycles, increases variability, reduces ovulation, "
            "and shortens or attenuates luteal progesterone patterns. Venturoli et al. supplied the initial "
            f"direction-setting evidence. {marker('C046')} An adolescent multicenter study and a two-year "
            f"daily hormone study later corroborated these directions. {marker('C047')}"
        ),
    )
    add_heading(document, "1.7.3 Perimenopause", 3)
    add_paragraph(
        document,
        (
            "The perimenopause profile increases cycle variability, lowers ovulation probability, "
            "attenuates progesterone, and increases hormone variability in accordance with menopause-"
            f"transition evidence. {marker('C009')}"
        ),
    )
    add_heading(document, "1.7.4 Combined oral contraceptives", 3)
    add_paragraph(
        document,
        (
            "Cyclic and continuous combined oral-contraceptive profiles use fixed 28-day regimen cycles "
            "with suppressed endogenous-equivalent hormone curves. The cyclic profile generates scheduled "
            "withdrawal bleeding. The continuous profile generates amenorrhea or breakthrough bleeding. "
            f"These behaviors follow the evidence summarized by Edelman et al. {marker('C010')}"
        ),
    )
    add_heading(document, "1.7.5 Intrauterine devices", 3)
    add_paragraph(
        document,
        (
            "The levonorgestrel-releasing intrauterine-device profile preserves ovulation in most cycles "
            "and reduces bleeding, with amenorrhea in a minority of cycles. Long-term hormone and ultrasound "
            f"studies support these directions. {marker('C011')} The copper intrauterine-device profile "
            "preserves ovarian cycling and increases bleeding duration. Hubacher et al. supplied contextual "
            f"evidence in the initial source registry. {marker('C051')} Hormone and measured-blood-loss "
            f"studies later corroborated that profile. {marker('C052')}"
        ),
    )
    add_heading(document, "1.7.6 Primary dysmenorrhea", 3)
    add_paragraph(
        document,
        (
            "The primary-dysmenorrhea profile preserves ovulation and applies a small bleeding-duration "
            "shift. This implementation reflects the predominantly prostaglandin-mediated pain phenotype "
            f"described in the clinical literature. {marker('C013')}"
        ),
    )
    new_section(document, "landscape")
    add_table_caption(
        document,
        "Appendix Table A4. ",
        (
            "Core HORMONE-CYCLE parameter catalogue. The implemented rules were transcribed from "
            "src/hormone_cycler/hormone_constants.py, model.py, and types.py. The final column distinguishes "
            "direct or derived published quantities, calibrated constants, investigator choices, and "
            "software rules. Published timing and hormone anchors refer to Li et al. (2023), Bull et al. "
            f"(2019), and Stricker et al. (2006) {marker('C003')}. These exact values describe version 0.1.0."
        ),
    )
    add_table(
        document,
        [
            "Parameter",
            "Hierarchy",
            "Implemented rule or value",
            "Bounds, rounding, or transformation",
            "Evidence classification",
        ],
        core_parameter_rows(),
        [1.35, 0.85, 3.15, 2.15, 1.5],
        font_size=7.2,
        left_columns=(0, 1, 2, 3, 4),
    )

    add_table_caption(
        document,
        "Appendix Table A5. ",
        (
            "Implemented modifier values and their evidence sources. The numerical multipliers are "
            "calibration constants unless the row states a direct published quantity. The cited papers "
            "support the direction or clinical context of each profile."
        ),
    )
    add_table(
        document,
        [
            "Modifier",
            "Cycle-length effect",
            "Cycle variability",
            "Ovulation effect",
            "Bleeding and hormone effects",
            "Published evidence",
        ],
        modifier_rows(),
        [1.45, 1.5, 1.15, 1.1, 2.55, 1.25],
        font_size=7.7,
        left_columns=(0, 1, 2, 3, 4, 5),
    )
    new_section(document, "portrait")

    add_heading(document, "1.8 Reproducibility and implementation limits", 2)
    add_paragraph(
        document,
        (
            "Fixed inputs and a fixed seed produce identical diaries. Reproducibility records should "
            "preserve the repository commit, configuration, seed, requested diary length, age and modifier "
            "inputs, and the validation-report hash. The current implementation omits pregnancy, postpartum "
            "physiology, medication changes, secular trends, missing diary entries, bleeding intensity, "
            "laboratory-assay error, and mechanistic endocrine feedback. Modifier combinations are coarse "
            "profile transformations that support synthetic stress testing. Causal treatment effects remain "
            "outside the scope of the simulator."
        ),
    )

    # Chapter 2
    document.add_page_break()
    add_heading(
        document,
        "Chapter 2. Internal calibration and validation status of HORMONE-CYCLE",
        1,
    )
    add_heading(document, "2.1 Evaluation terminology and design", 2)
    add_paragraph(
        document,
        (
            "This chapter reports internal target reproduction and directional software stress tests. "
            "The same aggregate sources informed parameter development and evaluation. Independent external "
            "validation requires held-out participant-level data and remains future work."
        ),
    )
    add_paragraph(
        document,
        (
            "The quality-control design follows established simulation-study reporting principles. "
            f"{marker('C014')} The baseline command was equivalent to “hormone_cycler validate "
            "--patients 10000 --days 365 --seed 7.” The acceptance rules were encoded in "
            "src/hormone_cycler/hormone_constants.py before this reported execution. The run compared "
            "19 distributional summaries from 10,000 synthetic participants with published aggregate "
            "targets. Fourteen hormone-anchor and four kinetic smoke checks used 16 age-balanced retained diaries. Eight modifier scenarios "
            "tested code-defined directional or bounded behavior."
        ),
    )
    add_paragraph(
        document,
        (
            "The primary sources used for parameter development were reused as evaluation targets. The "
            "results therefore establish internal distributional calibration and directional face validity. "
            "Independent external validation against unseen participant-level menstrual and hormone diaries "
            "remains outstanding. Table A1 records this source reuse explicitly."
        ),
    )

    add_heading(document, "2.2 Baseline calibration targets and tolerances", 2)
    add_paragraph(
        document,
        (
            "The 19 distributional checks comprise eight age-specific mean cycle lengths, eight age-specific "
            "irregularity probabilities, mean follicular and luteal intervals, and mean bleeding duration. "
            "Age-specific acceptance bands combine a minimum practical margin with the published "
            "confidence-interval half-width and a reconstruction allowance. The aggregate synthetic age "
            "distribution differs from the Bull et al. cohort, which was restricted to ages 18–45 and selected "
            "ovulatory cycles. The phase-duration comparison is therefore a structurally unmatched internal "
            "calibration check."
        ),
    )
    add_paragraph(
        document,
        (
            "Fourteen hormone-anchor checks compare seven estradiol and seven progesterone subphase summaries with "
            "Stricker-derived anchors. Four additional kinetic checks evaluate estradiol peak width, premenstrual progesterone withdrawal, the terminal-to-peak progesterone ratio, and cross-cycle progesterone continuity. "
            "The validation implementation retains two diaries from each of eight age bands (16 total). These checks confirm software output "
            "and approximate anchor reproduction in a small age-balanced sample. They do not establish "
            "population-level hormone validation."
        ),
    )
    new_section(document, "landscape")
    add_table_caption(
        document,
        "Appendix Table A6. ",
        (
            f"Complete internal calibration and hormone smoke-check results. Calibration targets are from "
            f"Li et al. (2023), Bull et al. (2019), and Stricker et al. (2006) {marker('C015')}. The 19 "
            "distributional checks use 10,000 synthetic participants, 365 days per participant, and seed 7. "
            "The 18 hormone anchor and kinetic checks use 16 retained diaries balanced across eight age bands. Every row "
            "met its code-defined acceptance band. The results do not constitute independent external "
            "validation. Duration metrics are in days; estradiol is in picograms per milliliter; progesterone "
            "is in nanograms per milliliter."
        ),
    )
    add_table(
        document,
        ["Metric", "Simulated", "Calibration target", "Acceptance range", "Result", "Source and location", "Simulation sample"],
        validation_rows(validation),
        [1.85, 0.7, 0.8, 1.25, 0.55, 1.45, 2.4],
        font_size=7.3,
        left_columns=(0, 5, 6),
    )

    add_heading(document, "2.3 Modifier-scenario checks", 2)
    add_paragraph(
        document,
        (
            "Each modifier was simulated separately against the same baseline reference. The checks assessed "
            "cycle length, irregularity, ovulation, bleeding, or amenorrhea according to the feature addressed "
            "by the source literature. Each scenario used 1,200 synthetic participants distributed equally "
            "across all age bands. The early-postmenarche and perimenopause scenarios are therefore not "
            "age-matched to their clinical populations. These results are software stress tests of the encoded "
            "modifier behavior."
        ),
    )
    add_table_caption(
        document,
        "Appendix Table A7. ",
        (
            "Modifier-scenario summary. Each scenario uses 1,200 synthetic participants distributed equally "
            "across all age bands. “Code-defined checks” counts the acceptance criteria listed individually "
            "in Table A8. Polycystic ovary syndrome is compared with Mortimer et al. (2026) "
            f"and Jarrett et al. (2020) {marker('C025')}. Early postmenarche is compared with the World Health "
            f"Organization adolescent study (1986) and Zhang et al. (2008) {marker('C026')}. Perimenopause "
            f"is compared with Santoro and Randolph (2011) {marker('C027')}. Combined oral contraceptives "
            f"are compared with Edelman et al. (2014) {marker('C028')}. The levonorgestrel-releasing "
            f"intrauterine device is compared with Xiao et al. (1995) and Barbosa et al. (1995) {marker('C029')}. "
            f"The copper intrauterine device is compared with Faundes et al. (1980) and Malmqvist et al. (1974) "
            f"{marker('C030')}. Primary dysmenorrhea is compared with Dawood (2006) {marker('C031')}."
        ),
    )
    add_table(
        document,
        [
            "Scenario",
            "Mean cycle (days)",
            "Ovulatory cycles",
            "Mean bleeding (days)",
            "Irregularity",
            "Amenorrhea",
            "Code-defined checks",
        ],
        subgroup_rows(validation),
        [1.9, 1.1, 1.1, 1.25, 1.0, 1.0, 1.15],
        font_size=7.8,
        left_columns=(0,),
    )

    add_table_caption(
        document,
        "Appendix Table A8. ",
        (
            "Auditable modifier-scenario criteria. Each comparison uses 1,200 synthetic participants "
            "distributed equally across the implemented age bands. Acceptance bounds were defined in the "
            "validation code and are investigator-selected unless the criterion names a direct published "
            "quantity. The supporting papers are named in the source column."
        ),
    )
    add_table(
        document,
        [
            "Scenario",
            "Code-defined criterion",
            "Simulated value",
            "Acceptance range",
            "Supporting paper",
            "Result",
        ],
        subgroup_check_rows(validation),
        [1.75, 2.6, 1.05, 1.35, 2.2, 0.55],
        font_size=7.2,
        left_columns=(0, 1, 4),
    )

    add_heading(document, "2.4 Graphical calibration summary", 2)
    add_figure(
        document,
        ASSET_DIR / "hormone_cycle_validation.png",
        8.0,
        "Appendix Figure A2. ",
        (
            "Internal observed-versus-target calibration. Panels A and B compare age-stratified cycle length and "
            "irregularity with Li et al. (2023); panel C compares follicular, luteal, and bleeding durations "
            "with Bull et al. (2019), including the published 4.0-day mean bleeding duration from Table 1; "
            "panel D compares estradiol and progesterone subphase summaries with Stricker et al. (2006) "
            f"{marker('C017')}. Panels A–C use the 10,000-participant run. Panel D uses 16 retained diaries, "
            "all from the younger-than-20 age band. Early follicular, mid-follicular, preovulatory, ovulatory, "
            "early luteal, midluteal, and late luteal phases are shown. These are internal calibration and "
            "software smoke checks using development sources."
        ),
    )
    new_section(document, "portrait")
    add_heading(document, "2.5 Interpretation and remaining validation needs", 2)
    add_paragraph(
        document,
        (
            "The current checks show that the implementation reproduces its selected marginal targets and "
            "expresses the intended modifier directions. Joint patient-level fidelity, transportability to "
            "other populations, assay-specific error, and individual-cycle prediction remain untested. "
            "A future external validation should use held-out daily bleeding, ovulation, and hormone "
            "records; align phase definitions prospectively; report calibration by age and reproductive stage; "
            "and examine multivariate dependence, serial correlation, and between-person heterogeneity."
        ),
    )

    # Chapter 3
    document.add_page_break()
    add_heading(
        document,
        "Chapter 3. Additional analyses and methods augmenting the main results",
        1,
    )
    add_paragraph(
        document,
        (
            "This chapter documents how HORMONE-CYCLE and the seizure-diary simulator were combined for "
            "the present false-positive study. It then reports the phase algorithms, ratio and regression "
            "methods, sensitivity analyses, and supplementary results."
        ),
    )
    add_heading(document, "3.1 Reproducibility, provenance, and simulator inputs", 2)
    add_paragraph(
        document,
        (
            "The completed primary simulation included 100,000 synthetic participants, 50,000 per cohort, "
            "followed for 36 months with master seed 20260505. HORMONE-CYCLE generated menstrual and hormone "
            f"diaries. {marker('C001')} The Cyclic Heterogeneous Overdispersed Clustered Open-source L-relationship Adjustable "
            "Temporally limited E-diary Simulator (CHOCOLATES) generated seizure diaries. The generators were "
            "run independently. HORMONE-CYCLE selected the first output day uniformly from the realized "
            "days of the first generated menstrual cycle and then continued forward without wrapping. The "
            f"two diaries were merged by calendar-day index. {marker('C033')} Direct alignment preserved "
            "the original day ordering throughout the merge."
        ),
    )
    add_paragraph(
        document,
        (
            f"The reporting framework follows simulation-study guidance. {marker('C014')} CHOCOLATES and "
            "HORMONE-CYCLE were designed to reproduce selected statistical features within and across "
            "synthetic patient populations through data-informed hierarchical sampling. Chapter 1 details "
            "the HORMONE-CYCLE evidence sources and transformations."
        ),
    )
    add_paragraph(
        document,
        (
            "The completed manifest stores Secure Hash Algorithm 256-bit fingerprints for the analysis "
            "code, configuration, and generated files. These fingerprints should accompany archived "
            "results. The study wrapper and HORMONE-CYCLE package declare version 0.1.0 and Python 3.11 or "
            "later. The full-run command was equivalent to “run_paper1_null_ce.py --config "
            "config_random_start_full.yaml --full,” and the archived machine-readable outputs are in "
            "outputs/random_start_full. The manifest records the complete configuration, assumptions, "
            "generated-file hashes, and dependency-facing analysis fingerprints."
        ),
    )
    new_section(document, "landscape")
    t0 = deepcopy(source_tables[0])
    t0[5][6] = "Available through generator diagnostics"
    t0[6][3] = "Healthy cohort: 18–45 years with ovulation required; heterogeneous cohort: 13–54.9 years"
    t0[11][5] = "Operational rule based on Herzog et al. (1997)"
    t0[11][6] = "Audited fraction designated as having an inadequate luteal phase"
    t0[15][6] = "Excluded from participant summary"
    t0[16][3] = "Available for future scenarios; excluded from the completed run"
    t0[16][6] = "Outside the completed simulation"
    t0[18][3] = "Complete synthetic diaries"
    t0[18][6] = "All requested days retained"
    t0[19][3] = (
        "Separately seeded generators; menstrual diary starts at a uniformly selected first-cycle day; "
        "diaries merged by matching calendar-day index without wrapping"
    )
    t0[21][3] = "Strict Herzog and luteal-anchored ovulatory definitions"
    t0[22][3] = "50,000 healthy ovulatory and 50,000 heterogeneous participants"
    add_table_caption(
        document,
        "Appendix Table S1. ",
        (
            "Simulator inputs, realized features, and evidence roles for the completed study. Literature "
            "sources are named in the source/rationale column; investigator-configured values are study "
            f"assumptions. The named HORMONE-CYCLE, CHOCOLATES, seizure-rhythm, menstrual-cycle, and Herzog "
            f"sources are cited here {marker('C062')}. “Simulator-native” means drawn directly by CHOCOLATES. "
            "“Multidien” denotes periods longer than 24 hours. “Audit spectrum” is a spectral estimate from "
            "retained daily audit data. “Interseizure constraints” are generator limits on simulated "
            "intervals between seizures. Hormone concentrations are written in nanograms per milliliter "
            "or picograms per milliliter."
        ),
    )
    add_table(
        document,
        t0[0],
        t0[1:],
        [1.0, 1.2, 0.9, 2.15, 0.8, 1.7, 1.25],
        font_size=7.5,
        left_columns=(0, 1, 3, 5, 6),
    )

    add_heading(document, "3.2 Assumption-by-assumption review", 2)
    add_paragraph(
        document,
        (
            "Within this section, menstrual (M), follicular (F), ovulatory (O), and luteal (L) identify "
            "cycle phases. Type C1 is the perimenstrual pattern, type C2 is the periovulatory pattern, and "
            "type C3 is the inadequate-luteal-phase pattern. Inadequate luteal phase (ILP) is the operational "
            "designation defined in Section 3.4."
        ),
    )
    assumption = deepcopy(source_tables[6])
    assumption[1][1] = "Defines the false-positive estimand through independent seizure and hormone random streams."
    assumption[1][2] = "True catamenial effects and false negatives fall outside this null simulation."
    assumption[2][2] = "Separate random streams establish independence, and direct alignment avoids a wraparound boundary."
    assumption[3][2] = "Medication changes and secular trends fall outside the data-generating model."
    assumption[4][2] = "This reference cohort represents ovulatory menstrual cycling under the configured age range."
    assumption[5][2] = "Configured modifier rates serve as stress-test inputs; population prevalence estimation falls outside scope."
    assumption[8][1] = "Uses the published distributional inflection points as the implemented thresholds."
    assumption[9][1] = "Reports a positive numerator over a zero comparator as +infinity and reports 0/0 as undefined."
    assumption[10][2] = "The cutoff captures one operational definition of luteal adequacy."
    assumption[11][2] = "Type C3 sensitivity depends directly on the simulated prevalence of anovulation."
    assumption[12][2] = "The four-month or six-cycle and four-seizure-day cutoffs are pragmatic sensitivity settings."
    assumption[13][2] = "This analysis is a statistical calibration check for the simulator-based study."
    assumption[14][1] = "Provides a separately specified follicular-versus-combined-phase contrast in ILP cycles."
    assumption[15][0] = "Complete-diary observation"
    assumption[16][0] = "Stationary treatment and background risk"
    add_table_caption(
        document,
        "Appendix Table S2. ",
        (
            "Assumptions, rationale, and implications for interpretation. The Herzog ratio definitions, "
            f"thresholds, and inadequate-luteal-phase context derive from Herzog and colleagues {marker('C022')}."
        ),
    )
    add_table(
        document,
        assumption[0],
        assumption[1:],
        [2.0, 3.45, 3.55],
        font_size=7.6,
        left_columns=(0, 1, 2),
    )
    new_section(document, "portrait")

    add_heading(document, "3.3 Menstrual-phase labeling and complete-cycle rules", 2)
    add_paragraph(
        document,
        (
            "Phase labels were assigned to each complete menstrual cycle before any analysis-window subset "
            "was taken. Forward day counts from menstrual onset, with day 1 as the first bleeding day. "
            "Backward day equals the forward day minus the cycle length plus one, so the final cycle day is "
            "−1. Label priority is menstrual, follicular, ovulatory, and then luteal. Days outside the listed "
            "ranges remain unlabeled in the strict definition."
        ),
    )
    phase_rows = [
        ["Menstrual (M)", "Forward days 1–3 or backward days −3 to −1", "Same"],
        ["Follicular (F)", "Forward days 4–9", "All days not assigned to M, O, or L"],
        ["Ovulatory (O)", "Forward day ≥10 and backward day ≤−13", "Backward days −16 to −13"],
        ["Luteal (L)", "Backward days −12 to −4", "Backward days −12 to −4"],
        ["Unlabeled", "All remaining days", "None"],
    ]
    add_table_caption(
        document,
        "Appendix Table S3. ",
        (
            "Implemented strict Herzog and luteal-anchored sensitivity phase boundaries. Forward day is "
            "counted from menstrual onset and backward day from the next onset. Priority rules prevent "
            "overlapping labels. The luteal-anchored definition fixes the ovulatory interval immediately "
            "before the luteal phase and assigns the remaining nonmenstrual days to the follicular phase."
        ),
    )
    add_table(
        document,
        ["Phase", "Strict Herzog definition", "Luteal-anchored sensitivity"],
        phase_rows,
        [1.25, 2.75, 2.5],
        font_size=8.4,
        left_columns=(0, 1, 2),
    )
    add_paragraph(
        document,
        (
            "A complete cycle contains every daily row from cycle day 1 through its recorded cycle length. "
            "The exact three-cycle Herzog analysis additionally requires three complete cycles, each 23–35 "
            "days long. For combined phases, average daily seizure frequency is calculated as the total "
            "seizure count across all included phases divided by the total number of labeled days across "
            "those phases."
        ),
    )

    add_heading(document, "3.4 Type C3 and inadequate-luteal-phase algorithm", 2)
    add_paragraph(
        document,
        (
            "Type C3 is the inadequate-luteal-phase catamenial pattern. The present adapter is a custom "
            "operationalization. Herzog et al. used a single day-22 progesterone measurement and a "
            f"5-nanogram-per-milliliter threshold. {marker('C019')} This study borrowed that threshold and "
            "defined the timing and aggregation algorithm separately."
        ),
    )
    add_paragraph(
        document,
        (
            "The adapter designates a nonovulatory cycle or a cycle without a valid ovulation day as "
            "inadequate. For an ovulatory cycle, it takes the maximum progesterone concentration from days "
            "5–9 after ovulation. If the interval reaches the diary boundary, it uses the available "
            "postovulation days. A maximum below 5 nanograms per milliliter is designated inadequate. "
            "A maximum of exactly 5 is therefore adequate in the implemented rule. "
            "Complete HORMONE-CYCLE daily outputs contain progesterone on every simulated day. The adapter "
            "also contains a defensive missing-field branch that designates the cycle inadequate; that branch "
            "was never exercised in this completed simulation and should be treated as an implementation "
            "limitation."
        ),
    )
    add_paragraph(
        document,
        (
            "Type C3 applicability is evaluated within each analysis window. A pooled window is applicable "
            "when it contains at least one designated cycle. The completed study evaluated type C3 only in "
            "the heterogeneous cohort as a study-design applicability rule. The healthy cohort was omitted "
            "from type C3 reporting; enforced ovulation does not establish that low luteal progesterone is "
            "biologically impossible."
        ),
    )

    add_heading(document, "3.5 Herzog seizure-frequency ratios and thresholds", 2)
    add_paragraph(
        document,
        (
            "Average daily seizure frequency (ADSF) is the number of seizures in a phase divided by the number of "
            "labeled days in that phase. Menstrual (M), follicular (F), ovulatory (O), and luteal (L) denote "
            "the locally defined phases. The three ratios are shown in Equations 7–9. The exact thresholds "
            f"1.69, 1.83, and 1.62 correspond to published distributional inflection points. {marker('C020')}"
        ),
    )
    add_equation(document, "C1 = ADSF(M) / ADSF(F + L)", 7)
    add_equation(document, "C2 = ADSF(O) / ADSF(F + L)", 8)
    add_equation(document, "C3 = ADSF(O + L + M) / ADSF(F)", 9)
    add_paragraph(
        document,
        (
            "A small comparator count can produce a large or undefined ratio. Positive infinity is defined "
            "and exceeds every finite "
            "threshold. Ratios of zero divided by zero and ratios with missing phase days are undefined."
        ),
    )
    add_paragraph(
        document,
        (
            f"The exact three-cycle analysis requires exactly three complete 23–35-day cycles and classifies "
            f"a participant when at least two cycles are positive. {marker('C021')}"
        ),
    )

    ratio_tables = [
        ("Panel A — Type C1", source_tables[7], [1.45, 1.1, 1.1, 1.25, 1.25]),
        ("Panel B — Type C2", source_tables[8], [1.45, 1.1, 1.1, 1.25, 1.25]),
        ("Panel C — Type C3", source_tables[9], [1.55, 1.35, 1.35]),
    ]
    add_table_caption(
        document,
        "Appendix Table S4. ",
        (
            "Cumulative distributions of simulated Herzog seizure-frequency ratios in strict "
            "three-complete-cycle windows. The layout follows the cumulative presentation used by Herzog "
            f"(2015), and the ratio definitions originate from Herzog et al. (1997) {marker('C022')}. "
            "Only the present simulations are tabulated. In panels A–C, the threshold-zero denominator "
            "includes finite and positive-infinite pooled ratios and excludes undefined ratios. It is smaller "
            "than 50,000 because zero-over-zero ratios, missing required phase days, and eligibility failures "
            "are undefined or unavailable. Positive infinity is included at every applicable threshold."
        ),
    )
    for panel, table_data, widths in ratio_tables:
        add_heading(document, panel, 3)
        add_table(
            document,
            table_data[0],
            table_data[1:],
            widths,
            font_size=8.2,
            left_columns=(0,),
        )

    add_heading(document, "Panel D — Sparse-comparator audit", 3)
    sparse_rows = [
        ["Type C1", "Healthy ovulatory", "41,624", "1,983", "6,393", "11,748", "1,983 (16.9%)"],
        ["Type C1", "Heterogeneous", "41,584", "2,069", "6,347", "11,967", "2,069 (17.3%)"],
        ["Type C2", "Healthy ovulatory", "41,612", "2,393", "5,995", "11,282", "2,393 (21.2%)"],
        ["Type C2", "Heterogeneous", "41,566", "2,783", "5,651", "11,554", "2,783 (24.1%)"],
        ["Type C3", "Heterogeneous", "9,662", "7,094", "33,244", "8,849", "7,094 (80.2%)"],
    ]
    add_table(
        document,
        [
            "Ratio",
            "Cohort",
            "Finite",
            "Positive infinity",
            "Undefined",
            "At or above threshold",
            "Positive-infinite share",
        ],
        sparse_rows,
        [0.7, 1.25, 0.85, 1.15, 0.9, 1.25, 1.3],
        font_size=7.8,
        left_columns=(0, 1),
    )
    add_paragraph(
        document,
        (
            "Panel D audits pooled ratio status in the same strict three-cycle data used for panels A–C. "
            "The final exact Herzog classification separately requires at least two of three individual "
            "cycles to be positive. “Positive-infinite share” is the number of positive-infinite ratios "
            "divided by all pooled ratios at or above the applicable threshold."
        ),
    )

    add_heading(document, "3.6 Type C3 window-duration sensitivity", 2)
    add_paragraph(
        document,
        (
            "Type C3 is the inadequate-luteal-phase pattern. This sensitivity analysis varies the amount "
            "of observation available for that pattern."
        ),
    )
    c3_window = deepcopy(source_tables[10])
    c3_window[0][4] = "Positive among classifiable, % (95% Wilson confidence interval)"
    c3_window = [
        row for row in c3_window if row[0] != "36-month full diary"
    ]
    add_table_caption(
        document,
        "Appendix Table S5. ",
        (
            "Windowed type C3 false-positive classifications in the heterogeneous cohort. "
            "All 50,000 heterogeneous-cohort participants were attempted. “Applicable” indicates at least "
            "one designated inadequate-luteal-phase cycle in the window. “Classifiable” additionally requires "
            "a defined ratio. “All-attempted rate” is positive divided by 50,000. The healthy ovulatory cohort "
            "was omitted under the study-design applicability rule described in Section 3.4."
        ),
    )
    add_table(
        document,
        c3_window[0],
        c3_window[1:],
        [1.2, 1.0, 1.0, 0.9, 1.55, 1.15],
        font_size=8.2,
        left_columns=(0,),
    )

    new_section(document, "landscape")
    add_heading(document, "3.7 Investigator-selected minimum-data sensitivity", 2)
    minimum = minimum_data_rows()
    add_paragraph(
        document,
        (
            "The minimum-data thresholds are investigator-selected sensitivity settings. They represent "
            "the minimum amount of seizure information an investigator is willing to accept before "
            "classification. The analysis reports every setting; investigators select the sensitivity "
            "appropriate to their scientific question."
        ),
    )
    add_table_caption(
        document,
        "Appendix Table S6. ",
        (
            "Classification across observation windows and minimum seizure-day requirements. Each cohort "
            "contains 50,000 attempted participants. “Positive among classifiable” is positive divided by "
            "classifiable with a 95% Wilson score confidence interval. “All-attempted rate” is positive divided "
            f"by 50,000 {marker('C034')}."
        ),
    )
    add_table(
        document,
        [
            "Cohort",
            "Window",
            "Minimum seizure days",
            "Classifiable",
            "Positive",
            "Positive among classifiable, % (95% Wilson confidence interval)",
            "All-attempted rate",
        ],
        minimum,
        [1.2, 0.95, 1.15, 1.0, 0.85, 2.45, 1.1],
        font_size=7.2,
        left_columns=(0, 1),
    )

    add_heading(document, "3.8 Model-concordant negative-binomial calibration checks", 2)
    add_paragraph(
        document,
        (
            "The primary model-concordant calibration check uses a log-link negative-binomial generalized "
            "linear model for daily seizure counts. Participant-full-diary method-of-moments dispersion is "
            "used, with cycle fixed effects when the required complete cycles are available. Type C1 and "
            f"type C2 one-sided Wald tests are adjusted by the Holm method within participant. {marker('C035')} "
            "Type C1 is the perimenstrual contrast, type C2 is the periovulatory contrast, and type C3 is "
            f"the inadequate-luteal-phase contrast. The Cyclic Heterogeneous Overdispersed Clustered Open-source "
            "L-relationship Adjustable Temporally limited E-diary Simulator (CHOCOLATES) uses an overdispersed "
            f"count process. {marker('C033')} Concordance is limited to the count family and log link. The "
            "regression does not reproduce CHOCOLATES clustering, multidien rhythms, or participant heterogeneity. "
            "The indicator "
            "IM equals 1 on menstrual days, the indicator IO equals 1 on ovulatory days, and both equal 0 "
            "otherwise."
        ),
    )
    add_equation(
        document,
        "log(μicd) = β0 + βM IM,icd + βO IO,icd + γc",
        10,
    )
    add_paragraph(
        document,
        (
            "In Equation 10, μicd is the expected seizure count for participant i, cycle c, and day d; "
            "IM and IO are the menstrual-day and ovulatory-day indicators defined above; and γc is a cycle "
            "fixed effect."
        ),
    )
    add_paragraph(
        document,
        (
            "The exploratory type C3 model restricts the data to complete simulator-designated inadequate-"
            "luteal-phase cycles and replaces the two phase indicators with one indicator for combined "
            "ovulatory, luteal, and menstrual days. Positivity requires a one-sided P value below .05 and "
            "a rate ratio of at least 1.62. The minimum data are four complete designated cycles and four "
            "seizure days. A robust Poisson model with heteroskedasticity-consistent covariance type 0 is "
            "attempted only when the negative-binomial fit raises an exception. Both models use a maximum "
            "of 100 iterations. Equation 11 defines the per-participant dispersion, where s² is the daily-count "
            "variance and ȳ is the daily-count mean. The dispersion is 1.0 when the variance does not exceed "
            "the mean. Each daily row has one day of exposure, so no offset is required."
        ),
    )
    add_equation(
        document,
        "αi = max{10⁻⁶, min[50, (sᵢ² − ȳᵢ)/ȳᵢ²]}",
        11,
    )
    add_paragraph(
        document,
        (
            "Equation 12 gives the exploratory type C3 model. The indicator IOLM equals 1 on ovulatory, "
            "luteal, or menstrual days and 0 on follicular days within complete cycles designated as having "
            "an inadequate luteal phase."
        ),
    )
    add_equation(
        document,
        "log(μicd) = β0 + βOLM IOLM,icd + γc",
        12,
    )
    add_paragraph(
        document,
        (
            "The retained daily audit sample was selected independently within each cohort without replacement "
            "at a 1% fraction. NumPy’s default random-number generator used a deterministic 32-bit seed derived "
            "from master seed 20260505, the cohort name, and “audit_sample.” The heterogeneous-cohort seed was "
            "529110050. The type C3 exploratory analysis attempted all 500 retained heterogeneous participants: "
            "431 had ratio-level type C3 applicability, 204 met regression data requirements, 271 had fewer "
            "than four complete inadequate-luteal-phase cycles, and 25 had fewer than four seizure days. "
            "All 204 classifiable participants used the negative-binomial fit; no robust-Poisson fallback or "
            "regression failure occurred. Eight were positive."
        ),
    )
    nb = source_tables[12]
    nb[0][4] = "Positive among classifiable, % (95% Wilson confidence interval)"
    add_table_caption(
        document,
        "Appendix Table S7. ",
        (
            "Exploratory type C3 negative-binomial calibration result in the retained 1% daily audit sample. "
            "All 500 retained heterogeneous participants were attempted; 431 were ratio-applicable, 204 were "
            "regression-classifiable, and 8 were positive. “All-attempted rate” is 8/500. The 95% Wilson "
            "interval describes Monte Carlo uncertainty under this configured simulation. No robust-Poisson "
            "fallback was used."
        ),
    )
    add_table(
        document,
        nb[0],
        nb[1:],
        [1.2, 1.4, 1.55, 1.0, 2.05, 1.8],
        font_size=8.0,
        left_columns=(0,),
    )

    add_heading(document, "3.9 Mutually exclusive pattern decomposition", 2)
    add_paragraph(
        document,
        (
            "Type C1 denotes the perimenstrual pattern, type C2 the periovulatory pattern, and type C3 the "
            "inadequate-luteal-phase pattern. Table S8 separates their mutually exclusive combinations."
        ),
    )
    decomposition = deepcopy(source_tables[13])
    decomposition = [
        [
            "Outside C3 applicability" if value == "Not applicable" else value
            for value in row
        ]
        for row in decomposition
    ]
    add_table_caption(
        document,
        "Appendix Table S8. ",
        (
            "Mutually exclusive participant-window categories in 36-month strict-Herzog windows. Each cohort "
            "contains 50,000 attempted participants. “Windowed Herzog” uses the full 36-month diary. "
            "“Minimum-data” adds a minimum of four seizure days. Type C3 follows the heterogeneous-cohort "
            "applicability rule. Percentages may differ from 100% by rounding."
        ),
    )
    add_table(
        document,
        decomposition[0],
        decomposition[1:],
        [1.9, 0.8, 1.1, 0.8, 0.8, 0.8, 1.1, 1.7],
        font_size=7.8,
        left_columns=(0,),
    )

    add_heading(document, "3.10 Simulator-feature associations", 2)
    add_paragraph(
        document,
        (
            "The machine-readable file tableS4_simulated_classification_associations.csv reports stratified "
            "false-positive classification "
            f"rates and 95% Wilson score confidence intervals {marker('C034')} across configured quintiles "
            "of age, seizure burden, "
            "cycle length and variability, ovulatory fraction, and dominant seizure-cycle period, plus "
            "configured binary modifiers. These descriptive results apply to the synthetic study population."
        ),
    )

    add_heading(document, "3.11 Realized simulation distributions", 2)
    add_paragraph(
        document,
        (
            "The figures in this section show simulated data. Polycystic ovary syndrome (PCOS) and inadequate "
            "luteal phase (ILP) are written in abbreviated form within selected panel labels."
        ),
    )
    figure_specs = [
        (
            "figS1_seizure_process_distributions.png",
            "Supplementary Figure S1. ",
            "Realized seizure-process distributions. Seizures per month and seizure days per month use all 50,000 participants in each cohort. Daily count variance divided by mean and next-day clustering propensity use the retained 1% daily audit sample of 500 participants per cohort. All panels display simulated data; monthly measures use 30-day months.",
            7.8,
        ),
        (
            "figS2_seizure_rhythm_distributions.png",
            "Supplementary Figure S2. ",
            "Realized seizure-rhythm distributions. Generator-reported dominant period and latent monthly seizure burden use all 50,000 participants per cohort. Spectral period and amplitude use the retained 1% daily audit sample of 500 participants per cohort. Period is in days, spectral amplitude is in seizures per day, and burden is in seizures per 30-day month. All panels display simulated data.",
            7.8,
        ),
        (
            "figS3_menstrual_cycle_distributions.png",
            "Supplementary Figure S3. ",
            "Realized menstrual-cycle distributions in all 50,000 participants per cohort. Panels show mean cycle length in days, within-participant cycle-length standard deviation in days, ovulatory-cycle fraction, and inadequate luteal phase (ILP) fraction. The healthy cohort was generated with ovulation required. All panels display simulated data.",
            7.8,
        ),
        (
            "figS4_age_and_modifier_distributions.png",
            "Supplementary Figure S4. ",
            "Age and configured modifiers among all 50,000 participants per cohort. Polycystic ovary syndrome (PCOS) is abbreviated within the panel. Ages are in years. Modifier frequencies are investigator-configured heterogeneity stress-test settings and do not estimate population prevalence. All panels display simulated data.",
            8.2,
        ),
        (
            "figS5_simulated_classification_associations.png",
            "Supplementary Figure S5. ",
            "Associations with false-positive classification under the independent null model across within-cohort feature quintiles. Each cohort contains 50,000 attempted participants; individual points use the classifiable denominator for the displayed definition. Error bars are 95% Wilson score confidence intervals. The display is descriptive and applies to the synthetic study population.",
            8.2,
        ),
    ]
    for filename, label, caption, width in figure_specs:
        add_figure(document, SUPPLEMENT_DIR / filename, width, label, caption)

    new_section(document, "portrait")
    add_heading(document, "3.12 Interpretation of the supplementary analyses", 2)
    add_paragraph(
        document,
        (
            "The supplementary analyses quantify apparent catamenial classification under independence. "
            "They separate the perimenstrual (type C1), periovulatory (type C2), and inadequate-luteal-phase "
            "(type C3) patterns; preserve classifiable and all-attempted "
            "denominators; and show how diary duration, minimum-data requirements, ratio denominators, and inadequate-"
            "luteal-phase prevalence affect the estimates. Medication changes and secular trends are outside "
            "the simulation scope and remain limitations for translation to longitudinal clinical diaries."
        ),
    )

    add_heading(document, "References", 1)
    add_paragraph(document, "[[BIBLIOGRAPHY]]")

    # Ensure final section geometry and set update-fields-on-open.
    configure_section(document.sections[-1], "portrait")
    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    document.core_properties.title = (
        "Appendix S1 — False Positive Catamenial Epilepsy Classification"
    )
    document.core_properties.subject = (
        "HORMONE-CYCLE development, internal calibration, validation status, and supplementary analyses"
    )
    document.core_properties.comments = (
        "Built from draft_v6_appendix_s1_hormone_cycle_expanded; source preserved."
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    visible_text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    )
    used_codes = sorted(set(re.findall(r"\[\[CITE:(C\d+)\]\]", visible_text)))
    document.save(output)
    REVIEW_DIR.mkdir(parents=True, exist_ok=True)
    (REVIEW_DIR / "citation_map.json").write_text(
        json.dumps(
            {
                "document": str(output),
                "citation_style": "NLM citation-sequence",
                "collection": "CERES",
                "citations": {
                    marker(code): CITATION_SPECS[code] for code in used_codes
                },
            },
            indent=2,
        )
    )


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    build_document(args.source, args.output)
    print(
        json.dumps(
            {
                "source": str(args.source),
                "output": str(args.output),
                "citation_markers": len(CITATION_SPECS),
                "citation_map": str(REVIEW_DIR / "citation_map.json"),
            },
            indent=2,
        )
    )


if __name__ == "__main__":
    main()
